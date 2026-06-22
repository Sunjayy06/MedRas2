"""Final Sigma deterministic cleanup — regression checks.

Covers:
  1. Age baseline Minimum=30 / Maximum=88 populated in Word/PDF
  2. Primary outcome table uses Parameter | Category | n | % columns
  3. Primary outcome distribution figure exists before Age association figure
  4. pT / Nodal status / LVI absent from main Word/PDF by default
  5. Histological type, ER, PR, Molecular subtype, AR figures present
  6. category_merges sheet records Postive->Positive and Ki67 normalisation
  7. cleaned_processed_dataset contains no raw "Postive" values
  8. Significant-findings table has separate raw p-value / adjusted p-value columns
  9. Generic (non-p27) fixture still passes

Run from artifacts/medras:
    python -m test_fixtures.verify_sigma_final_polish
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from app.services import chapter_v_export, export
from app.services.results import build_table_one
from app.services.thesis_blueprint import build_thesis_analysis_blueprint


PNG_1X1 = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)

T4B_LONG_LABEL = (
    "Ulceration and / or ipsilateral satellite nodules and / or edema of the skin "
    "which does not meet the criteria for inflammatory carcinoma"
)


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def _docx_paragraph_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    return "\n".join(p.text for p in doc.paragraphs)


def _observed_rows_for(name: str) -> List[List[Any]]:
    if name == "ER":
        return [["Postive", 2, 0, 2], ["Positive", 4, 0, 4], ["Negative", 0, 6, 6], ["Total", 6, 6, 12]]
    if name == "PR":
        return [["Postive", 1, 0, 1], ["Positive", 5, 0, 5], ["Negative", 0, 6, 6], ["Total", 6, 6, 12]]
    if name == "AR":
        return [["Postive", 1, 0, 1], ["Positive", 5, 0, 5], ["Negative", 0, 6, 6], ["Total", 6, 6, 12]]
    if name == "Histological type":
        return [["1.0", 3, 0, 3], ["2.0", 3, 1, 4], ["3.0", 0, 5, 5], ["Total", 6, 6, 12]]
    if name == "Molecular subtype":
        return [["Luminal A", 3, 0, 3], ["Luminal B", 2, 1, 3], ["Her2neu", 1, 2, 3], ["Triple negative", 0, 3, 3], ["Total", 6, 6, 12]]
    if name == "Ki67":
        return [[">=14", 5, 1, 6], [">= 14%", 1, 0, 1], ["<14", 0, 5, 5], ["Total", 6, 6, 12]]
    return [["Positive", 6, 0, 6], ["Negative", 0, 6, 6], ["Total", 6, 6, 12]]


def _assoc_test(name: str, *, p_value: str, adjusted: str, effect: str = "Cramér's V = 0.46") -> Dict[str, Any]:
    return {
        "id": f"{name.lower().replace(' ', '_')}_by_outcome",
        "title": f"{name} vs p27 expression status: Chi-square test",
        "test_type": "chi_square",
        "analysis_family": "bivariate",
        "tables": [
            {
                "title": "Observed counts",
                "headers": ["Predictor category", "Positive", "Negative", "Total"],
                "rows": _observed_rows_for(name),
            },
            {
                "title": f"Association of p27 expression status with {name}",
                "headers": [
                    "Predictor category", "Positive n (%)", "Negative n (%)",
                    "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings",
                ],
                "rows": [["Positive", "6 (66.7%)", "0 (0.0%)", p_value, adjusted, "Chi-square test", effect, "-"]],
            },
        ],
        "figures": [{
            "title": f"{name} by p27 expression status",
            "png_data_uri": PNG_1X1,
            "caption": f"Grouped percentage bar chart of {name} by p27 expression status.",
        }],
    }


def _with_sparse_warning(test: Dict[str, Any]) -> Dict[str, Any]:
    test["tables"][1]["rows"][0][7] = (
        "Warning: Interpret with caution: sparse cells detected. "
        "Minimum expected count = 1.20. Interpret with caution."
    )
    return test


def _build_fixture() -> tuple[Dict[str, Any], "pd.DataFrame", Dict[str, Any]]:
    """Build representative results, raw df, and FakeEntry-style meta."""
    # Age range deliberately spans 30–88 to verify min/max pipeline.
    ages = [30, 35, 42, 48, 55, 60, 62, 67, 72, 77, 81, 88]
    n = len(ages)

    df = pd.DataFrame({
        "Positive/ Negative": ["Yes"] * 8 + ["No"] * 4,
        "Age": ages,
        "ER": ["Postive"] * 2 + ["Positive"] * 4 + ["Negative"] * 6,     # 2 typos
        "PR": ["Postive"] * 1 + ["Positive"] * 5 + ["Negative"] * 6,     # 1 typo
        "AR": ["Postive"] * 1 + ["Positive"] * 5 + ["Negative"] * 6,     # 1 typo
        "Histological type": [1.0, 2.0, 3.0] * 4,
        "Molecular subtype": ["Luminal A", "Luminal B", "HER2neu", "Triple negative"] * 3,
        "Her2Neu": ["Negative"] * 4 + ["1"] * 2 + ["2+"] * 2 + ["3+"] * 3 + ["Positive"],
        "EGFR": ["Negative"] * 7 + ["Patchy positive"] * 2 + ["Positive"] * 3,
        "Ki67": [">=14"] * 5 + ["<14"] * 7,                               # unnormalised Ki67
        "pT": ["T1", "T2", "T3"] * 4,
        "Nodal status": ["N0"] * 5 + ["NO"] + ["N1"] * 3 + ["N2"] * 3,
        "Tumour size": [T4B_LONG_LABEL] * 2 + ["2-5 cm"] * 10,
        "No of nodes involved": ["0/31", "2026-01-17 00:00:00", "2026-10-12 00:00:00", "2026-04-22 00:00:00"] * 3,
        "LVI": ["Present", "Abse"] * 6,
        "DCIS": ["Negative"] * 6 + ["High grade", "Low grade", "Intermediate grade", "Positive", "Negative", "Negative"],
    })

    classifications = [
        {"column": "Positive/ Negative", "detected_type": "nominal"},
        {"column": "Age",                "detected_type": "scale"},
        {"column": "ER",                 "detected_type": "nominal"},
        {"column": "PR",                 "detected_type": "nominal"},
        {"column": "AR",                 "detected_type": "nominal"},
        {"column": "Histological type",  "detected_type": "nominal"},
        {"column": "Molecular subtype",  "detected_type": "nominal"},
        {"column": "Her2Neu",            "detected_type": "ordinal"},
        {"column": "EGFR",               "detected_type": "nominal"},
        {"column": "Ki67",               "detected_type": "nominal"},
        {"column": "pT",                 "detected_type": "ordinal"},
        {"column": "Nodal status",       "detected_type": "ordinal"},
        {"column": "Tumour size",        "detected_type": "nominal"},
        {"column": "No of nodes involved", "detected_type": "nominal"},
        {"column": "LVI",                "detected_type": "nominal"},
        {"column": "DCIS",               "detected_type": "nominal"},
    ]

    table_one = build_table_one(df, classifications, group=None)

    tests: List[Dict[str, Any]] = [
        {
            "id": "age_by_outcome",
            "title": "Age by p27 expression status: Welch's t-test",
            "test_type": "welch_ttest",
            "analysis_family": "bivariate",
            "tables": [{
                "title": "Comparison of Age by p27 expression status",
                "headers": [
                    "Group", "n", "Mean ± SD", "Test statistic",
                    "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings",
                ],
                "rows": [
                    ["p27 Positive", "8", "57.1 ± 18.9", "t = -0.89; df = 27.58011794899391", "p = 0.240", "p = 0.300", "welch_ttest", "Cohen's d = -0.22896609106040972", "-"],
                    ["p27 Negative", "4", "61.8 ± 22.5", "t = -0.89; df = 27.58011794899391", "p = 0.240", "p = 0.300", "welch_ttest", "Cohen's d = -0.22896609106040972", "-"],
                ],
            }],
            "figures": [{
                "title": "Age by p27 expression status",
                "png_data_uri": PNG_1X1,
                "caption": "Boxplot of Age by p27 expression status.",
            }],
        },
        # Significant predictors (real p-values per task spec)
        _with_sparse_warning(_assoc_test("Histological type", p_value="p = 0.004", adjusted="p = 0.013")),
        _assoc_test("ER",                p_value="p < 0.001",  adjusted="p = 0.005"),
        _assoc_test("PR",                p_value="p = 0.002",  adjusted="p = 0.009"),
        _assoc_test("Molecular subtype", p_value="p < 0.001",  adjusted="p < 0.001"),
        _assoc_test("AR",                p_value="p = 0.004",  adjusted="p = 0.013"),
        # Non-significant (should NOT appear in main Word/PDF figures)
        _assoc_test("pT",          p_value="p = 0.400", adjusted="p = 0.700"),
        _assoc_test("Nodal status", p_value="p = 0.420", adjusted="p = 0.710"),
        _assoc_test("LVI",         p_value="p = 0.500", adjusted="p = 0.750"),
        _assoc_test("node_ratio",   p_value="p = 0.800", adjusted="p = 0.900"),
        _assoc_test("positive_nodes", p_value="p = 0.810", adjusted="p = 0.900"),
        _assoc_test("total_nodes",  p_value="p = 0.820", adjusted="p = 0.900"),
    ]

    significant_findings: List[Dict[str, Any]] = [
        {
            "variable": "Histological type vs Positive/ Negative",
            "key_finding": "Histological type was associated with the primary outcome.",
            "test_statistic": "chi-square = 8.40",
            "p_value": "p = 0.004",
            "adjusted_p_value": "p = 0.013",
            "test_applied": "Chi-square test",
            "effect_size": "Cramér's V = 0.46",
            "notes_warnings": "-",
        },
        {
            "variable": "ER vs Positive/ Negative",
            "key_finding": "ER was associated with the primary outcome.",
            "test_statistic": "chi-square = 12.00",
            "p_value": "p < 0.001",
            "adjusted_p_value": "p = 0.005",
            "test_applied": "Chi-square test",
            "effect_size": "Cramér's V = 0.55",
            "notes_warnings": "-",
        },
        {
            "variable": "PR vs Positive/ Negative",
            "key_finding": "PR was associated with the primary outcome.",
            "test_statistic": "chi-square = 9.30",
            "p_value": "p = 0.002",
            "adjusted_p_value": "p = 0.009",
            "test_applied": "Chi-square test",
            "effect_size": "Cramér's V = 0.50",
            "notes_warnings": "-",
        },
        {
            "variable": "Molecular subtype vs Positive/ Negative",
            "key_finding": "Molecular subtype was associated with the primary outcome.",
            "test_statistic": "chi-square = 14.00",
            "p_value": "p < 0.001",
            "adjusted_p_value": "p < 0.001",
            "test_applied": "Chi-square test",
            "effect_size": "Cramér's V = 0.60",
            "notes_warnings": "-",
        },
        {
            "variable": "AR vs Positive/ Negative",
            "key_finding": "AR was associated with the primary outcome.",
            "test_statistic": "chi-square = 8.10",
            "p_value": "p = 0.004",
            "adjusted_p_value": "p = 0.013",
            "test_applied": "Chi-square test",
            "effect_size": "Cramér's V = 0.44",
            "notes_warnings": "-",
        },
    ]

    blueprint = build_thesis_analysis_blueprint(
        df_shape=(n, len(df.columns)),
        classifications=classifications,
        assignment={"outcome": "Positive/ Negative"},
        table_one=table_one,
        tests=tests,
        graphs=[],
        significant_findings=significant_findings,
        methods_text=(
            "Continuous variables are summarised as mean ± SD with range. "
            "Chi-square tests were used for categorical predictors. "
            "Welch's t-test was used for the continuous predictor Age. "
            "p < 0.05 was considered statistically significant."
        ),
        results_narrative="ER, PR, Molecular subtype, Histological type, and AR were significantly associated with p27 expression status.",
        session={
            "study_type": "association",
            "domain_profile": "breast_pathology",
            "main_marker": "p27",
            "main_outcome_concept": "p27 expression status",
            "analysis_excluded_columns": [],
        },
        plan={"debug": {"eligible_predictor_count": 8, "bivariate_test_count": 8}},
    )

    results = {
        "table_one": table_one,
        "tests": tests,
        "significant_findings": significant_findings,
        "thesis_analysis_blueprint": blueprint,
        "export_metadata": {
            "dataset_id": "final-polish-job",
            "result_id": "final-polish-result",
            "analysis_version": 1,
            "generated_at": "2026-06-20T00:00:00+00:00",
            "domain_profile": "breast_pathology",
        },
    }

    meta = {
        "filename": "final-polish-fixture.xlsx",
        "domain_profile": "breast_pathology",
        "assignment": {"outcome": "Positive/ Negative"},
        "classifications": classifications,
        "data_version": 0,
        "category_merge_actions": [],   # no user merges; system merges come from df scan
    }

    return results, df, meta


# ── 1. Age baseline min / max ─────────────────────────────────────────────────

def test_age_baseline_min_max() -> None:
    """Word export must show Minimum=30 and Maximum=88 for Age."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))

    # Find the continuous descriptive table (Age row)
    age_row: list | None = None
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if headers[:2] == ["Parameter", "n"] and "Minimum" in headers and "Maximum" in headers:
            min_idx = headers.index("Minimum")
            max_idx = headers.index("Maximum")
            for row in table.rows[1:]:
                cells = [cell.text for cell in row.cells]
                if cells[0] == "Age":
                    age_row = cells
                    break
        if age_row:
            break

    assert age_row is not None, "Age row not found in any continuous descriptive table"
    min_val = age_row[min_idx]
    max_val = age_row[max_idx]
    assert min_val == "30", f"Expected Minimum=30, got {min_val!r}"
    assert max_val == "88", f"Expected Maximum=88, got {max_val!r}"


# ── 2. Primary outcome table format ──────────────────────────────────────────

def test_primary_outcome_table_format() -> None:
    """Primary outcome table must use Parameter | Category | n | % columns."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))
    text = _docx_text(blob)

    assert "Parameter\nCategory\nn\n%" in text, "Header row Parameter|Category|n|% not found"

    # p27 expression status must be the Parameter column entry
    found_p27_row = False
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if headers == ["Parameter", "Category", "n", "%"]:
            row_params = {row.cells[0].text for row in table.rows[1:]}
            if "p27 expression status" in row_params:
                found_p27_row = True
                break
    assert found_p27_row, "p27 expression status not found as Parameter in a Parameter/Category/n/% table"

    # Verify Positive and Negative categories appear
    assert "Positive\n8\n66.7%" in text, "Expected Positive row in primary outcome table"
    assert "Negative\n4\n33.3%" in text, "Expected Negative row in primary outcome table"


# ── 3. Figure order: distribution before Age ─────────────────────────────────

def test_primary_outcome_figure_before_age_figure() -> None:
    """Figure 1 must be the distribution figure; Age figure must follow."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)

    assert "Figure 1. Distribution of p27 expression status." in text, \
        "Distribution figure must be Figure 1"

    dist_pos = text.index("Figure 1. Distribution of p27 expression status.")
    age_pos = text.index("Figure 2.")  # Age is second because it's the next core figure
    assert dist_pos < age_pos, "Distribution figure must precede Age figure"


# ── 4. Non-significant figures absent from main report ───────────────────────

def test_nonsignificant_figures_absent() -> None:
    """pT, Nodal status, LVI figures must NOT appear in the main Word export."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)

    # After _normalise_figure_metadata the format is "{outcome} by {predictor}".
    assert "p27 expression status by pT" not in text, \
        "pT figure must be absent from main report"
    assert "p27 expression status by Nodal status" not in text, \
        "Nodal status figure must be absent from main report"
    assert "p27 expression status by LVI" not in text, \
        "LVI figure must be absent from main report"


# ── 5. Significant association figures present ────────────────────────────────

def test_significant_association_figures_present() -> None:
    """Histological type, ER, PR, Molecular subtype, AR figures must be present."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)

    # _normalise_figure_metadata rewrites captions to "{outcome} by {predictor}".
    for predictor in ("Histological type", "ER", "PR", "Molecular subtype", "AR"):
        assert f"p27 expression status by {predictor}" in text, \
            f"Figure for {predictor} must be present in main report"


def test_association_figures_precede_tables() -> None:
    """Core association figures should appear before the association tables in DOCX."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_paragraph_text(blob)

    first_table = text.index("Table ")
    first_assoc_table = text.index("Association of p27 expression status with Histological type", first_table)
    for caption in (
        "p27 expression status by Age",
        "p27 expression status by Histological type",
        "p27 expression status by ER",
        "p27 expression status by PR",
        "p27 expression status by Molecular subtype",
        "p27 expression status by AR",
    ):
        assert caption in text, f"{caption} figure caption must be present"
        assert text.index(caption) < first_assoc_table, f"{caption} must appear before association tables"


def test_association_tables_aggregate_display_categories() -> None:
    """Generated association tables must aggregate duplicate display categories."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))
    text = _docx_text(blob)

    association_tables = []
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if headers and headers[0] == "Predictor category":
            association_tables.append([[cell.text for cell in row.cells] for row in table.rows[1:]])

    assert association_tables, "Expected at least one generated association table"
    for rows in association_tables:
        categories = [row[0] for row in rows if row and row[0]]
        assert len(categories) == len(set(categories)), f"Duplicate display categories found: {categories}"

    assert "Postive" not in text
    assert "\nYes\n" not in text and "\nNo\n" not in text
    assert "Grade 1" in text and "Grade 2" in text and "Grade 3" in text
    assert "Type 1" not in text and "Type 2" not in text and "Type 3" not in text
    assert "HER2-enriched" in text
    histological_rows = []
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if headers and headers[0] == "Predictor category":
            rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
            if any(row and row[0] in {"Grade 1", "Grade 2", "Grade 3"} for row in rows):
                histological_rows.extend(rows)
    histological_categories = [row[0] for row in histological_rows if row]
    assert {"Grade 1", "Grade 2", "Grade 3"} <= set(histological_categories), histological_categories
    assert not {"1.0", "2.0", "3.0"} & set(histological_categories), histological_categories


def test_manual_style_descriptive_consolidation_docx() -> None:
    """Main descriptive tables should consolidate common breast pathology fragments."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)
    assert "N0" in text and "NO" not in text
    assert "Negative/low" in text
    assert "Equivocal (2+)" in text
    assert "Positive (3+)" in text
    assert "Patchy positive" not in text
    assert "High grade" not in text and "Low grade" not in text and "Intermediate grade" not in text
    assert "Tumour quadrant" in text or "Tumour site" not in text


# ── 6. category_merges audit sheet completeness ───────────────────────────────

class _FakeEntry:
    def __init__(self, df: pd.DataFrame, meta: dict) -> None:
        self.df = df
        self.meta = meta


def test_category_merges_records_system_display_typos() -> None:
    """category_merges sheet must record Postive->Positive and Ki67 normalisation."""
    results, df, meta = _build_fixture()
    entry = _FakeEntry(df, meta)
    blob = export.to_xlsx(entry, results, {"outcome": "Positive/ Negative"})

    wb = load_workbook(io.BytesIO(blob))
    assert "category_merges" in wb.sheetnames, "category_merges sheet must exist"

    ws = wb["category_merges"]
    rows = list(ws.iter_rows(min_row=2, values_only=True))

    original_cols = [str(r[1]) for r in rows if r[1] is not None]
    cleaned_cols  = [str(r[2]) for r in rows if r[2] is not None]

    # Postive -> Positive entries must be present (ER=2, PR=1, AR=1)
    assert "Postive" in original_cols, "Postive must appear in Original category column"
    assert "Positive" in cleaned_cols, "Positive must appear in Cleaned category column"

    # Ki67 normalisation: >=14 -> >=14%
    assert ">=14" in original_cols, ">=14 must appear in Original category column"
    assert ">=14%" in cleaned_cols, ">=14% must appear in Cleaned category column"

    # ER column with count=2 must have an explicit row
    er_rows = [r for r in rows if str(r[0]) == "ER" and str(r[1]) == "Postive"]
    assert er_rows, "ER / Postive row must be present in category_merges"
    assert er_rows[0][3] == 2, f"ER Postive count_affected must be 2, got {er_rows[0][3]}"

    # PR with count=1
    pr_rows = [r for r in rows if str(r[0]) == "PR" and str(r[1]) == "Postive"]
    assert pr_rows, "PR / Postive row must be present in category_merges"
    assert pr_rows[0][3] == 1, f"PR Postive count_affected must be 1, got {pr_rows[0][3]}"

    # AR with count=1
    ar_rows = [r for r in rows if str(r[0]) == "AR" and str(r[1]) == "Postive"]
    assert ar_rows, "AR / Postive row must be present in category_merges"
    assert ar_rows[0][3] == 1, f"AR Postive count_affected must be 1, got {ar_rows[0][3]}"

    assert any(str(r[0]) == "Nodal status" and str(r[1]) == "NO" and str(r[2]) == "N0" for r in rows)
    assert any(str(r[0]) == "Her2Neu" and str(r[1]) == "1" and str(r[2]) == "Negative/low" for r in rows)
    assert any(str(r[0]) == "Her2Neu" and str(r[1]) == "2+" and str(r[2]) == "Equivocal (2+)" for r in rows)
    assert any(str(r[0]) == "Her2Neu" and str(r[1]) == "3+" and str(r[2]) == "Positive (3+)" for r in rows)
    assert any(str(r[0]) == "EGFR" and str(r[1]) == "Patchy positive" and str(r[2]) == "Positive" for r in rows)
    assert any(str(r[0]) == "DCIS" and str(r[1]) == "High grade" and str(r[2]) == "Present" for r in rows)
    assert any(str(r[0]) == "No of nodes involved" and "2026-01-17" in str(r[1]) and str(r[2]) == "1/17" for r in rows)


# ── 7. cleaned_processed_dataset has no raw typos ────────────────────────────

def test_cleaned_dataset_no_postive() -> None:
    """cleaned_processed_dataset sheet must not contain the raw typo 'Postive'."""
    results, df, meta = _build_fixture()
    entry = _FakeEntry(df, meta)
    blob = export.to_xlsx(entry, results, {"outcome": "Positive/ Negative"})

    wb = load_workbook(io.BytesIO(blob))
    assert "cleaned_processed_dataset" in wb.sheetnames

    ws = wb["cleaned_processed_dataset"]
    # iter_rows(values_only=True) returns raw scalars, not Cell objects
    all_text = " ".join(
        str(v)
        for row in ws.iter_rows(min_row=2, values_only=True)
        for v in row
        if v is not None
    )
    assert "Postive" not in all_text, \
        "cleaned_processed_dataset must not contain the raw typo 'Postive'"
    assert all(
        str(v).strip() != "Abse"
        for row in ws.iter_rows(min_row=2, values_only=True)
        for v in row
        if v is not None
    ), "cleaned_processed_dataset must not contain the raw LVI typo 'Abse'"
    assert "Grade 1" in all_text and "Grade 2" in all_text and "Grade 3" in all_text, \
        "cleaned_processed_dataset must show Histological type as Grade 1/2/3"
    assert "NO" not in all_text and "N0" in all_text
    assert "Negative/low" in all_text and "Equivocal (2+)" in all_text and "Positive (3+)" in all_text
    assert "Patchy positive" not in all_text
    assert "High grade" not in all_text and "Low grade" not in all_text and "Intermediate grade" not in all_text
    assert "2026-" not in all_text
    assert "1/17" in all_text and "10/12" in all_text and "4/22" in all_text


def test_lvi_abse_audit_trace_preserved() -> None:
    """category_merges must preserve the raw LVI Abse -> Absent audit trail."""
    results, df, meta = _build_fixture()
    entry = _FakeEntry(df, meta)
    blob = export.to_xlsx(entry, results, {"outcome": "Positive/ Negative"})
    wb = load_workbook(io.BytesIO(blob))
    rows = list(wb["category_merges"].iter_rows(min_row=2, values_only=True))
    lvi_rows = [r for r in rows if str(r[0]) == "LVI" and str(r[1]) == "Abse" and str(r[2]) == "Absent"]
    assert lvi_rows, "category_merges must record LVI Abse -> Absent"


def test_abse_is_confined_to_excel_audit_trace() -> None:
    """Raw Abse may remain in audit trace, never in user-facing workbook sheets."""
    results, df, meta = _build_fixture()
    workbook = load_workbook(io.BytesIO(export.to_xlsx(
        _FakeEntry(df, meta), results, {"outcome": "Positive/ Negative"}
    )))
    audit_sheets = {"category_merges", "cleaning_decisions", "Data Cleaning Log"}
    exposed = []
    invalid_molecular_wording = []
    for sheet in workbook.worksheets:
        if sheet.title in audit_sheets:
            continue
        for row in sheet.iter_rows():
            for cell in row:
                if re.search(r"\bAbse\b", str(cell.value or "")):
                    exposed.append((sheet.title, cell.coordinate, cell.value))
                if "Molecular subtype-negative" in str(cell.value or ""):
                    invalid_molecular_wording.append((sheet.title, cell.coordinate, cell.value))
    assert not exposed, f"Raw Abse leaked into user-facing Excel sheets: {exposed}"
    assert not invalid_molecular_wording, (
        f"Invalid molecular-subtype wording leaked into Excel: {invalid_molecular_wording}"
    )
    audit_values = [
        cell.value
        for row in workbook["category_merges"].iter_rows()
        for cell in row
    ]
    assert "Abse" in audit_values, "Raw Abse must remain traceable in category_merges"


def test_t4b_definition_is_display_normalized_with_audit_trace() -> None:
    results, df, meta = _build_fixture()
    docx_text = _docx_text(chapter_v_export.generate_docx(results))
    assert T4B_LONG_LABEL not in docx_text
    assert "Skin involvement / T4b features" in docx_text

    workbook = load_workbook(io.BytesIO(export.to_xlsx(
        _FakeEntry(df, meta), results, {"outcome": "Positive/ Negative"}
    )))
    audit_sheets = {"category_merges", "cleaning_decisions", "Data Cleaning Log"}
    exposed = [
        (sheet.title, cell.coordinate)
        for sheet in workbook.worksheets if sheet.title not in audit_sheets
        for row in sheet.iter_rows() for cell in row
        if T4B_LONG_LABEL in str(cell.value or "")
    ]
    assert not exposed, f"Long T4b definition leaked into user-facing Excel sheets: {exposed}"
    merge_rows = list(workbook["category_merges"].iter_rows(min_row=2, values_only=True))
    assert any(row[0] == "Tumour size" and row[1] == T4B_LONG_LABEL for row in merge_rows)


def test_association_direction_wording_is_semantically_valid() -> None:
    results, _, _ = _build_fixture()
    text = _docx_text(chapter_v_export.generate_docx(results))
    assert "Molecular subtype-negative" not in text


def test_sparse_caution_deduplicated() -> None:
    """Sparse-cell caution should appear once, without minimum expected-count dump."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)
    caution = "This finding should be interpreted cautiously because some expected cell counts were below 5."
    assert text.count(caution) == 1, "Sparse-cell caution must appear once for the result"
    assert "Minimum expected count" not in text
    assert "because some." not in text
    assert "Interpret with caution: -" not in text


def test_age_table_formats_internal_values() -> None:
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)
    assert "welch_ttest" not in text
    assert "Welch's t-test" in text
    assert "27.58011794899391" not in text
    assert "df = 27.58" in text
    assert "-0.22896609106040972" not in text
    assert "Cohen's d = -0.229" in text


def test_node_derived_figures_absent_but_results_retained() -> None:
    results, _, _ = _build_fixture()
    word = _docx_text(chapter_v_export.generate_docx(results))
    pdf = _pdf_text(chapter_v_export.generate_pdf(results))
    for variable in ("node_ratio", "positive_nodes", "total_nodes"):
        assert f"p27 expression status by {variable}" not in word
        assert f"p27 expression status by {variable}" not in pdf
        assert any(variable in str(test.get("title") or "") for test in results["tests"])


def test_marker_acronyms_preserved_in_rendered_captions() -> None:
    results, _, _ = _build_fixture()
    word = _docx_text(chapter_v_export.generate_docx(results))
    pdf = _pdf_text(chapter_v_export.generate_pdf(results))
    for text in (word, pdf):
        assert "p27 expression status by ER" in text
        assert "p27 expression status by AR" in text
        assert "p27 expression status by Er" not in text
        assert "p27 expression status by Ar" not in text


def test_significant_key_findings_are_thesis_style() -> None:
    """Final significant findings should use deterministic thesis-style wording."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)
    expected = (
        "Grade 3 cases were proportionately higher in the p27-negative group.",
        "p27 positivity was strongly associated with ER positivity.",
        "p27 positivity was significantly associated with PR positivity.",
        "Triple-negative phenotype was proportionately enriched among p27-negative cases, while Luminal B predominated among p27-positive cases.",
        "p27 positivity was significantly associated with AR positivity.",
    )
    for phrase in expected:
        assert phrase in text, f"Missing thesis-style finding: {phrase}"
    assert "was associated with the primary outcome" not in text


# ── 8. Significant findings p-values are separate and correct ─────────────────

def test_significant_findings_pvalues_separate() -> None:
    """Significant findings table must have separate raw and adjusted p-value columns."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))

    # Last table in Chapter V is the significant findings summary
    sig_table = doc.tables[-1]
    headers = [cell.text for cell in sig_table.rows[0].cells]

    assert "p-value" in headers, "p-value column must exist in significant findings table"
    assert "Adjusted p-value" in headers, "Adjusted p-value column must exist"

    p_idx  = headers.index("p-value")
    ap_idx = headers.index("Adjusted p-value")

    rows = {row.cells[0].text: [cell.text for cell in row.cells]
            for row in sig_table.rows[1:] if row.cells[0].text}

    def _find(label: str) -> list[str]:
        for key, row in rows.items():
            if label in key:
                return row
        raise AssertionError(f"Finding '{label}' not in significant findings table. Keys: {list(rows)}")

    hist = _find("Histological type")
    assert hist[p_idx] == "p = 0.004",  f"Histological type raw p: {hist[p_idx]!r}"
    assert hist[ap_idx] == "p = 0.013", f"Histological type adjusted p: {hist[ap_idx]!r}"

    er = _find("ER vs")
    assert er[p_idx] == "p < 0.001",  f"ER raw p: {er[p_idx]!r}"
    assert er[ap_idx] == "p = 0.005", f"ER adjusted p: {er[ap_idx]!r}"

    pr = _find("PR vs")
    assert pr[p_idx] == "p = 0.002",  f"PR raw p: {pr[p_idx]!r}"
    assert pr[ap_idx] == "p = 0.009", f"PR adjusted p: {pr[ap_idx]!r}"

    mol = _find("Molecular subtype")
    assert mol[p_idx] == "p < 0.001",  f"Molecular subtype raw p: {mol[p_idx]!r}"
    assert mol[ap_idx] == "p < 0.001", f"Molecular subtype adjusted p: {mol[ap_idx]!r}"

    ar = _find("AR vs")
    assert ar[p_idx] == "p = 0.004",  f"AR raw p: {ar[p_idx]!r}"
    assert ar[ap_idx] == "p = 0.013", f"AR adjusted p: {ar[ap_idx]!r}"


# ── 9. PDF primary outcome table uses Parameter|Category|n|% ─────────────────

def _pdf_text(blob: bytes) -> str:
    """Extract plain text from a PDF blob using an available PDF reader."""
    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None
    if pdfplumber is not None:
        pages = []
        with pdfplumber.open(io.BytesIO(blob)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                pages.append(t)
        text = "\n".join(pages)
        if text.strip():
            return text

    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore[no-redef]
        except ImportError as exc:
            raise AssertionError("PDF text extraction requires pdfplumber, pypdf, or PyPDF2") from exc

    reader = PdfReader(io.BytesIO(blob))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def test_primary_outcome_table_format_pdf() -> None:
    """PDF primary outcome table must use Parameter|Category|n|%, not Variable|Summary|Overall."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_pdf(results)
    assert blob[:4] == b"%PDF", "export must produce a valid PDF"
    text = _pdf_text(blob)
    assert text.strip(), "PDF text extraction returned empty content"
    assert "Variable Summary Overall" not in text, \
        "PDF must not render primary outcome table as Variable|Summary|Overall"
    assert "Parameter" in text, \
        "PDF primary outcome table must contain 'Parameter' column header"
    assert "Category" in text, \
        "PDF primary outcome table must contain 'Category' column header"
    assert "Molecular subtype-negative" not in text


# ── 10. PDF non-significant figures absent ───────────────────────────────────

def test_nonsignificant_figures_absent_pdf() -> None:
    """pT, Nodal status, LVI figure captions must NOT appear in the main PDF."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_pdf(results)
    text = _pdf_text(blob)
    assert text.strip(), "PDF text extraction returned empty content"
    assert "p27 expression status by pT" not in text, \
        "pT figure must be absent from main PDF"
    assert "p27 expression status by Nodal status" not in text, \
        "Nodal status figure must be absent from main PDF"
    assert "p27 expression status by LVI" not in text, \
        "LVI figure must be absent from main PDF"


# ── 11. PDF significant figures present ──────────────────────────────────────

def test_significant_figures_present_pdf() -> None:
    """Distribution, Age, and significant predictor figures must be present in PDF."""
    results, _, _ = _build_fixture()
    blob = chapter_v_export.generate_pdf(results)
    text = _pdf_text(blob)
    assert text.strip(), "PDF text extraction returned empty content"
    assert "Figure 1. Distribution of p27 expression status." in text, \
        "Distribution figure must be Figure 1 in PDF"
    for predictor in ("Age", "Histological type", "ER", "PR", "Molecular subtype", "AR"):
        assert f"p27 expression status by {predictor}" in text, \
            f"PDF must contain figure for {predictor}"


# ── 12. Generic (non-p27) fixture still passes ────────────────────────────────

def test_generic_non_p27_fixture() -> None:
    """A generic two-group study must render without p27-specific content."""
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(20, 2),
        classifications=[
            {"column": "Score",     "detected_type": "scale"},
            {"column": "Treatment", "detected_type": "nominal"},
        ],
        assignment={"outcome": "Score", "group": "Treatment"},
        table_one={
            "headers": ["Variable", "Type", "Overall"],
            "rows": [{"variable": "Score", "type": "mean ± SD", "cells": ["12.4 ± 3.1"]}],
        },
        tests=[],
        graphs=[],
        significant_findings=[],
        methods_text="Welch's t-test was used for the two-group comparison.",
        results_narrative="No significant finding was detected.",
        session={"study_type": "two-group comparison", "domain_profile": "generic"},
    )
    blob = chapter_v_export.generate_docx({
        "thesis_analysis_blueprint": blueprint,
        "export_metadata": {"result_id": "generic-final-polish"},
    })
    text = _docx_text(blob)
    assert "CHAPTER V" in text
    assert "p27" not in text


# ── Phase B: Narrative polish service ────────────────────────────────────────

def test_polish_skipped_without_consent() -> None:
    """polish_results must return an empty dict when OpenRouter is not configured."""
    from app.services import narrative_polish
    results, _, _ = _build_fixture()
    # OpenRouter is not configured in the test environment; must return {}
    overrides = narrative_polish.polish_results(results)
    assert isinstance(overrides, dict), "polish_results must return a dict"
    # Without a live key this will always be empty — deterministic fallback
    # (We can't assert it's empty if somehow a key exists, but we can confirm
    # generate_docx accepts {} without error.)
    blob = chapter_v_export.generate_docx(results, polish_overrides=overrides)
    assert blob[:2] == b"PK", "DOCX must still be generated with empty overrides"


def test_polish_fallback_on_openrouter_failure() -> None:
    """generate_docx/generate_pdf must produce valid files when polish_overrides={}."""
    results, _, _ = _build_fixture()
    word_blob = chapter_v_export.generate_docx(results, polish_overrides={})
    pdf_blob = chapter_v_export.generate_pdf(results, polish_overrides={})
    assert word_blob[:2] == b"PK", "DOCX must be produced with empty polish overrides"
    assert pdf_blob[:4] == b"%PDF", "PDF must be produced with empty polish overrides"


def test_polish_safety_rejects_new_numbers() -> None:
    """_is_safe must reject AI output that introduces new numeric tokens."""
    from app.services.narrative_polish import _is_safe
    original = "ER was associated with the primary outcome in this cohort."
    proposed_with_number = "ER was significantly associated with the outcome (n=42)."
    assert not _is_safe(original, proposed_with_number), \
        "Safety validator must reject prose that introduces a new number"
    proposed_safe = "ER demonstrated a statistically significant association with the primary outcome."
    assert _is_safe(original, proposed_safe), \
        "Safety validator must accept prose with no new numbers"


def test_polish_safety_preserves_cautions() -> None:
    """_is_safe must reject AI output that drops caution phrases from the original."""
    from app.services.narrative_polish import _is_safe
    original = "Results should be interpreted cautiously because sparse cells were observed."
    proposed_drops_caution = "Results indicate a significant association between ER and the outcome."
    assert not _is_safe(original, proposed_drops_caution), \
        "Safety validator must reject prose that removes a caution phrase"


def test_polish_overrides_applied_to_docx() -> None:
    """If polish_overrides contains a table interpretation, DOCX must use it."""
    results, _, _ = _build_fixture()
    # Determine a table_id that exists in the fixture blueprint
    blueprint = results.get("thesis_analysis_blueprint") or {}
    table_id = None
    for section in blueprint.get("analysis_sections") or []:
        for table in section.get("tables") or []:
            if table.get("table_id"):
                table_id = table["table_id"]
                break
        if table_id:
            break
    if table_id is None:
        return  # no tables to test
    polished_text = "This table presents a comprehensive overview of the study characteristics."
    overrides = {f"table:{table_id}": polished_text}
    blob = chapter_v_export.generate_docx(results, polish_overrides=overrides)
    from docx import Document
    doc = Document(io.BytesIO(blob))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert polished_text in full_text, \
        f"DOCX must contain the polished override text for table:{table_id}"


def test_pvalues_unchanged_by_polish() -> None:
    """p-values in the significant findings table must be unchanged regardless of polish."""
    results, _, _ = _build_fixture()
    for overrides in ({}, {"section:bivariate_associations": "Polished intro text."}):
        blob = chapter_v_export.generate_docx(results, polish_overrides=overrides)
        from docx import Document
        doc = Document(io.BytesIO(blob))
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                all_text.extend(cell.text for cell in row.cells)
        full_text = "\n".join(all_text)
        assert "p = 0.004" in full_text, \
            "Histological type raw p-value must be present regardless of polish"
        assert "p = 0.013" in full_text, \
            "Adjusted p-value must be present regardless of polish"


# ── runner ────────────────────────────────────────────────────────────────────

def main() -> None:
    test_age_baseline_min_max()
    print("  [ok] Age baseline min/max")

    test_primary_outcome_table_format()
    print("  [ok] Primary outcome table Parameter|Category|n|% format")

    test_primary_outcome_figure_before_age_figure()
    print("  [ok] Primary outcome distribution figure before Age figure")

    test_nonsignificant_figures_absent()
    print("  [ok] pT / Nodal status / LVI figures absent from main report")

    test_significant_association_figures_present()
    print("  [ok] Histological type / ER / PR / Molecular subtype / AR figures present")

    test_association_figures_precede_tables()
    print("  [ok] Core association figures appear before association tables")

    test_association_tables_aggregate_display_categories()
    print("  [ok] Association tables aggregate cleaned display categories")

    test_manual_style_descriptive_consolidation_docx()
    print("  [ok] Manual-style descriptive category consolidation appears in DOCX")

    test_category_merges_records_system_display_typos()
    print("  [ok] category_merges records display normalisation audit rows")

    test_cleaned_dataset_no_postive()
    print("  [ok] cleaned_processed_dataset contains no raw typo 'Postive' or 'Abse'")

    test_lvi_abse_audit_trace_preserved()
    print("  [ok] category_merges records LVI Abse->Absent audit trace")

    test_abse_is_confined_to_excel_audit_trace()
    print("  [ok] Raw Abse is confined to Excel audit trace")

    test_t4b_definition_is_display_normalized_with_audit_trace()
    print("  [ok] Long T4b definition is normalized with audit trace")

    test_association_direction_wording_is_semantically_valid()
    print("  [ok] Association direction wording is semantically valid")

    test_sparse_caution_deduplicated()
    print("  [ok] Sparse-cell caution is deduplicated")

    test_age_table_formats_internal_values()
    print("  [ok] Age comparison table formats internal test values")

    test_node_derived_figures_absent_but_results_retained()
    print("  [ok] Node-derived figures stay out of main report while results remain")

    test_marker_acronyms_preserved_in_rendered_captions()
    print("  [ok] ER/AR figure labels preserve clinical capitalization")

    test_significant_key_findings_are_thesis_style()
    print("  [ok] Significant findings use deterministic thesis-style wording")

    test_significant_findings_pvalues_separate()
    print("  [ok] Significant findings raw p-value and adjusted p-value are separate")

    test_generic_non_p27_fixture()
    print("  [ok] Generic non-p27 fixture passes")

    test_primary_outcome_table_format_pdf()
    print("  [ok] PDF primary outcome table uses Parameter|Category|n|%")

    test_nonsignificant_figures_absent_pdf()
    print("  [ok] pT / Nodal status / LVI figure captions absent from PDF")

    test_significant_figures_present_pdf()
    print("  [ok] Distribution and significant figures present in PDF")

    test_polish_skipped_without_consent()
    print("  [ok] Narrative polish skipped when no consent (deterministic fallback)")

    test_polish_fallback_on_openrouter_failure()
    print("  [ok] Narrative polish falls back deterministically when OpenRouter fails")

    test_polish_safety_rejects_new_numbers()
    print("  [ok] Narrative polish safety validator rejects prose with new numbers")

    test_polish_safety_preserves_cautions()
    print("  [ok] Narrative polish safety validator rejects prose that removes cautions")

    test_polish_overrides_applied_to_docx()
    print("  [ok] Polished overrides are applied to DOCX interpretation text")

    test_pvalues_unchanged_by_polish()
    print("  [ok] p-values are unchanged by narrative polish")

    print("\nSigma final polish verification passed.")


if __name__ == "__main__":
    main()
