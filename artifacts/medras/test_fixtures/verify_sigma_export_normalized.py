from __future__ import annotations

import io
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.services.export import to_docx, to_pdf, to_xlsx  # noqa: E402


PNG_1X1 = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


class FakeEntry:
    def __init__(self) -> None:
        self.df = pd.DataFrame({
            "value": [1.0, 2.0, 3.0, 4.0],
            "group": ["A", "A", "B", "B"],
        })
        self.meta = {
            "filename": "normalized-export-fixture.csv",
            "classifications": [
                {"column": "value", "detected_type": "scale"},
                {"column": "group", "detected_type": "nominal"},
            ],
            "normality": {"columns": []},
        }


def _fixture_results() -> dict:
    return {
        "table_one": {
            "headers": ["Variable", "Summary", "Overall"],
            "rows": [
                {"variable": "value", "type": "Mean +/- SD", "cells": ["2.50 +/- 1.29"]},
            ],
        },
        "tests": [
            {
                "id": "normalized_ttest",
                "title": "Normalized t-test",
                "plan_name": "Normalized t-test",
                "test_type": "t_test_independent",
                "tables": [
                    {
                        "title": "Normalized Summary",
                        "headers": ["Statistic", "Value"],
                        "rows": [
                            ["Mean difference", "2.00"],
                            ["p-value", "0.010"],
                        ],
                    }
                ],
                "figures": [
                    {"title": "Tiny Normalized Figure", "png_data_uri": PNG_1X1}
                ],
                "rows": [],
                "narrative": "Normalized narrative is preserved.",
            },
            {
                "id": "legacy_kappa",
                "title": "Legacy Kappa",
                "plan_name": "Legacy Kappa",
                "test_type": "kappa",
                "rows": [{"label": "Legacy statistic", "value": "0.80"}],
                "kappa": 0.8,
                "narrative": "Legacy fallback narrative is preserved.",
            },
        ],
        "graphs": [],
    }


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main() -> None:
    entry = FakeEntry()
    results = _fixture_results()
    assignment = {"outcome": "value", "group": "group", "covariates": []}

    docx_blob = to_docx(entry, results, assignment)
    docx_text = _docx_text(docx_blob)
    assert "Normalized Summary" in docx_text
    assert "Mean difference" in docx_text
    assert "Tiny Normalized Figure" in docx_text
    assert "Legacy Kappa" in docx_text

    pdf_blob = to_pdf(entry, results, assignment)
    assert pdf_blob.startswith(b"%PDF")
    assert len(pdf_blob) > 1000

    xlsx_blob = to_xlsx(entry, results, assignment)
    wb = load_workbook(io.BytesIO(xlsx_blob))
    assert "Normalized t-test" in wb.sheetnames
    norm_values = [
        cell
        for row in wb["Normalized t-test"].iter_rows(values_only=True)
        for cell in row
        if cell is not None
    ]
    assert "Normalized Summary" in norm_values
    assert "Mean difference" in norm_values
    assert "Tiny Normalized Figure" in norm_values

    assert "Legacy Kappa" in wb.sheetnames
    legacy_values = [
        cell
        for row in wb["Legacy Kappa"].iter_rows(values_only=True)
        for cell in row
        if cell is not None
    ]
    assert "Legacy statistic" in legacy_values
    assert "0.80" in legacy_values

    print("Sigma normalized export verification passed.")


if __name__ == "__main__":
    main()
