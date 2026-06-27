"""Verify Sigma Chapter V export is driven by thesis_analysis_blueprint.

Run:
    python -m test_fixtures.verify_sigma_chapter_v_export
"""

from __future__ import annotations

import asyncio
import io
from pathlib import Path
from unittest.mock import patch

import pandas as pd
from docx import Document

from app.api import stats
from app.services import chapter_v_export, export, narrative_polish
from app.services.thesis_blueprint import build_thesis_analysis_blueprint


PNG_1X1 = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


class FakeEntry:
    def __init__(self):
        self.df = pd.DataFrame({
            "Positive/ Negative": ["Yes", "No"] * 6,
            "Age": [42, 44, 51, 53, 55, 56, 60, 61, 63, 65, 67, 69],
            "ER": ["Positive", "Negative"] * 6,
            "PR": ["Positive", "Negative"] * 6,
            "AR": ["Positive", "Negative"] * 6,
            "Histological type": ["Type 1", "Type 2", "Type 3"] * 4,
            "Molecular subtype": ["Luminal A", "Luminal B", "HER2neu", "Triple negative"] * 3,
            "pT": ["T1", "T2", "T3"] * 4,
            "Nodal status": ["N0", "N1", "N2"] * 4,
            "LVI": ["Present", "Absent"] * 6,
            "Interpretation-site": ["Nuclear", "Cytoplasmic"] * 6,
            "Staining Result": ["Strong", "Weak"] * 6,
        })
        self.meta = {
            "filename": "chapter-v-fixture.xlsx",
            "domain_profile": "breast_pathology",
            "assignment": {"outcome": "Positive/ Negative"},
            "classifications": [
                {"column": "Positive/ Negative", "detected_type": "nominal"},
                {"column": "Age", "detected_type": "scale"},
                {"column": "ER", "detected_type": "nominal"},
                {"column": "PR", "detected_type": "nominal"},
                {"column": "AR", "detected_type": "nominal"},
                {"column": "Histological type", "detected_type": "nominal"},
                {"column": "Molecular subtype", "detected_type": "nominal"},
                {"column": "pT", "detected_type": "ordinal"},
                {"column": "Nodal status", "detected_type": "ordinal"},
                {"column": "LVI", "detected_type": "nominal"},
                {"column": "Interpretation-site", "detected_type": "nominal"},
                {"column": "Staining Result", "detected_type": "ordinal"},
            ],
            "data_version": 0,
        }


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def _representative_results() -> dict:
    classifications = [
        {"column": "Positive/ Negative", "detected_type": "nominal"},
        {"column": "Age", "detected_type": "scale"},
        {"column": "ER", "detected_type": "nominal"},
        {"column": "PR", "detected_type": "nominal"},
        {"column": "AR", "detected_type": "nominal"},
        {"column": "Histological type", "detected_type": "nominal"},
        {"column": "Molecular subtype", "detected_type": "nominal"},
        {"column": "pT", "detected_type": "ordinal"},
        {"column": "Nodal status", "detected_type": "ordinal"},
        {"column": "LVI", "detected_type": "nominal"},
        {"column": "Interpretation-site", "detected_type": "nominal"},
        {"column": "Staining Result", "detected_type": "ordinal"},
    ]
    table_one = {
        "headers": ["Variable", "Type", "Overall"],
        "rows": [
            {"variable": "Positive/ Negative", "type": "n (%)", "cells": ["Yes: 9 (75.0%); No: 3 (25.0%)"]},
            {"variable": "Age", "type": "mean ± SD", "cells": ["57.2 ± 8.6"]},
            {"variable": "ER", "type": "n (%)", "cells": ["Postive: 6 (50.0%); Negative: 6 (50.0%)"]},
            {"variable": "PR", "type": "n (%)", "cells": ["Positive: 6 (50.0%); Negative: 6 (50.0%)"]},
            {"variable": "AR", "type": "n (%)", "cells": ["Positive: 6 (50.0%); Negative: 6 (50.0%)"]},
            {"variable": "Histological type", "type": "n (%)", "cells": ["Type 1: 4 (33.3%); Type 2: 4 (33.3%); Type 3: 4 (33.3%)"]},
            {"variable": "Molecular subtype", "type": "n (%)", "cells": ["Luminal A: 3 (25.0%); Luminal B: 3 (25.0%); HER2neu: 3 (25.0%); Triple negative: 3 (25.0%)"]},
            {"variable": "pT", "type": "n (%)", "cells": ["T1: 4 (33.3%); T2: 4 (33.3%); T3: 4 (33.3%)"]},
            {"variable": "Nodal status", "type": "n (%)", "cells": ["N0: 4 (33.3%); N1: 4 (33.3%); N2: 4 (33.3%)"]},
            {"variable": "LVI", "type": "n (%)", "cells": ["Present: 6 (50.0%); Absent: 6 (50.0%)"]},
            {"variable": "Interpretation-site", "type": "n (%)", "cells": ["Nuclear: 6 (50.0%); Cytoplasmic: 6 (50.0%)"]},
            {"variable": "Staining Result", "type": "n (%)", "cells": ["Strong: 6 (50.0%); Weak: 6 (50.0%)"]},
        ],
    }
    for row in table_one["rows"]:
        if row.get("variable") == "Age":
            row["type"] = "Mean ± SD"
            row["cells"] = ["57.2 ± 8.6 (n=12); Missing: 0 (0.0%)"]
        if row.get("variable") == "ER":
            row["cells"] = ["Postive: 6 (50.0%); Positive: 2 (16.7%); Negative: 4 (33.3%)"]

    def assoc_test(name: str, *, p_value: str, adjusted: str, effect: str = "Cramer's V = 0.46") -> dict:
        return {
            "id": f"{name.lower().replace(' ', '_')}_by_outcome",
            "title": f"{name} vs p27 expression status: Chi-square test",
            "test_type": "chi_square",
            "analysis_family": "bivariate",
            "tables": [{
                "title": f"Association of p27 expression status with {name}",
                "headers": ["Predictor category", "Positive n (%)", "Negative n (%)", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
                "rows": [["Positive", "6 (66.7%)", "0 (0.0%)", p_value, adjusted, "Chi-square test", effect, "-"]],
            }],
            "figures": [{
                "title": f"{name} by p27 expression status",
                "png_data_uri": PNG_1X1,
                "caption": f"Grouped percentage bar chart of {name} by p27 expression status.",
            }],
        }

    tests = [
        {
            "id": "age_by_outcome",
            "title": "Age by p27 expression status: Welch's t-test",
            "test_type": "welch_ttest",
            "analysis_family": "bivariate",
            "tables": [{
                "title": "Comparison of Age by p27 expression status",
                "headers": ["Group", "n", "Mean ± SD", "Test statistic", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
                "rows": [["Positive", "9", "56.1 ± 8.0", "t = 1.20", "p = 0.240", "p = 0.300", "Welch's t-test", "Cohen's d = 0.42", "-"]],
            }],
            "figures": [],
        },
        {
            "id": "er_by_outcome",
            "title": "ER vs p27 expression status: Chi-square test",
            "test_type": "chi_square",
            "analysis_family": "bivariate",
            "tables": [{
                "title": "Association of p27 expression status with ER",
                "headers": ["Predictor category", "Positive n (%)", "Negative n (%)", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
                "rows": [["Positive", "6 (66.7%)", "0 (0.0%)", "p = 0.005", "p = 0.018", "Chi-square test", "Cramer's V = 0.46", "-"]],
            }],
            "figures": [{
                "title": "ER by p27 expression status",
                "png_data_uri": PNG_1X1,
                "caption": "Grouped percentage bar chart of ER by p27 expression status.",
            }],
        },
    ]
    tests = [
        {
            "id": "age_by_outcome",
            "title": "Age by p27 expression status: Welch's t-test",
            "test_type": "welch_ttest",
            "analysis_family": "bivariate",
            "tables": [{
                "title": "Comparison of Age by p27 expression status",
                "headers": ["Group", "n", "Mean Â± SD", "Test statistic", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
                "rows": [
                    ["p27 Positive", "9", "56.1 Â± 8.0", "t = 1.20", "p = 0.240", "p = 0.300", "Welch's t-test", "Cohen's d = 0.42", "-"],
                    ["p27 Negative", "3", "59.0 Â± 9.1", "t = 1.20", "p = 0.240", "p = 0.300", "Welch's t-test", "Cohen's d = 0.42", "-"],
                ],
            }],
            "figures": [{
                "title": "Age by p27 expression status",
                "png_data_uri": PNG_1X1,
                "caption": "Boxplot of Age by p27 expression status.",
            }],
        },
        assoc_test("Histological type", p_value="p = 0.004", adjusted="p = 0.013"),
        assoc_test("ER", p_value="p = 0.005", adjusted="p = 0.018"),
        assoc_test("PR", p_value="p = 0.002", adjusted="p = 0.012", effect="Cramer's V = 0.50"),
        assoc_test("Molecular subtype", p_value="p = 0.006", adjusted="p = 0.019"),
        assoc_test("AR", p_value="p = 0.030", adjusted="p = 0.045"),
        assoc_test("pT", p_value="p = 0.400", adjusted="p = 0.700"),
        assoc_test("Nodal status", p_value="p = 0.420", adjusted="p = 0.710"),
        assoc_test("LVI", p_value="p = 0.500", adjusted="p = 0.750"),
    ]
    significant_findings = [
        {
            "variable": "Histological type vs Positive/ Negative",
            "key_finding": "Histological type was associated with the primary outcome.",
            "test_statistic": "chi-square = 8.40",
            "p_value": "p = 0.004",
            "adjusted_p_value": "p = 0.013",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.46",
            "notes_warnings": "-",
        },
        {
            "variable": "ER vs Positive/ Negative",
            "key_finding": "ER status was associated with the primary outcome.",
            "test_statistic": "chi-square = 8.00",
            "p_value": "p = 0.005",
            "adjusted_p_value": "p = 0.018",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.46",
            "notes_warnings": "-",
        },
        {
            "variable": "PR vs Positive/ Negative",
            "key_finding": "PR status was associated with the primary outcome.",
            "test_statistic": "chi-square = 9.30",
            "p_value": "p = 0.002",
            "adjusted_p_value": "p = 0.012",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.50",
            "notes_warnings": "-",
        },
        {
            "variable": "Molecular subtype vs Positive/ Negative",
            "key_finding": "Molecular subtype was associated with the primary outcome.",
            "test_statistic": "chi-square = 7.80",
            "p_value": "p = 0.006",
            "adjusted_p_value": "p = 0.019",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.46",
            "notes_warnings": "-",
        },
        {
            "variable": "AR vs Positive/ Negative",
            "key_finding": "AR status was associated with the primary outcome.",
            "test_statistic": "chi-square = 4.70",
            "p_value": "p = 0.030",
            "adjusted_p_value": "p = 0.045",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.32",
            "notes_warnings": "-",
        },
        {
            "variable": "Interpretation-site vs Positive/ Negative",
            "key_finding": "Marker component association.",
            "test_statistic": "chi-square = 7.00",
            "p_value": "p = 0.008",
            "adjusted_p_value": "p = 0.020",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.44",
            "notes_warnings": "-",
        },
        {
            "variable": "Staining Result vs Positive/ Negative",
            "key_finding": "Marker component association.",
            "test_statistic": "chi-square = 7.10",
            "p_value": "p = 0.007",
            "adjusted_p_value": "p = 0.020",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.45",
            "notes_warnings": "-",
        },
    ]
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(12, 6),
        classifications=classifications,
        assignment={"outcome": "Positive/ Negative"},
        table_one=table_one,
        tests=tests,
        graphs=[],
        significant_findings=significant_findings,
        methods_text=(
            "Continuous variables are summarised as mean ± SD. "
            "Descriptive statistics were used for baseline variables. Welch's t-test "
            "and chi-square tests were used for completed bivariate analyses. "
            "A p-value threshold of p &lt; 0.05 was used."
        ),
        results_narrative="ER and PR were significantly associated with p27 expression status.",
        session={
            "study_type": "association",
            "domain_profile": "breast_pathology",
            "main_marker": "p27",
            "main_outcome_concept": "p27 expression status",
            "analysis_excluded_columns": ["positive_nodes"],
        },
        plan={"debug": {"eligible_predictor_count": 4, "bivariate_test_count": 2}},
    )
    blueprint["tables"].append({
        "table_id": "internal_observed",
        "title": "Observed counts",
        "table_type": "internal_detail",
        "columns": ["A", "B"],
        "rows": [["x", "y"]],
        "thesis_ready": True,
        "priority": "detailed_report_only",
        "detailed_report_only": True,
        "warnings": [],
    })
    blueprint["figures"].append({
        "figure_id": "internal_detail_figure",
        "title": "Internal detailed figure",
        "graph_type": "bar_chart",
        "caption": "Internal detailed figure.",
        "thesis_ready": True,
        "priority": "detailed_report_only",
        "optional": True,
        "detailed_report_only": True,
        "warnings": [],
    })
    return {
        "table_one": table_one,
        "tests": tests,
        "significant_findings": significant_findings,
        "thesis_analysis_blueprint": blueprint,
        "export_metadata": {
            "dataset_id": "chapter-v-job",
            "result_id": "chapter-v-result",
            "analysis_version": 1,
            "generated_at": "2026-06-20T00:00:00+00:00",
            "domain_profile": "breast_pathology",
        },
    }


def verify_service_docx() -> None:
    results = _representative_results()
    blob = chapter_v_export.generate_docx(results)
    assert blob and len(blob) > 1000
    text = _docx_text(blob)
    assert "CHAPTER V" in text
    assert "OBSERVATION AND RESULTS" in text
    assert "The present chapter summarises" in text
    assert "5.1 Study Summary" in text
    assert "5.2 Statistical Methods" in text
    assert "5.3 Key Findings" in text
    assert "Section I - Baseline Characteristics" in text
    assert "Section II - Nodal and Prognostic Pathology" in text
    assert "Section III - Immunophenotype" in text
    assert "Section IV - p27 expression status / Marker Expression" in text
    assert "Section V - Statistical Associations" in text
    assert "Section VI - Summary of Tested Associations" in text
    assert "Summary of Significant Associations" in text
    assert "Significant Findings Highlight" in text
    assert "Detailed Association Tables" in text
    assert "Baseline demographic and tumour profile" in text
    assert "Nodal and adverse pathological features" in text
    assert "Tumour burden and pathological stage" in text
    assert "Immunophenotype and molecular subtype profile" in text
    assert "p27 expression and marker components" in text
    assert "Adjusted p-value summary for key associations" in text
    assert "Descriptive Visual Profile" in text
    assert "Age summary" in text
    assert "Distribution of Pathological T stage" in text
    assert "Positive/present profile" in text
    assert "Distribution of Molecular subtype" in text
    assert "Limitations and Interpretation Notes" in text
    assert "Analysis Audit Summary" in text
    assert "Interpretation: Interpretation:" not in text
    assert "This table summarises the analysed sample." not in text
    assert "p27 positivity was observed in 9 cases (75.0%), while 3 cases (25.0%) were p27-negative. This distribution formed the basis for subsequent association analyses." in text
    assert "Missing\n0\n0.0%" not in text
    assert "p27 expression status" in text
    assert "Positive\n9\n75.0%" in text
    assert "No:" not in text and "Yes:" not in text
    assert "Postive" not in text
    assert "Domain-profile grouping is descriptive" not in text
    assert "selected domain profile" not in text
    assert "Variable\nType\nOverall" not in text
    assert "Parameter\nCategory\nn\n%" in text
    assert "Table 1. Table 1." not in text
    assert "Dataset ID" not in text
    assert "Result ID" not in text
    assert "Generated at" not in text
    assert "breast_pathology" not in text
    assert "cross_sectional_association" not in text
    assert "&gt;" not in text and "&lt;" not in text and "&amp;" not in text
    assert "p < 0.05" in text
    assert "mean ± SD" in text
    assert "Variable / parameter" in text
    assert "p-value" in text and "Adjusted p-value" in text
    doc = Document(io.BytesIO(blob))

    def _table_rows(table) -> list[list[str]]:
        return [[cell.text for cell in row.cells] for row in table.rows]

    def _is_descriptive_table(table) -> bool:
        if not table.rows:
            return False
        header = [cell.text for cell in table.rows[0].cells]
        return header[:2] == ["Parameter", "n"] or header == ["Parameter", "Category", "n", "%"]

    descriptive_tables = [table for table in doc.tables if _is_descriptive_table(table)]
    assert len(descriptive_tables) <= 6
    continuous_table = next(
        table for table in descriptive_tables
        if [cell.text for cell in table.rows[0].cells][:2] == ["Parameter", "n"]
        and any(row.cells[0].text == "Age" for row in table.rows[1:])
    )
    continuous_headers = [cell.text for cell in continuous_table.rows[0].cells]
    continuous_values = next(
        [cell.text for cell in row.cells] for row in continuous_table.rows[1:] if row.cells[0].text == "Age"
    )
    assert continuous_headers[0:2] == ["Parameter", "n"]
    assert continuous_headers[2].startswith("Mean") and "SD" in continuous_headers[2]
    assert continuous_headers[3:] == ["Median", "Minimum", "Maximum", "Missing n (%)"]
    assert continuous_values[0] == "Age" and continuous_values[1] == "12"
    assert "57.2" in continuous_values[2] and "8.6" in continuous_values[2]
    assert continuous_values[-1] == "0 (0.0%)"
    immuno_rows = [
        row for table in descriptive_tables
        for row in _table_rows(table)
        if row[:1] == ["ER"]
    ]
    assert ["ER", "Positive", "8", "66.7%"] in immuno_rows
    outcome_table_count = sum(
        1
        for table in descriptive_tables
        if any(row[0] == "p27 expression status" for row in _table_rows(table)[1:])
    )
    assert outcome_table_count == 1
    final_table = next(
        table for table in doc.tables
        if table.rows
        and "Variable / parameter" in [cell.text for cell in table.rows[0].cells]
        and "Adjusted p-value" in [cell.text for cell in table.rows[0].cells]
    )
    final_rows = [[cell.text for cell in row.cells] for row in final_table.rows]
    histology = next(row for row in final_rows if row[0] == "Histological type vs p27 expression status")
    assert histology[3] == "p = 0.004" and histology[4] == "p = 0.013"
    assert "p27 Positive" in text and "p27 Negative" in text
    assert "56.1" in text and "59.0" in text
    assert "Histological type vs p27 expression status" in text
    assert "ER vs p27 expression status" in text
    assert "PR vs p27 expression status" in text
    assert "Molecular subtype vs p27 expression status" in text
    assert "AR vs p27 expression status" in text
    assert "Figure 1. Distribution of p27 expression status." in text
    assert "Figure 2." in text and "p27 expression status by Age" in text
    assert "p27 expression status by Histological type" in text
    assert "p27 expression status by ER" in text
    assert "p27 expression status by PR" in text
    assert "p27 expression status by Molecular subtype" in text
    assert "p27 expression status by AR" in text
    assert "p27 expression status by pT" not in text
    assert "p27 expression status by Nodal status" not in text
    assert "p27 expression status by LVI" not in text
    assert "Graph preview not generated yet" not in text
    final_section = text.split("Section VI - Summary of Tested Associations", 1)[1].split("Warnings and Interpretation Notes", 1)[0]
    assert "Interpretation-site" not in final_section
    assert "Staining Result" not in final_section
    assert "Observed counts" not in text
    assert "Expected counts" not in text
    assert "Row percentages" not in text
    assert "Column percentages" not in text
    assert "section_id" not in text and "source_result_id" not in text
    assert "cross_sectional_association" not in text
    assert "Internal detailed figure" not in text

    def _canonical_member(value: str) -> str:
        text = str(value or "").strip()
        aliases = {
            "Her2Neu": "HER2",
            "Her2neu": "HER2",
            "HER2neu": "HER2",
            "Ki67": "Ki-67",
            "Tumour site / quadrant": "Tumour quadrant",
            "Tumor site / quadrant": "Tumour quadrant",
            "Tumour quadrant / quadrant": "Tumour quadrant",
            "pT": "Pathological T stage",
            "Interpretation-site": "p27 staining localization",
            "Interpretation - site": "p27 staining localization",
            "Staining Result": "p27 staining score pattern",
        }
        return aliases.get(text, text)

    descriptive_table_members = []
    for table in descriptive_tables:
        members = set()
        for row in table.rows[1:]:
            value = row.cells[0].text
            if value:
                members.add(_canonical_member(value))
        descriptive_table_members.extend(members)
    for variable in ["Age", "ER", "PR", "AR", "pT", "Nodal status", "LVI", "Interpretation-site", "Staining Result"]:
        count = descriptive_table_members.count(_canonical_member(variable))
        assert count == 1, f"{variable} appears {count} time(s) in descriptive sections: {descriptive_table_members}"

    regular_word = export.to_docx(FakeEntry(), results, {"outcome": "Positive/ Negative"})
    regular_text = _docx_text(regular_word)
    assert "CHAPTER V" in regular_text
    assert "Observed counts" not in regular_text
    pdf_blob = chapter_v_export.generate_pdf(results)
    assert pdf_blob.startswith(b"%PDF") and len(pdf_blob) > 1000


def verify_generic_docx() -> None:
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(20, 2),
        classifications=[
            {"column": "Score", "detected_type": "scale"},
            {"column": "Treatment", "detected_type": "nominal"},
        ],
        assignment={"outcome": "Score", "group": "Treatment"},
        table_one={
            "headers": ["Variable", "Type", "Overall"],
            "rows": [{"variable": "Score", "type": "mean ± SD", "cells": ["12.4 ± 3.1"]}],
        },
        tests=[],
        graphs=[],
        significant_findings=[],
        methods_text="Welch's t-test was used for the two-group comparison.",
        results_narrative="No significant finding was detected.",
        session={"study_type": "two-group comparison", "domain_profile": "generic"},
    )
    blueprint["analysis_sections"].append({
        "section_id": "bivariate_associations",
        "title": "Bivariate Associations / Group Comparisons",
        "purpose": "Summarise predictor-by-outcome tests.",
        "source_results": ["generic_assoc"],
        "tables": [{
            "table_id": "generic_assoc_table",
            "title": "Association of Score with Treatment",
            "table_type": "categorical_association_thesis_table",
            "columns": ["Predictor", "p-value", "Test applied"],
            "rows": [["Treatment", "p = 0.120", "Chi-square test"]],
            "source_variables": ["Treatment", "Score"],
            "source_test_ids": ["generic_assoc"],
            "interpretation": "Treatment did not show a statistically significant association with Score.",
            "thesis_ready": True,
            "priority": "thesis_ready_primary",
            "optional": False,
            "detailed_report_only": False,
            "warnings": [],
        }],
        "figures": [],
        "interpretation": "Bivariate associations were reviewed.",
    })
    blueprint["tested_associations"] = [{
        "predictor": "Treatment",
        "test_applied": "Chi-square test",
        "p_value": "p = 0.120",
        "adjusted_p_value": "-",
        "effect_size": "-",
        "significance_status": "Not statistically significant",
        "notes_warnings": "",
    }]
    blob = chapter_v_export.generate_docx({"thesis_analysis_blueprint": blueprint, "export_metadata": {"result_id": "generic-result"}})
    text = _docx_text(blob)
    assert "CHAPTER V" in text
    assert "Two Group Comparison" in text or "two-group comparison" in text.lower()
    assert "p27" not in text
    assert "Percentages in detailed association tables are calculated within predictor categories unless otherwise stated." in text
    assert "marker/category distribution within p27 expression groups" not in text


def verify_p27_fallback_is_strictly_gated() -> None:
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(12, 2),
        classifications=[
            {"column": "p27 expression status", "detected_type": "nominal"},
            {"column": "Education group", "detected_type": "nominal"},
        ],
        assignment={"outcome": "p27 expression status"},
        table_one={
            "headers": ["Variable", "Type", "Overall"],
            "rows": [{"variable": "Education group", "type": "n (%)", "cells": ["A: 6 (50.0%); B: 6 (50.0%)"]}],
        },
        tests=[],
        graphs=[],
        significant_findings=[{
            "variable": "Education group vs p27 expression status",
            "key_finding": "Education group was associated with the outcome.",
            "test_statistic": "chi-square = 4.00",
            "p_value": "p = 0.046",
            "adjusted_p_value": "-",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.20",
            "notes_warnings": "",
        }],
        methods_text="Chi-square testing was used.",
        results_narrative="Education group was associated with the outcome.",
        session={"study_type": "association", "domain_profile": "generic"},
    )
    findings = blueprint["significant_findings"]
    assert findings
    key_finding = findings[0]["key_finding"]
    assert key_finding == "Education group was associated with the outcome."
    assert "p27 positivity" not in key_finding
    assert "p27-negative" not in key_finding


async def verify_post_endpoint() -> None:
    entry = FakeEntry()
    original_get = stats.dataset_store.get
    stats.dataset_store.get = lambda job_id: entry
    try:
        published = stats._publish_results(entry, "chapter-v-job", _representative_results())
        result_id = published["export_metadata"]["result_id"]
        response = await stats.export_chapter_v(stats.ChapterVExportRequest(
            job_id="chapter-v-job",
            result_id=result_id,
            format="docx",
        ))
        assert response.media_type.endswith("wordprocessingml.document")
        text = _docx_text(response.body)
        assert "CHAPTER V" in text
        assert "p27 expression status" in text
    finally:
        stats.dataset_store.get = original_get


async def verify_consent_driven_ai_polish_audit_states() -> None:
    """The visible Word/PDF export route (/export/{job_id}/{fmt}) must only
    attempt AI narration polish when the caller sends
    X-Narrative-Polish-Consent: true, and the exported audit line must
    correctly distinguish "never requested" from "requested but unavailable
    or rejected" from "requested and applied". Excel must never request
    polish at all (it has no narrative prose)."""
    entry = FakeEntry()
    original_get = stats.dataset_store.get
    stats.dataset_store.get = lambda job_id: entry
    try:
        published = stats._publish_results(entry, "consent-job", _representative_results())
        result_id = published["export_metadata"]["result_id"]

        # 1. No consent header -> deterministic only, narrative_polish never called.
        with patch.object(narrative_polish, "polish_results") as mocked:
            response = await stats.export("consent-job", "word", result_id=result_id)
            mocked.assert_not_called()
        text = _docx_text(response.body)
        assert "AI polish: deterministic only." in text
        assert "AI polish: OpenRouter narration polish applied." not in text
        assert "AI polish: deterministic fallback used." not in text

        # 2. Consent given, AI polish mocked successful -> applied.
        with patch.object(
            narrative_polish, "polish_results",
            return_value={"results_synthesis": "A polished, fact-preserving synthesis paragraph."},
        ) as mocked:
            response = await stats.export(
                "consent-job", "word", result_id=result_id,
                x_narrative_polish_consent="true",
            )
            mocked.assert_called_once()
        text = _docx_text(response.body)
        assert "AI polish: OpenRouter narration polish applied." in text
        assert "A polished, fact-preserving synthesis paragraph." in text

        # 3. Consent given, but AI polish unavailable/rejected (empty overrides)
        #    -> deterministic fallback used, not "deterministic only".
        with patch.object(narrative_polish, "polish_results", return_value={}) as mocked:
            response = await stats.export(
                "consent-job", "word", result_id=result_id,
                x_narrative_polish_consent="true",
            )
            mocked.assert_called_once()
        text = _docx_text(response.body)
        assert "AI polish: deterministic fallback used." in text
        assert "AI polish: deterministic only." not in text
        assert "AI polish: OpenRouter narration polish applied." not in text

        # 4. Excel never requests AI polish, even with consent given.
        with patch.object(narrative_polish, "polish_results") as mocked:
            await stats.export(
                "consent-job", "excel", result_id=result_id,
                x_narrative_polish_consent="true",
            )
            mocked.assert_not_called()
    finally:
        stats.dataset_store.get = original_get


def verify_frontend_wiring() -> None:
    root = Path(__file__).resolve().parents[1]
    html = (root / "public/analysis.html").read_text(encoding="utf-8")
    js = (root / "public/js/analysis.js").read_text(encoding="utf-8")
    assert "Download Word Report" in html
    assert "Download PDF Report" in html
    assert "Download Cleaned Excel" in html
    assert '<div class="se-export-feature is-hidden" hidden>' in html
    assert "/export/chapter-v" in js
    assert 'data-testid="button-chapter-v-word">Download Chapter V DOCX' in html
    assert "include_detailed_appendix: false" in js
    assert "include_optional_figures: false" in js

    # AI-polish consent checkbox: present near export controls, default
    # unchecked (no "checked" attribute), with the required safety note.
    assert 'id="ai-polish-consent-checkbox"' in html
    assert 'type="checkbox" id="ai-polish-consent-checkbox"' in html, \
        "AI polish checkbox must default OFF (no checked attribute)"
    assert "Polish narrative with AI" in html
    assert "Only narrative text is polished; statistics and tables remain deterministic." in html
    assert "X-Narrative-Polish-Consent" not in html  # backend-only header, not user-facing text
    for secret_name in ("OPENROUTER_API_KEY", "OPENAI_API_KEY", "GEMINI_API_KEY"):
        assert secret_name not in html

    # JS must read the checkbox and only send the consent header when checked,
    # never send it for the Excel format (no narrative prose to polish).
    assert "function aiPolishConsentRequested" in js
    assert "ai-polish-consent-checkbox" in js
    assert 'const aiRequested = format !== "excel" && aiPolishConsentRequested();' in js
    assert 'X-Narrative-Polish-Consent"] = "true"' in js
    assert 'if (aiRequested)' in js
    for secret_name in ("OPENROUTER_API_KEY", "OPENAI_API_KEY", "GEMINI_API_KEY"):
        assert secret_name not in js


def main() -> None:
    from test_fixtures._network_guard import block_real_openrouter_calls
    with block_real_openrouter_calls():
        verify_service_docx()
        verify_generic_docx()
        verify_p27_fallback_is_strictly_gated()
        asyncio.run(verify_post_endpoint())
        asyncio.run(verify_consent_driven_ai_polish_audit_states())
        verify_frontend_wiring()
    print("Sigma Chapter V export verification passed.")


if __name__ == "__main__":
    main()
