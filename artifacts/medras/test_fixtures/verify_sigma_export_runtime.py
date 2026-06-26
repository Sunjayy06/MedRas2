"""Directly verify every Sigma export endpoint against normalized Step 7 state."""

from __future__ import annotations

import asyncio
import io
import time
from pathlib import Path
from unittest.mock import patch

import pandas as pd
from docx import Document
from fastapi import HTTPException
from openpyxl import load_workbook

from app.api import stats


PNG_1X1 = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


class FakeEntry:
    def __init__(self):
        self.df = pd.DataFrame({
            "PR": ["Positive", "Negative"] * 10,
            "ER": ["Positive", "Positive", "Negative", "Negative"] * 5,
        })
        self.meta = {
            "filename": "export-runtime.csv",
            "assignment": {"outcome": "PR", "group": None, "covariates": []},
            "classifications": [
                {"column": "PR", "detected_type": "nominal"},
                {"column": "ER", "detected_type": "nominal"},
            ],
            "cleanup_notes": {"PR": "Whitespace normalized."},
            "quality_log": {"reviewed": 1},
        }


def _results():
    payload = {
        "table_one": {
            "headers": ["Variable", "Type", "Overall"],
            "rows": [{"variable": "PR", "type": "n (%)", "cells": ["20 (100.0%)"]}],
        },
        "tests": [{
            "id": "association_0",
            "title": "ER vs PR: Chi-square test",
            "plan_name": "ER vs PR: Chi-square / Fisher's exact",
            "test_type": "chi_square",
            "analysis_family": "bivariate",
            "actual_test_used": "Chi-square test",
            "p": 0.04,
            "warning": "Sparse cells should be interpreted cautiously.",
            "tables": [{
                "title": "Test summary",
                "headers": ["Test used", "p-value", "Effect size"],
                "rows": [["Chi-square test", "p = 0.040", "Cramer's V = 0.30"]],
            }],
            "figures": [{"title": "ER by PR (%)", "png_data_uri": PNG_1X1}],
            "narrative": "ER was associated with PR.",
        }],
        "graphs": [],
        "methods_md": "Categorical associations used chi-square or Fisher as appropriate.",
        "results_md": "ER was associated with PR.",
    }
    payload["thesis_analysis_blueprint"] = {
        "title": "Observation and Results",
        "study_summary": {"n": 20, "domain_profile": "generic"},
        "study_design": "cross_sectional_association",
        "primary_outcome": "PR",
        "analysis_sections": [
            {
                "section_id": "baseline_characteristics",
                "title": "Baseline and Study Characteristics",
                "purpose": "Describe the analysed sample.",
                "source_results": [],
                "tables": [{
                    "table_id": "table_one",
                    "title": "Table 1. Baseline Characteristics",
                    "table_type": "descriptive_table",
                    "columns": ["Variable", "Type", "Overall"],
                    "rows": [{"variable": "PR", "type": "n (%)", "cells": ["20 (100.0%)"]}],
                    "source_variables": ["PR"],
                    "source_test_ids": [],
                    "interpretation": "Baseline characteristics are summarised.",
                    "thesis_ready": True,
                    "priority": "thesis_ready_primary",
                    "optional": False,
                    "detailed_report_only": False,
                    "warnings": [],
                }],
                "figures": [],
                "interpretation": "Baseline characteristics are summarised.",
            },
            {
                "section_id": "bivariate_associations",
                "title": "Bivariate Associations / Group Comparisons",
                "purpose": "Summarise predictor-by-outcome tests.",
                "source_results": ["association_0"],
                "tables": [{
                    "table_id": "association_0_thesis",
                    "title": "Association of PR with ER",
                    "table_type": "categorical_association_thesis_table",
                    "columns": ["Predictor category", "PR n (%)", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
                    "rows": [["Positive", "10 (50.0%)", "p = 0.040", "-", "Chi-square test", "Cramer's V = 0.30", "Sparse cells should be interpreted cautiously."]],
                    "source_variables": ["ER", "PR"],
                    "source_test_ids": ["association_0"],
                    "interpretation": "ER was associated with PR.",
                    "thesis_ready": True,
                    "priority": "thesis_ready_primary",
                    "optional": False,
                    "detailed_report_only": False,
                    "warnings": ["Sparse cells should be interpreted cautiously."],
                }],
                "figures": [{
                    "figure_id": "association_0_figure",
                    "title": "ER by PR (%)",
                    "graph_type": "grouped_or_stacked_bar",
                    "source_variables": ["ER", "PR"],
                    "source_result_id": "association_0",
                    "caption": "ER by PR (%).",
                    "png_data_uri": PNG_1X1,
                    "interpretation": "Grouped percentage bar chart for ER by PR.",
                    "thesis_ready": True,
                    "priority": "thesis_ready_primary",
                    "optional": False,
                    "detailed_report_only": False,
                    "warnings": [],
                }],
                "interpretation": "Bivariate analyses compared eligible predictors against PR.",
            },
        ],
        "tables": [],
        "figures": [],
        "significant_findings": [{
            "variable": "ER vs PR",
            "key_finding": "ER was associated with PR.",
            "test_statistic": "chi-square = 4.20",
            "p_value": "p = 0.040",
            "adjusted_p_value": "-",
            "test_applied": "Chi-square test",
            "effect_size": "Cramer's V = 0.30",
            "notes_warnings": "Sparse cells should be interpreted cautiously.",
        }],
        "tested_associations": [{
            "predictor": "ER",
            "test_applied": "Chi-square test",
            "test_statistic": "chi-square = 4.20",
            "p_value": "p = 0.040",
            "adjusted_p_value": "-",
            "effect_size": "Cramer's V = 0.30",
            "significance_status": "Statistically significant",
            "notes_warnings": "Sparse cells should be interpreted cautiously.",
        }],
        "methods_text": "Categorical associations used chi-square or Fisher as appropriate.",
        "results_narrative": "ER was associated with PR.",
        "warnings": ["Sparse cells should be interpreted cautiously."],
        "unavailable_or_recommended_only": [],
        "thesis_ready": True,
    }
    return payload


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    text = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            text.extend(cell.text for cell in row.cells)
    return "\n".join(text)


async def _verify_endpoints():
    entry = FakeEntry()
    original_get = stats.dataset_store.get
    stats.dataset_store.get = lambda job_id: entry
    try:
        published = stats._publish_results(entry, "runtime-job", _results())
        result_id = published["export_metadata"]["result_id"]

        for missing_or_stale_id in (None, "stale-result-id"):
            try:
                await stats.export("runtime-job", "pdf", result_id=missing_or_stale_id)
                raise AssertionError("Missing/stale result_id should block export.")
            except HTTPException as exc:
                assert exc.status_code == 409
                detail = str(exc.detail).lower()
                if missing_or_stale_id is None:
                    assert "result_id" in detail
                else:
                    assert "stale" in detail

        cases = [
            ("word", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            ("pdf", "application/pdf"),
            ("excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        ]
        responses = {}
        for fmt, media_type in cases:
            response = await stats.export("runtime-job", fmt, result_id=result_id)
            assert response.media_type == media_type
            assert response.body
            assert "attachment;" in response.headers["content-disposition"]
            responses[fmt] = response.body

        chapter_word = await stats.export_chapter_v_word("runtime-job", result_id=result_id)
        chapter_pdf = await stats.export_chapter_v_pdf("runtime-job", result_id=result_id)
        assert chapter_word.body and chapter_word.media_type.endswith("wordprocessingml.document")
        assert chapter_pdf.body.startswith(b"%PDF")

        for blob in (responses["word"], chapter_word.body):
            text = _docx_text(blob)
            assert "CHAPTER V" in text
            assert "Observation" in text or "OBSERVATION" in text
            assert "Baseline Characteristics" in text
            assert "Association of PR with ER" in text
            assert "Chi-square test" in text
            assert "PR by ER" in text
            assert "Test summary" not in text
            assert "Data Cleaning Log" not in text
            assert "No primary inferential test ran successfully" not in text

        assert responses["pdf"].startswith(b"%PDF")
        assert len(responses["pdf"]) > 1000
        assert len(chapter_pdf.body) > 1000
        assert b"No primary inferential test ran successfully" not in responses["pdf"]
        assert b"No primary inferential test ran successfully" not in chapter_pdf.body

        def slow_polish(_consent, _res):
            time.sleep(0.05)
            return True, {"results_synthesis": "AI text that should miss the deadline."}, "applied"

        with patch.object(stats, "_EXPORT_AI_POLISH_WAIT_SECONDS", 0.01), \
             patch.object(stats, "_chapter_v_polish_overrides", side_effect=slow_polish):
            timed_out_word = await stats.export(
                "runtime-job", "word", result_id=result_id,
                x_narrative_polish_consent="true",
            )
            timed_out_pdf = await stats.export(
                "runtime-job", "pdf", result_id=result_id,
                x_narrative_polish_consent="true",
            )
        assert timed_out_word.body[:2] == b"PK"
        assert timed_out_pdf.body.startswith(b"%PDF")
        assert timed_out_word.headers["x-medras-ai-polish"] == "fallback"
        assert timed_out_pdf.headers["x-medras-ai-polish"] == "fallback"

        with patch.object(stats, "_chapter_v_polish_overrides", side_effect=AssertionError("Excel must not call AI polish")):
            excel_response = await stats.export("runtime-job", "excel", result_id=result_id)
        assert excel_response.body
        assert excel_response.headers["x-medras-ai-polish"] == "deterministic"

        wb = load_workbook(io.BytesIO(responses["excel"]))
        assert {
            "cleaned_processed_dataset", "variable_classification",
            "cleaning_decisions", "category_merges", "missing_data_decisions",
            "excluded_variables", "analysis_summary", "significant_findings",
            "tested_associations", "detailed_results", "Table 1", "Data Cleaning Log", "Narrative",
        } <= set(wb.sheetnames)
        assert list(next(wb["cleaned_processed_dataset"].iter_rows(values_only=True))) == ["PR", "ER"]
        all_excel_text = "\n".join(
            str(value)
            for sheet in wb.worksheets
            for row in sheet.iter_rows(values_only=True)
            for value in row
            if value is not None
        )
        assert "p27" not in all_excel_text.lower()
        assert "marker/category distribution within p27 expression groups" not in all_excel_text
        assert "Percentages in detailed association tables are calculated within predictor categories unless otherwise stated." in all_excel_text
        result_sheets = [
            name for name in wb.sheetnames
            if name not in {
                "Cover", "Variables", "Data Cleaning Log", "Table 1", "Narrative",
                "cleaned_processed_dataset", "variable_classification",
                "cleaning_decisions", "category_merges", "missing_data_decisions",
                "excluded_variables", "analysis_summary", "significant_findings",
                "tested_associations", "detailed_results",
            }
        ]
        assert result_sheets
        assert all(":" not in name and "/" not in name for name in result_sheets)
        values = [
            value
            for row in wb[result_sheets[0]].iter_rows(values_only=True)
            for value in row if value is not None
        ]
        assert "Test summary" in values
        assert "ER by PR (%)" in values
    finally:
        stats.dataset_store.get = original_get


def main():
    asyncio.run(_verify_endpoints())
    root = Path(__file__).resolve().parents[1]
    js = (root / "public/js/analysis.js").read_text(encoding="utf-8")
    html = (root / "public/analysis.html").read_text(encoding="utf-8")
    for fmt in ("word", "pdf", "excel"):
        assert f'data-action="download" data-format="{fmt}"' in html
    assert "Download Word Report" in html
    assert "Download PDF Report" in html
    assert "Download Cleaned Excel" in html
    assert '<div class="se-export-feature is-hidden" hidden>' in html
    assert "exportErrorMessage(res)" in js
    assert "downloadBlob(blob" in js
    assert "X-Narrative-Polish-Consent" in js
    assert "X-MedRAS-AI-Polish" in js
    assert "AI polish unavailable; deterministic narration used." in js
    assert "finally" in js and "button.disabled = false" in js
    assert "Narrative: deterministic" in html
    assert "Narrative polish: deterministic fallback" in js
    assert "Narrative polish: AI-polished" in js
    print("Sigma export runtime/completeness verification passed.")


if __name__ == "__main__":
    main()
