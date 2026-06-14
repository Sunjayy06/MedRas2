"""Verify breast-pathology result presentation and export correctness.

Run:
    python -m test_fixtures.verify_sigma_breast_report_correctness
"""

import io

import pandas as pd
from docx import Document

from app.services import dataset_store, export, plan, results, variable_classifier


PROFILE = "breast_pathology"


def _fixture() -> pd.DataFrame:
    n = 48
    return pd.DataFrame({
        "PR": ["Positive", "Negative"] * (n // 2),
        "Age": list(range(32, 32 + n)),
        "Laterality": ["Left", "Right"] * (n // 2),
        "ENE": ["Present", "Absent", "Absent", "Present"] * (n // 4),
        "ER": ["Positive"] * 34 + ["Postive"] + ["Negative"] * 13,
        "pT": ["T1c", "T2", "T3", "T4b"] * (n // 4),
        "Nodal status": ["N0", "N1a", "N2a", "N3a"] * (n // 4),
        "Tumour site": [
            "Upper outer quadrant - extensive multifocal disease",
            "Upper inner quadrant - central extension",
            "Lower outer quadrant - lateral extension",
            "Lower inner quadrant - medial extension",
            "Central quadrant involving nipple areolar complex",
            "Overlapping sites of breast",
            None,
            "nan",
        ] * 6,
    })


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def verify_report_presentation() -> None:
    df = _fixture()
    classes = variable_classifier.classify_dataframe(df, profile=PROFILE)
    assignment = {"outcome": "PR", "group": None, "covariates": []}
    session = {
        "study_type": "association",
        "study_type_confirmed": True,
        "analysis_predictors": [
            "Age", "Laterality", "ENE", "ER", "pT", "Nodal status", "Tumour site",
        ],
        "domain_profile": PROFILE,
    }
    plan_dict = plan.generate_plan(df, classes, assignment, {"columns": []}, session=session)
    predictor_warning = next(
        item for item in plan_dict["suggestions"]
        if item["id"] == "predictor_duplicate_labels_ER"
    )
    assert predictor_warning["blocking"] is False
    assert "Postive" in predictor_warning["warning"]
    assert not any(
        item["id"] == "predictor_duplicate_labels_Age"
        for item in plan_dict["suggestions"]
    )

    output = results.run_plan(
        df, classes, assignment, plan_dict, session=session,
    )
    contingency = [
        item for item in output["tests"]
        if item.get("actual_test_used")
    ]
    assert contingency
    assert all("Chi-square / Fisher" not in item["title"] for item in contingency)
    assert all(item["plan_name"] == item["title"] for item in contingency)
    assert results._categorical_graph_uses_horizontal_layout(
        df["Tumour site"].dropna().astype(str).tolist()
    )
    assert results._stacked_bar(df, "PR", "Tumour site")

    job_id = dataset_store.put(df, {
        "filename": "breast_report.xlsx",
        "domain_profile": PROFILE,
        "domain_profile_locked": True,
        "classifications": classes,
        "plan": plan_dict,
        "cleanup_notes": {
            "pT": "Auto-extracted numeric values from text; converted to numeric and treated as scale.",
            "Nodal status": "Auto-extracted numeric values from text; converted to numeric and treated as scale.",
        },
    })
    entry = dataset_store.get(job_id)
    entry.meta["results"] = output

    word = export.to_docx(entry, output, assignment)
    chapter_word = export.generate_chapter_v_word(entry, output, assignment)
    pdf = export.to_pdf(entry, output, assignment)
    chapter_pdf = export.generate_chapter_v_pdf(entry, output, assignment)
    for blob in (word, chapter_word, pdf, chapter_pdf):
        assert blob and len(blob) > 1000

    for text in (_docx_text(word), _docx_text(chapter_word)):
        assert "T1c" in text and "T4b" in text
        assert "N0" in text and "N3a" in text
        assert "Auto-extracted numeric values from text" not in text
        assert "Likely duplicate predictor labels were detected" in text
        assert "Chi-square / Fisher's exact" not in text


if __name__ == "__main__":
    verify_report_presentation()
    print("Sigma breast report correctness verification passed.")
