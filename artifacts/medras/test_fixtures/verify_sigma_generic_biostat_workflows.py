"""Verify Sigma stays generic across non-p27 biostat workflows.

Run:
    python -m test_fixtures.verify_sigma_generic_biostat_workflows
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any, Dict, List

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from app.services import chapter_v_export, export
from app.services.plan import generate_plan
from app.services.thesis_blueprint import build_thesis_analysis_blueprint


PNG_1X1 = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _classes(**types: str) -> List[Dict[str, Any]]:
    return [{"column": column, "detected_type": detected} for column, detected in types.items()]


def _normality(*normal_columns: str) -> Dict[str, Any]:
    return {"columns": [{"column": column, "decision": "normal"} for column in normal_columns]}


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def _pdf_text(blob: bytes) -> str:
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(io.BytesIO(blob)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        try:
            from pypdf import PdfReader
        except Exception:
            from PyPDF2 import PdfReader  # type: ignore[no-redef]
        reader = PdfReader(io.BytesIO(blob))
        return "\n".join((page.extract_text() or "") for page in reader.pages)


def _excel_text(blob: bytes) -> str:
    wb = load_workbook(io.BytesIO(blob))
    return "\n".join(
        str(value)
        for sheet in wb.worksheets
        for row in sheet.iter_rows(values_only=True)
        for value in row
        if value is not None
    )


def _table_one(df: pd.DataFrame, types: Dict[str, str]) -> Dict[str, Any]:
    rows: List[Dict[str, Any]] = []
    for column, detected in types.items():
        series = df[column]
        if detected == "scale":
            rows.append({
                "variable": column,
                "type": "Mean ± SD",
                "cells": [f"{series.mean():.1f} ± {series.std(ddof=1):.1f} (n={series.notna().sum()}); Missing: {series.isna().sum()} (0.0%)"],
            })
        elif detected != "id":
            counts = series.fillna("Missing").astype(str).value_counts(dropna=False)
            total = int(counts.sum()) or 1
            summary = "; ".join(f"{label}: {count} ({count / total * 100:.1f}%)" for label, count in counts.items())
            rows.append({"variable": column, "type": "n (%)", "cells": [summary]})
    return {"headers": ["Variable", "Type", "Overall"], "rows": rows}


def _generic_test(
    *,
    test_id: str,
    title: str,
    family: str,
    table_title: str,
    source_variables: List[str],
    test_applied: str,
    p_value: str = "p = 0.120",
    adjusted: str = "p = 0.240",
    graph_title: str | None = None,
    graph_type: str = "grouped_or_stacked_bar",
) -> Dict[str, Any]:
    return {
        "id": test_id,
        "title": title,
        "test_type": test_applied.lower().replace(" ", "_"),
        "analysis_family": family,
        "tables": [{
            "title": table_title,
            "headers": ["Predictor", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
            "rows": [[source_variables[0], p_value, adjusted, test_applied, "-", "-"]],
        }],
        "figures": ([{
            "title": graph_title,
            "graph_type": graph_type,
            "source_variables": source_variables,
            "source_result_id": test_id,
            "caption": f"{graph_title}.",
            "png_data_uri": PNG_1X1,
        }] if graph_title else []),
    }


@dataclass
class Entry:
    df: pd.DataFrame
    meta: Dict[str, Any]


def _results_payload(
    *,
    name: str,
    df: pd.DataFrame,
    classifications: List[Dict[str, Any]],
    assignment: Dict[str, Any],
    session: Dict[str, Any],
    tests: List[Dict[str, Any]],
    tested_associations: List[Dict[str, Any]] | None = None,
) -> Dict[str, Any]:
    types = {row["column"]: row["detected_type"] for row in classifications}
    table_one = _table_one(df, types)
    tested_associations = tested_associations or [{
        "predictor": (test.get("columns") or test.get("source_variables") or ["Predictor"])[0],
        "test_applied": (test.get("actual_test_used") or test.get("tables", [{}])[0].get("rows", [["", "", "", "Test"]])[0][3]),
        "test_statistic": "-",
        "p_value": "p = 0.120",
        "adjusted_p_value": "p = 0.240",
        "effect_size": "-",
        "significance_status": "Not significant after multiple-testing correction",
        "notes_warnings": "",
    } for test in tests]
    blueprint = build_thesis_analysis_blueprint(
        df_shape=df.shape,
        classifications=classifications,
        assignment=assignment,
        plan={"tests": [{"id": test["id"]} for test in tests], "debug": {"eligible_predictor_count": len(tested_associations), "bivariate_test_count": len(tested_associations)}},
        table_one=table_one,
        tests=tests,
        graphs=[figure for test in tests for figure in test.get("figures", [])],
        significant_findings=[],
        tested_associations=tested_associations,
        methods_text=session.get("methods_text", "Descriptive and inferential analyses were selected from variable roles and measurement level."),
        results_narrative=f"{name} analysis was completed using generic Sigma sections.",
        session=session,
    )
    return {
        "table_one": table_one,
        "tests": tests,
        "graphs": [figure for test in tests for figure in test.get("figures", [])],
        "significant_findings": [],
        "tested_associations": tested_associations,
        "methods_md": blueprint["methods_text"],
        "results_md": blueprint["results_narrative"],
        "thesis_analysis_blueprint": blueprint,
        "export_metadata": {"result_id": f"{name.lower().replace(' ', '-')}-result"},
    }


def _assert_generic_outputs(results: Dict[str, Any], entry: Entry, assignment: Dict[str, Any]) -> None:
    docx_text = _docx_text(chapter_v_export.generate_docx(results))
    pdf_blob = chapter_v_export.generate_pdf(results)
    assert pdf_blob.startswith(b"%PDF") and len(pdf_blob) > 1000
    pdf_text = _pdf_text(pdf_blob)
    xlsx_text = _excel_text(export.to_xlsx(entry, results, assignment))
    combined = "\n".join([docx_text, pdf_text, xlsx_text])
    assert "p27" not in combined.lower()
    assert "breast" not in combined.lower()
    assert "Nodal and Prognostic Pathology" not in combined
    assert "Immunophenotype" not in combined
    assert "Marker Expression" not in combined
    assert "marker/category distribution within p27 expression groups" not in combined
    assert "Percentages in detailed association tables are calculated within predictor categories unless otherwise stated." in combined
    for marker in ("HER2", "Ki-67"):
        assert marker not in combined
    for marker in ("ER", "PR", "AR"):
        assert not re.search(rf"\b{marker}\b", combined)


def test_diabetes_metabolic_association() -> None:
    df = pd.DataFrame({
        "Age": [42, 55, 61, 39, 47, 52, 68, 44, 58, 63, 49, 57],
        "Sex": ["F", "M"] * 6,
        "BMI": [23.4, 31.2, 29.8, 22.1, 27.5, 35.4, 30.2, 24.1, 33.0, 28.6, 26.0, 32.4],
        "Fasting glucose": [91, 145, 132, 88, 104, 160, 150, 96, 142, 138, 101, 155],
        "HbA1c": [5.4, 7.8, 7.1, 5.2, 6.0, 8.3, 7.6, 5.5, 7.4, 7.2, 5.8, 8.0],
        "Diabetes status": ["No", "Yes", "Yes", "No", "No", "Yes", "Yes", "No", "Yes", "Yes", "No", "Yes"],
        "Hypertension": ["No", "Yes", "Yes", "No", "No", "Yes", "Yes", "No", "Yes", "No", "No", "Yes"],
        "Treatment group": ["Lifestyle", "Medication"] * 6,
    })
    classifications = _classes(
        Age="scale", Sex="nominal", BMI="scale", **{
            "Fasting glucose": "scale", "HbA1c": "scale", "Diabetes status": "nominal",
            "Hypertension": "nominal", "Treatment group": "nominal",
        }
    )
    predictors = ["Age", "Sex", "BMI", "Fasting glucose", "HbA1c", "Hypertension", "Treatment group"]
    sigma_plan = generate_plan(
        df, classifications, {"outcome": "Diabetes status"},
        _normality("Age", "BMI", "Fasting glucose", "HbA1c"),
        {"study_type": "association", "study_type_confirmed": True, "objective": "Association of metabolic predictors with diabetes status", "analysis_predictors": predictors},
    )
    titles = " ".join(test["title"] for test in sigma_plan["tests"])
    assert "BMI by Diabetes status" in titles
    assert "HbA1c by Diabetes status" in titles
    assert "Hypertension vs Diabetes status" in titles
    tests = [_generic_test(
        test_id="diabetes_bmi",
        title="BMI by Diabetes status: Welch's t-test",
        family="bivariate",
        table_title="Comparison of BMI by Diabetes status",
        source_variables=["BMI", "Diabetes status"],
        test_applied="Welch's t-test",
        graph_title="BMI by Diabetes status",
        graph_type="boxplot",
    )]
    session = {
        "study_type": "association",
        "domain_profile": "generic",
        "objective": "Association of metabolic predictors with diabetes status",
        "methods_text": "Metabolic predictors were compared by diabetes status with FDR correction where applicable.",
    }
    results = _results_payload(
        name="Diabetes metabolic association",
        df=df,
        classifications=classifications,
        assignment={"outcome": "Diabetes status"},
        session=session,
        tests=tests,
    )
    assert results["thesis_analysis_blueprint"]["primary_outcome"] == "Diabetes status"
    _assert_generic_outputs(results, Entry(df, {"classifications": classifications, "domain_profile": "generic"}), {"outcome": "Diabetes status"})


def test_treatment_control_comparison() -> None:
    df = pd.DataFrame({
        "Group": ["Treatment", "Control"] * 12,
        "Baseline score": [50 + (i % 5) for i in range(24)],
        "Follow-up score": [58 + (i % 2) * 4 + (i % 5) for i in range(24)],
        "Response status": ["Responder", "Non-responder"] * 12,
        "Adverse event": ["No", "Yes", "No", "No"] * 6,
    })
    classifications = _classes(**{
        "Group": "nominal", "Baseline score": "scale", "Follow-up score": "scale",
        "Response status": "nominal", "Adverse event": "nominal",
    })
    sigma_plan = generate_plan(
        df, classifications, {"outcome": "Follow-up score", "group": "Group"},
        _normality("Follow-up score"),
        {"study_type": "comparison", "study_type_confirmed": True, "objective": "Compare treatment versus control follow-up score"},
    )
    assert any("Welch" in test["title"] or "t-test" in test["title"] for test in sigma_plan["tests"])
    tests = [_generic_test(
        test_id="followup_group",
        title="Follow-up score by Group: Welch's t-test",
        family="bivariate",
        table_title="Comparison of Follow-up score by Group",
        source_variables=["Follow-up score", "Group"],
        test_applied="Welch's t-test",
        graph_title="Follow-up score by Group",
        graph_type="boxplot",
    )]
    results = _results_payload(
        name="Treatment control comparison",
        df=df,
        classifications=classifications,
        assignment={"outcome": "Follow-up score", "group": "Group"},
        session={"study_type": "two-group comparison", "domain_profile": "generic", "objective": "Treatment versus control comparison"},
        tests=tests,
    )
    assert results["thesis_analysis_blueprint"]["primary_outcome"] == "Follow-up score"
    _assert_generic_outputs(results, Entry(df, {"classifications": classifications, "domain_profile": "generic"}), {"outcome": "Follow-up score", "group": "Group"})


def test_pre_post_paired_study() -> None:
    df = pd.DataFrame({
        "Patient ID": list(range(1, 21)),
        "Pre score": [10 + i for i in range(20)],
        "Post score": [12 + i for i in range(20)],
        "Sex": ["F", "M"] * 10,
        "Age": [30 + i for i in range(20)],
    })
    classifications = _classes(**{
        "Patient ID": "id", "Pre score": "scale", "Post score": "scale",
        "Sex": "nominal", "Age": "scale",
    })
    sigma_plan = generate_plan(
        df, classifications, {"outcome": "Pre score"},
        _normality("Pre score", "Post score"),
        {
            "study_type": "repeated_measures", "study_type_confirmed": True,
            "objective": "Pre/post change in score in the same patients",
            "paired": True, "paired_col1": "Pre score", "paired_col2": "Post score",
        },
    )
    assert any((test.get("_phase_b") or {}).get("test_type") == "t_test_paired" for test in sigma_plan["tests"])
    tests = [_generic_test(
        test_id="pre_post_paired",
        title="Pre score vs Post score: paired t-test",
        family="repeated_measures",
        table_title="Paired comparison of Pre score and Post score",
        source_variables=["Pre score", "Post score"],
        test_applied="Paired t-test",
        graph_title="Pre score and Post score",
        graph_type="line_chart",
    )]
    results = _results_payload(
        name="Pre post paired study",
        df=df,
        classifications=classifications,
        assignment={"outcome": "Pre score"},
        session={"study_type": "pre-post study", "domain_profile": "generic", "objective": "Pre/post score comparison"},
        tests=tests,
        tested_associations=[{
            "predictor": "Post score",
            "test_applied": "Paired t-test",
            "test_statistic": "-",
            "p_value": "p = 0.080",
            "adjusted_p_value": "-",
            "effect_size": "-",
            "significance_status": "Not statistically significant",
            "notes_warnings": "",
        }],
    )
    _assert_generic_outputs(results, Entry(df, {"classifications": classifications, "domain_profile": "generic"}), {"outcome": "Pre score"})


def test_diagnostic_accuracy_stays_diagnostic() -> None:
    df = pd.DataFrame({
        "Gold standard disease status": ["Disease", "No disease"] * 15,
        "Test result": ["Positive", "Negative"] * 15,
        "Age": [40 + i for i in range(30)],
    })
    classifications = _classes(**{
        "Gold standard disease status": "nominal", "Test result": "nominal", "Age": "scale",
    })
    sigma_plan = generate_plan(
        df, classifications, {"outcome": "Gold standard disease status", "group": "Test result"},
        _normality(),
        {
            "study_type": "diagnostic", "study_type_confirmed": True,
            "objective": "Diagnostic accuracy, sensitivity, specificity, ROC AUC of test result against gold standard disease status",
            "disease_col": "Gold standard disease status", "test_result_col": "Test result",
        },
    )
    assert any((test.get("_phase_b") or {}).get("test_type") == "diagnostic_accuracy" for test in sigma_plan["tests"])
    tests = [{
        "id": "diagnostic_accuracy",
        "title": "Test result diagnostic accuracy",
        "test_type": "diagnostic_accuracy",
        "analysis_family": "diagnostic_accuracy",
        "tables": [{
            "title": "Diagnostic accuracy estimates",
            "headers": ["Metric", "Estimate", "95% CI"],
            "rows": [["Sensitivity", "0.90", "-"], ["Specificity", "0.85", "-"], ["AUC", "0.88", "-"]],
        }],
        "figures": [{
            "title": "ROC curve for Test result",
            "graph_type": "roc_curve",
            "source_variables": ["Gold standard disease status", "Test result"],
            "source_result_id": "diagnostic_accuracy",
            "caption": "ROC curve for Test result.",
            "png_data_uri": PNG_1X1,
        }],
    }]
    results = _results_payload(
        name="Diagnostic accuracy study",
        df=df,
        classifications=classifications,
        assignment={"outcome": "Gold standard disease status", "group": "Test result"},
        session={"study_type": "diagnostic accuracy", "domain_profile": "generic", "objective": "Diagnostic accuracy of Test result"},
        tests=tests,
        tested_associations=[{
            "predictor": "Test result",
            "test_applied": "Diagnostic accuracy",
            "test_statistic": "AUC = 0.88",
            "p_value": "-",
            "adjusted_p_value": "-",
            "effect_size": "-",
            "significance_status": "Diagnostic estimates reported",
            "notes_warnings": "",
        }],
    )
    assert results["thesis_analysis_blueprint"]["study_design"] == "diagnostic_accuracy"
    _assert_generic_outputs(results, Entry(df, {"classifications": classifications, "domain_profile": "generic"}), {"outcome": "Gold standard disease status", "group": "Test result"})


def main() -> None:
    test_diabetes_metabolic_association()
    test_treatment_control_comparison()
    test_pre_post_paired_study()
    test_diagnostic_accuracy_stays_diagnostic()
    print("Sigma generic biostat workflow verification passed.")


if __name__ == "__main__":
    main()
