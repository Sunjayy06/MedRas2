"""Verify breast-profile cleaning protections and result labels/figures."""

from __future__ import annotations

from pathlib import Path
import io
import re

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from app.services import category_merger, chapter_v_export, export, plan, results, variable_classifier


def _classes(df):
    return variable_classifier.classify_dataframe(df, profile="breast_pathology")


def _normality(*columns):
    return {"columns": [{"column": column, "decision": "normal"} for column in columns]}


def verify_outcome_typo_blocks_then_routes_binary():
    df = pd.DataFrame({
        "PR": ["Positive"] * 35 + ["Negative"] * 34 + ["Postive"],
        "Age": list(range(30, 100)),
        "ER": ["Positive", "Negative"] * 35,
    })
    session = {
        "domain_profile": "breast_pathology",
        "study_type": "association",
        "study_type_confirmed": True,
        "analysis_predictors": ["Age", "ER"],
    }
    dirty_plan = plan.generate_plan(df, _classes(df), {"outcome": "PR"}, _normality("Age"), session)
    warning = next(item for item in dirty_plan["suggestions"] if item["id"] == "outcome_duplicate_labels")
    assert warning["blocking"] is True
    assert "Postive" in warning["warning"] and "Positive" in warning["warning"]

    proposals = category_merger.detect_category_duplicates(
        df["PR"], profile="breast_pathology"
    )
    typo = next(item for item in proposals["borderline"] if "Postive" in item["members"])
    cleaned, _ = category_merger.apply_merges(
        df,
        [{"column": "PR", "canonical": typo["canonical"], "members": typo["members"]}],
        profile="breast_pathology",
    )
    clean_plan = plan.generate_plan(
        cleaned, _classes(cleaned), {"outcome": "PR"}, _normality("Age"), session
    )
    age_test = next(test for test in clean_plan["tests"] if "Age" in test["columns"])
    assert age_test["_phase_b"]["function"] in {"run_pairwise_welch", "run_pairwise_mann_whitney"}
    assert age_test["_phase_b"]["function"] != "run_pairwise_anova"
    assert not any(item.get("blocking") for item in clean_plan["suggestions"])

    generic = pd.DataFrame({
        "Outcome": ["Luminal A", "Luminal B"] * 20,
        "Score": list(range(40)),
    })
    generic_plan = plan.generate_plan(
        generic,
        [{"column": "Outcome", "detected_type": "nominal"},
         {"column": "Score", "detected_type": "scale"}],
        {"outcome": "Outcome"},
        _normality("Score"),
        {
            "domain_profile": "generic",
            "study_type": "association",
            "study_type_confirmed": True,
            "analysis_predictors": ["Score"],
        },
    )
    assert not any(item.get("blocking") for item in generic_plan["suggestions"])


def verify_clinical_labels_and_protected_merges():
    df = pd.DataFrame({
        "pT": ["T1c", "T2", "T3", "T4b"],
        "Nodal status": ["N0", "N1a", "N2a", "N3a"],
        "Molecular subtype": ["Luminal A", "Luminal B", "Her2Neu", "Her2neu"],
    })
    cleaned, notes = variable_classifier.clean_numeric_like_columns(
        df, profile="breast_pathology"
    )
    assert notes == {}
    assert cleaned["pT"].tolist() == df["pT"].tolist()
    assert cleaned["Nodal status"].tolist() == df["Nodal status"].tolist()

    proposals = category_merger.detect_all_columns(
        df, _classes(df), profile="breast_pathology"
    )
    groups = []
    for result in proposals.values():
        groups.extend(result.get("obvious") or [])
        groups.extend(result.get("borderline") or [])
    assert not any({"Luminal A", "Luminal B"} <= set(group["members"]) for group in groups)
    assert any({"Her2Neu", "Her2neu"} <= set(group["members"]) for group in groups)

    protected, actions = category_merger.apply_merges(
        df,
        [{
            "column": "Molecular subtype",
            "canonical": "Luminal B",
            "members": ["Luminal A", "Luminal B"],
        }],
        profile="breast_pathology",
    )
    assert actions == []
    assert protected["Molecular subtype"].tolist() == df["Molecular subtype"].tolist()


def verify_actual_test_label_and_association_figure():
    df = pd.DataFrame({
        "PR": ["Positive", "Negative"] * 30,
        "ER": ["Positive", "Positive", "Negative", "Negative"] * 15,
        "Subtype": ["A", "B", "C"] * 20,
    })
    classes = [
        {"column": "PR", "detected_type": "nominal"},
        {"column": "ER", "detected_type": "nominal"},
        {"column": "Subtype", "detected_type": "nominal"},
    ]
    session = {
        "domain_profile": "breast_pathology",
        "study_type": "association",
        "study_type_confirmed": True,
        "analysis_predictors": ["ER", "Subtype"],
    }
    sigma_plan = plan.generate_plan(df, classes, {"outcome": "PR"}, _normality(), session)
    output = results.run_plan(
        df, classes, {"outcome": "PR"}, sigma_plan,
        confirmed_graph_ids=[], session=session,
    )
    assert output["tests"]
    assert all(test.get("actual_test_used") in {"Chi-square test", "Fisher's exact test"}
               for test in output["tests"])
    assert all(test["actual_test_used"] in test["title"] for test in output["tests"])
    assert any(test.get("figures") for test in output["tests"])
    assert any(
        figure.get("png_data_uri", "").startswith("data:image/png;base64,")
        for test in output["tests"] for figure in test.get("figures") or []
    )

    sparse_rxc = pd.DataFrame({
        "Predictor": ["A"] * 16 + ["B"] * 2 + ["C"] * 2,
        "Outcome": ["Yes", "No"] * 10,
    })
    sparse_result = results.run_chi_or_fisher("Predictor", "Outcome", {}, sparse_rxc)
    assert sparse_result["actual_test_used"] == "Chi-square test"
    assert (
        "This finding should be interpreted cautiously because some expected cell counts were below 5."
        in sparse_result["note"]
    )
    assert "because some." not in sparse_result["note"]
    assert "Interpret with caution: -" not in sparse_result["note"]

    sparse_2x2 = pd.DataFrame({
        "Predictor": ["A"] * 8 + ["B"] * 2,
        "Outcome": ["Yes"] * 8 + ["No"] * 2,
    })
    fisher_result = results.run_chi_or_fisher("Predictor", "Outcome", {}, sparse_2x2)
    assert fisher_result["actual_test_used"] == "Fisher's exact test"


def verify_missing_exclusion_and_significant_summary():
    df = pd.DataFrame({
        "Positive/ Negative": ["Positive"] * 8 + ["Negative"] * 8,
        "PR": ["Positive"] * 8 + ["Negative"] * 8,
        "ER": ["Positive", "Negative"] * 8,
        "Age": [51, 55, 49, 61, 58, 63, 47, 54, 60, 62, 59, 66, 57, 65, 64, 68],
        "positive_nodes": [1, None, None, 2, None, None, 3, None, None, None, 1, None, None, 2, None, None],
        "total_nodes": [12, None, None, 18, None, None, 20, None, None, None, 13, None, None, 15, None, None],
        "node_ratio": [0.08, None, None, 0.11, None, None, 0.15, None, None, None, 0.08, None, None, 0.13, None, None],
    })
    classes = [
        {"column": "Positive/ Negative", "detected_type": "nominal"},
        {"column": "PR", "detected_type": "nominal"},
        {"column": "ER", "detected_type": "nominal"},
        {"column": "Age", "detected_type": "scale"},
        {"column": "positive_nodes", "detected_type": "exclude"},
        {"column": "total_nodes", "detected_type": "exclude"},
        {"column": "node_ratio", "detected_type": "exclude"},
    ]
    excluded = {"positive_nodes", "total_nodes", "node_ratio"}
    assert excluded <= {row["column"] for row in classes if row["detected_type"] == "exclude"}
    session = {
        "study_type": "association",
        "study_type_confirmed": True,
        "analysis_predictors": ["PR", "ER", "Age", "positive_nodes", "total_nodes", "node_ratio"],
        "domain_profile": "breast_pathology",
    }
    sigma_plan = plan.generate_plan(
        df, classes, {"outcome": "Positive/ Negative", "group": None, "covariates": []},
        _normality("Age"), session,
    )
    analysis_payload = {
        "tests": sigma_plan.get("tests"),
        "graphs": sigma_plan.get("graphs"),
        "analysis_layers": sigma_plan.get("analysis_layers"),
        "descriptive_plan": sigma_plan.get("descriptive_plan"),
        "bivariate_plan": sigma_plan.get("bivariate_plan"),
        "multivariate_plan": sigma_plan.get("multivariate_plan"),
    }
    analysis_text = str(analysis_payload)
    assert "positive_nodes" not in analysis_text
    assert "total_nodes" not in analysis_text
    assert "node_ratio" not in analysis_text
    assert any("PR vs Positive/ Negative" in test["title"] for test in sigma_plan["tests"])
    for graph in sigma_plan.get("graphs", []):
        assert graph.get("graph_id")
        assert graph.get("recommended_chart_type")
        assert graph.get("caption")
        assert graph.get("interpretation")
        assert graph.get("why_recommended")
        assert graph.get("thesis_ready") is True
        assert graph.get("outcome") == "Positive/ Negative"
        assert not (excluded & set(graph.get("variables") or []))

    output = results.run_plan(
        df, classes, {"outcome": "Positive/ Negative", "group": None, "covariates": []},
        sigma_plan, session=session,
    )
    rendered_payload = {
        "tests": output.get("tests"),
        "graphs": output.get("graphs"),
        "table_one": output.get("table_one"),
        "significant_findings": output.get("significant_findings"),
        "summary": output.get("summary"),
    }
    rendered = str(rendered_payload)
    assert "positive_nodes" not in rendered
    assert "total_nodes" not in rendered
    assert "node_ratio" not in rendered
    assert any("PR vs Positive/ Negative" in test["title"] for test in output["tests"])
    assert not any(" by PR" in test["title"] or "vs PR" in test["title"] for test in output["tests"])
    assert all(test.get("compact_summary") for test in output["tests"])
    assert output["significant_findings"]
    assert all(row["p_numeric"] < 0.05 for row in output["significant_findings"])
    assert any(row["variable"].startswith("PR vs Positive/ Negative") for row in output["significant_findings"])


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def verify_age_receives_adjusted_pvalue_in_fdr_family():
    """Issue 1: Age is dispatched through the phase-B run_pairwise_welch
    runner, which only sets 'p_value' (not 'p'). apply_correction_if_needed
    filters strictly on 'p', so Age was silently excluded from the FDR
    family while chi-square predictors (which natively set 'p') were
    corrected — Age showed a raw p-value but '-' for adjusted p-value
    in Word/PDF/Excel. Verify Age now receives an adjusted p-value
    everywhere, and that the FDR test count matches the methods text."""
    n = 40
    df = pd.DataFrame({
        "Positive/ Negative": ["Yes"] * 24 + ["No"] * 16,
        "Age": list(range(30, 54)) + list(range(55, 71)),
        "ER": (["Positive"] * 20 + ["Negative"] * 4) + (["Positive"] * 4 + ["Negative"] * 12),
        "PR": (["Positive"] * 19 + ["Negative"] * 5) + (["Positive"] * 5 + ["Negative"] * 11),
    })
    classes = [
        {"column": "Positive/ Negative", "detected_type": "nominal"},
        {"column": "Age", "detected_type": "scale"},
        {"column": "ER", "detected_type": "nominal"},
        {"column": "PR", "detected_type": "nominal"},
    ]
    session = {
        "study_type": "association",
        "study_type_confirmed": True,
        "analysis_predictors": ["Age", "ER", "PR"],
        "domain_profile": "breast_pathology",
        "main_marker": "p27",
        "main_outcome_concept": "p27 expression status",
    }
    assignment = {"outcome": "Positive/ Negative", "group": None, "covariates": []}
    sigma_plan = plan.generate_plan(df, classes, assignment, _normality("Age"), session)
    output = results.run_plan(df, classes, assignment, sigma_plan, session=session)

    age_test = next(test for test in output["tests"] if test.get("test_type") == "welch_ttest")
    assert age_test.get("p") is not None, "Age test must carry a 'p' key for FDR correction to find it"
    assert age_test.get("p_corrected") is not None, "Age must receive an adjusted p-value like every other completed inferential test"

    completed = [
        row for row in output["tested_associations"]
        if row.get("p_value") not in (None, "-", "")
    ]
    assert len(completed) >= 3
    for row in completed:
        assert row.get("adjusted_p_value") not in (None, "-", ""), (
            f"{row['predictor']} is a completed inferential test but has no adjusted p-value: {row}"
        )
    age_row = next(row for row in completed if row["predictor"] == "Age")
    assert age_row["adjusted_p_value"] != "-"

    correction_info = session.get("correction_info") or {}
    n_tests_methods_match = re.search(r"across (\d+) inferential tests", output.get("methods_md") or "")
    if n_tests_methods_match:
        assert int(n_tests_methods_match.group(1)) == correction_info.get("n_tests") == len(completed)

    blueprint = output["thesis_analysis_blueprint"]
    doc_bytes = chapter_v_export.generate_docx({
        "thesis_analysis_blueprint": blueprint, "tests": output["tests"],
    })
    doc = Document(io.BytesIO(doc_bytes))
    tested_assoc_table = next(
        table for table in doc.tables
        if [cell.text for cell in table.rows[0].cells][:1] == ["Predictor"]
    )
    age_docx_row = next(
        [cell.text for cell in row.cells] for row in tested_assoc_table.rows[1:]
        if row.cells[0].text == "Age"
    )
    assert age_docx_row[4] != "-", age_docx_row

    class _FakeEntry:
        def __init__(self, df):
            self.df = df
            self.meta = {
                "filename": "age_fdr.xlsx", "domain_profile": "breast_pathology",
                "assignment": assignment,
                "classifications": classes,
                "data_version": 0,
            }

    xlsx_blob = export.to_xlsx(_FakeEntry(df), output, assignment)
    workbook = load_workbook(io.BytesIO(xlsx_blob))
    ws = workbook["tested_associations"]
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    age_excel_row = next(row for row in rows if row[0] == "Age")
    assert age_excel_row[4] not in (None, "-", ""), age_excel_row


def main():
    verify_outcome_typo_blocks_then_routes_binary()
    verify_clinical_labels_and_protected_merges()
    verify_actual_test_label_and_association_figure()
    verify_missing_exclusion_and_significant_summary()
    verify_age_receives_adjusted_pvalue_in_fdr_family()
    root = Path(__file__).resolve().parents[1]
    js_source = (root / "public/js/analysis.js").read_text(encoding="utf-8")
    api_source = (root / "app/api/stats.py").read_text(encoding="utf-8")
    assert "hasBlockingSuggestion" in js_source
    assert "escapeHtml(test.title)" in js_source
    assert "if blocking:" in api_source
    print("Sigma breast preprocessing/result correctness verification passed.")


if __name__ == "__main__":
    main()
