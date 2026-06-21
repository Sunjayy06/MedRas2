"""Verify Sigma Chapter V Word/PDF readability upgrade: split tables + richer interpretation.

Covers:
  1. Section II ("Clinical and Pathology Characteristics") does not collapse into one
     giant table; tumour quadrant / size / pT stage / nodal status / nodal burden /
     adverse pathology features appear as separate focused tables.
  2. Section III ("Immunophenotype and Marker Characteristics") splits into hormone
     receptor / HER2+proliferation / AR / molecular subtype tables.
  3. Tumour quadrant collapses rare categories into "Other" in the main body while
     the Excel audit workbook still preserves every raw category.
  4. p27 marker-component section splits into localization / score sub-tables.
  5. Descriptive interpretations carry >= 2 meaningful sentences for key sections.
  6. Association interpretations mention direction/pattern for significant predictors.
  7. Non-significant node-derived association tables/figures stay out of the main body.
  8. No raw machine phrases ("because some.", "Interpret with caution: -",
     "welch_ttest") and no oversized thesis-facing descriptive tables.
  9. Generic (non-breast) fixture still exports.

Run:
    python -m test_fixtures.verify_sigma_chapter_v_section_split
"""

from __future__ import annotations

import io
from typing import Any, Dict, List

from docx import Document
from openpyxl import load_workbook
import pandas as pd

from app.services import chapter_v_export, export
from app.services.results import build_table_one
from app.services.thesis_blueprint import build_thesis_analysis_blueprint


PNG_1X1 = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)

ROW_THRESHOLD = 12


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def _observed(rows: List[List[Any]], headers: List[str]) -> Dict[str, Any]:
    return {"title": "Observed counts", "headers": headers, "rows": rows}


def _chi_square_test(
    name: str,
    *,
    p: float,
    p_corrected: float,
    observed_rows: List[List[Any]],
    outcome_headers: List[str] = ("Positive", "Negative"),
    effect: float = 0.46,
    warning: str = "-",
) -> Dict[str, Any]:
    return {
        "id": f"{name.lower().replace(' ', '_')}_vs_outcome",
        "title": f"{name} vs Positive/ Negative",
        "test_type": "chi_square",
        "actual_test_used": "Chi-square test",
        "analysis_family": "bivariate",
        "p": p,
        "p_corrected": p_corrected,
        "cramers_v": effect,
        "warning": warning,
        "tables": [
            _observed(observed_rows, ["Predictor category", *outcome_headers, "Total"]),
            {
                "title": f"Association of p27 expression status with {name}",
                "headers": ["Predictor category", *[f"{h} n (%)" for h in outcome_headers], "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
                "rows": [],
            },
        ],
        "figures": [{
            "title": f"{name} by p27 expression status",
            "png_data_uri": PNG_1X1,
            "caption": f"Grouped percentage bar chart of {name} by p27 expression status.",
        }],
    }


def _group_comparison_test(name: str, *, p: float, p_corrected: float) -> Dict[str, Any]:
    return {
        "id": f"{name.lower().replace(' ', '_')}_vs_outcome",
        "title": f"{name} vs Positive/ Negative",
        "test_type": "welch_ttest",
        "actual_test_used": "Welch's t-test",
        "analysis_family": "bivariate",
        "p": p,
        "p_corrected": p_corrected,
        "cohens_d": 0.05,
        "warning": "-",
        "tables": [{
            "title": f"Comparison of {name} by p27 expression status",
            "headers": ["Group", "n", "Mean ± SD", "Test statistic", "p-value", "Adjusted p-value", "Test applied", "Effect size", "Warnings"],
            "rows": [
                ["Positive", "70", "5.0 ± 2.0", "t = 0.10", f"p = {p:.3f}", f"p = {p_corrected:.3f}", "Welch's t-test", "Cohen's d = 0.05", "-"],
                ["Negative", "46", "5.1 ± 2.1", "t = 0.10", f"p = {p:.3f}", f"p = {p_corrected:.3f}", "Welch's t-test", "Cohen's d = 0.05", "-"],
            ],
        }],
        "figures": [],
    }


def _table_one() -> Dict[str, Any]:
    rows = [
        ("Positive/ Negative", "n (%)", "Yes: 70 (60.3%); No: 46 (39.7%)"),
        ("Age", "Mean ± SD", "55.1 ± 12.0 (n=116); Missing: 0 (0.0%)"),
        ("Tumour quadrant", "n (%)", (
            "Central: 25 (21.6%); Upper outer: 25 (21.6%); Upper inner: 10 (8.6%); "
            "Lower outer: 15 (12.9%); Lower inner: 10 (8.6%); Axillary tail: 8 (6.9%); "
            "Diffuse/Mixed: 23 (19.8%)"
        )),
        ("Tumour size", "n (%)", "<=2 cm: 20 (17.2%); >2 but <=5 cm: 73 (62.9%); >5 cm: 23 (19.8%)"),
        ("pT", "n (%)", "T1: 30 (25.9%); T2: 40 (34.5%); T3: 30 (25.9%); T4: 16 (13.8%)"),
        ("Nodal status", "n (%)", "N0: 59 (50.9%); N1: 28 (24.1%); N2: 18 (15.5%); N3: 11 (9.5%)"),
        ("positive_nodes", "Mean ± SD", "2.1 ± 3.4 (n=116); Missing: 0 (0.0%)"),
        ("total_nodes", "Mean ± SD", "14.5 ± 6.2 (n=116); Missing: 0 (0.0%)"),
        ("node_ratio", "Mean ± SD", "0.15 ± 0.12 (n=116); Missing: 0 (0.0%)"),
        ("ENE", "n (%)", "Present: 30 (25.9%); Absent: 86 (74.1%)"),
        ("DCIS", "n (%)", "Present: 40 (34.5%); Absent: 76 (65.5%)"),
        ("LVI", "n (%)", "Present: 35 (30.2%); Absent: 81 (69.8%)"),
        ("Necrosis", "n (%)", "Present: 20 (17.2%); Absent: 96 (82.8%)"),
        ("ER", "n (%)", "Positive: 74 (63.8%); Negative: 42 (36.2%)"),
        ("PR", "n (%)", "Positive: 72 (62.1%); Negative: 44 (37.9%)"),
        ("Her2Neu", "n (%)", "Negative/low: 60 (51.7%); Equivocal (2+): 20 (17.2%); Positive (3+): 36 (31.0%)"),
        ("Ki67", "n (%)", ">=14%: 50 (43.1%); <14%: 66 (56.9%)"),
        ("EGFR", "n (%)", "Positive: 30 (25.9%); Negative: 86 (74.1%)"),
        ("AR", "n (%)", "Positive: 65 (56.0%); Negative: 51 (44.0%)"),
        ("Molecular subtype", "n (%)", "Luminal A: 40 (34.5%); Luminal B: 35 (30.2%); HER2neu: 20 (17.2%); Triple negative: 21 (18.1%)"),
        ("Interpretation-site", "n (%)", "Nuclear: 60 (51.7%); Cytoplasmic: 30 (25.9%); Nuclear with cytoplasmic: 20 (17.2%); Negative stain: 6 (5.2%)"),
        ("Staining Result", "n (%)", "4+2: 30 (25.9%); 5+2: 40 (34.5%); 6+1: 25 (21.6%); 3+3: 21 (18.1%)"),
    ]
    return {
        "headers": ["Variable", "Type", "Overall"],
        "rows": [{"variable": var, "type": kind, "cells": [summary]} for var, kind, summary in rows],
    }


def _classifications() -> List[Dict[str, Any]]:
    nominal = [
        "Positive/ Negative", "Tumour quadrant", "ER", "PR", "EGFR", "AR",
        "Molecular subtype", "ENE", "DCIS", "LVI", "Necrosis",
        "Interpretation-site", "Staining Result",
    ]
    ordinal = ["Tumour size", "pT", "Nodal status", "Her2Neu", "Ki67"]
    scale = ["Age", "positive_nodes", "total_nodes", "node_ratio"]
    out = []
    for col in nominal:
        out.append({"column": col, "detected_type": "nominal"})
    for col in ordinal:
        out.append({"column": col, "detected_type": "ordinal"})
    for col in scale:
        out.append({"column": col, "detected_type": "scale"})
    return out


def _build_blueprint() -> Dict[str, Any]:
    table_one = _table_one()
    tests = [
        _chi_square_test(
            "Histological type", p=0.004, p_corrected=0.013,
            observed_rows=[["1", 25, 5, 30], ["2", 20, 20, 40], ["3", 5, 25, 30], ["Total", 50, 50, 100]],
            warning="Chi-square test was used because some expected cell counts were below 5.",
        ),
        _chi_square_test(
            "ER", p=0.0009, p_corrected=0.005,
            observed_rows=[["Positive", 50, 10, 60], ["Negative", 5, 45, 50], ["Total", 55, 55, 110]],
        ),
        _chi_square_test(
            "PR", p=0.002, p_corrected=0.009,
            observed_rows=[["Positive", 48, 12, 60], ["Negative", 6, 44, 50], ["Total", 54, 56, 110]],
        ),
        _chi_square_test(
            "Molecular subtype", p=0.0008, p_corrected=0.0009,
            observed_rows=[
                ["Luminal A", 35, 5, 40], ["Luminal B", 30, 5, 35],
                ["HER2neu", 5, 15, 20], ["Triple negative", 2, 19, 21],
                ["Total", 72, 44, 116],
            ],
        ),
        _chi_square_test(
            "AR", p=0.004, p_corrected=0.013,
            observed_rows=[["Positive", 45, 20, 65], ["Negative", 10, 41, 51], ["Total", 55, 61, 116]],
        ),
        _chi_square_test(
            "pT", p=0.400, p_corrected=0.700,
            observed_rows=[["T1", 18, 12, 30], ["T2", 24, 16, 40], ["T3", 18, 12, 30], ["Total", 60, 40, 100]],
        ),
        _chi_square_test(
            "Nodal status", p=0.420, p_corrected=0.710,
            observed_rows=[["N0", 35, 24, 59], ["N1", 17, 11, 28], ["Total", 52, 35, 87],
                            ["N2", 18, 18, 36]],  # noisy but harmless extra row
        ),
        _chi_square_test(
            "LVI", p=0.500, p_corrected=0.750,
            observed_rows=[["Present", 21, 14, 35], ["Absent", 49, 32, 81], ["Total", 70, 46, 116]],
        ),
        _group_comparison_test("positive_nodes", p=0.800, p_corrected=0.900),
        _group_comparison_test("total_nodes", p=0.820, p_corrected=0.900),
        _group_comparison_test("node_ratio", p=0.810, p_corrected=0.900),
    ]
    significant_findings = [
        {
            "variable": "Histological type vs Positive/ Negative",
            "key_finding": "Histological type was associated with the primary outcome.",
            "test_statistic": "chi-square = 8.40", "p_value": "p = 0.004", "adjusted_p_value": "p = 0.013",
            "test_applied": "Chi-square test", "effect_size": "Cramer's V = 0.46", "notes_warnings": "-",
        },
        {
            "variable": "ER vs Positive/ Negative",
            "key_finding": "ER was associated with the primary outcome.",
            "test_statistic": "chi-square = 12.00", "p_value": "p < 0.001", "adjusted_p_value": "p = 0.005",
            "test_applied": "Chi-square test", "effect_size": "Cramer's V = 0.55", "notes_warnings": "-",
        },
        {
            "variable": "PR vs Positive/ Negative",
            "key_finding": "PR was associated with the primary outcome.",
            "test_statistic": "chi-square = 9.30", "p_value": "p = 0.002", "adjusted_p_value": "p = 0.009",
            "test_applied": "Chi-square test", "effect_size": "Cramer's V = 0.50", "notes_warnings": "-",
        },
        {
            "variable": "Molecular subtype vs Positive/ Negative",
            "key_finding": "Molecular subtype was associated with the primary outcome.",
            "test_statistic": "chi-square = 14.00", "p_value": "p < 0.001", "adjusted_p_value": "p < 0.001",
            "test_applied": "Chi-square test", "effect_size": "Cramer's V = 0.60", "notes_warnings": "-",
        },
        {
            "variable": "AR vs Positive/ Negative",
            "key_finding": "AR was associated with the primary outcome.",
            "test_statistic": "chi-square = 8.10", "p_value": "p = 0.004", "adjusted_p_value": "p = 0.013",
            "test_applied": "Chi-square test", "effect_size": "Cramer's V = 0.44", "notes_warnings": "-",
        },
    ]
    return build_thesis_analysis_blueprint(
        df_shape=(116, len(_classifications())),
        classifications=_classifications(),
        assignment={"outcome": "Positive/ Negative"},
        table_one=table_one,
        tests=tests,
        graphs=[],
        significant_findings=significant_findings,
        methods_text="Chi-square tests were used for categorical predictors. p < 0.05 was considered significant.",
        results_narrative="ER, PR, Molecular subtype, Histological type, and AR were significantly associated with p27 expression status.",
        session={
            "study_type": "association",
            "domain_profile": "breast_pathology",
            "main_marker": "p27",
            "main_outcome_concept": "p27 expression status",
            "analysis_excluded_columns": [],
        },
        plan={"debug": {"eligible_predictor_count": 12, "bivariate_test_count": 11}},
    )


def _build_results() -> Dict[str, Any]:
    blueprint = _build_blueprint()
    return {
        "table_one": _table_one(),
        "thesis_analysis_blueprint": blueprint,
        "export_metadata": {
            "dataset_id": "section-split-job",
            "result_id": "section-split-result",
            "analysis_version": 1,
            "generated_at": "2026-06-22T00:00:00+00:00",
            "domain_profile": "breast_pathology",
        },
    }


def _descriptive_tables(doc: Document):
    out = []
    for table in doc.tables:
        if not table.rows:
            continue
        header = [cell.text for cell in table.rows[0].cells]
        if header[:2] == ["Parameter", "n"] or header == ["Parameter", "Category", "n", "%"]:
            out.append(table)
    return out


def test_section_ii_splits_into_focused_tables() -> None:
    blueprint = _build_blueprint()
    clinical = next(s for s in blueprint["analysis_sections"] if s["section_id"] == "clinical_study_characteristics")
    titles = {table["title"] for table in clinical["tables"]}
    expected = {
        "Tumour quadrant distribution",
        "Tumour size distribution",
        "Pathological T stage distribution",
        "Nodal status distribution",
        "Nodal burden summary",
        "Adverse pathological features",
    }
    assert expected <= titles, titles
    assert len(clinical["tables"]) > 1, "Section II must not collapse into one giant table"
    for table in clinical["tables"]:
        assert len(table["rows"]) <= 8, f"{table['title']} carries too many raw rows for a focused table"


def test_section_iii_splits_into_focused_tables() -> None:
    blueprint = _build_blueprint()
    immuno = next(s for s in blueprint["analysis_sections"] if s["section_id"] == "immunophenotype_characteristics")
    titles = {table["title"] for table in immuno["tables"]}
    expected = {
        "Hormone receptor profile",
        "HER2 and proliferation marker profile",
        "AR expression profile",
        "Molecular subtype distribution",
    }
    assert expected <= titles, titles
    assert len(immuno["tables"]) > 1


def test_marker_components_split_localization_and_score() -> None:
    blueprint = _build_blueprint()
    marker = next(s for s in blueprint["analysis_sections"] if s["section_id"] == "marker_outcome_components")
    titles = {table["title"] for table in marker["tables"]}
    assert any("staining localization" in title for title in titles), titles
    assert any("staining score pattern" in title for title in titles), titles
    assert len(marker["tables"]) >= 2


def test_word_export_no_giant_descriptive_table_and_no_raw_machine_phrases() -> None:
    results = _build_results()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)
    doc = Document(io.BytesIO(blob))
    for table in _descriptive_tables(doc):
        data_rows = len(table.rows) - 1
        assert data_rows <= ROW_THRESHOLD, f"Descriptive table exceeds row threshold: {data_rows} rows"
    assert "because some." not in text
    assert "Interpret with caution: -" not in text
    assert "welch_ttest" not in text
    for token in text.split():
        # crude long-raw-float guard: reject 6+ digit decimal fragments like 27.58011794899391
        if "." in token:
            frac = token.split(".")[-1].strip(".,;()%")
            assert len(frac) < 6, f"Found a long raw float fragment: {token}"


def test_tumour_quadrant_collapses_other_in_main_body() -> None:
    results = _build_results()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))
    quadrant_table = None
    for table in _descriptive_tables(doc):
        rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
        if any(row[0] == "Tumour quadrant" for row in rows):
            quadrant_table = rows
            break
    assert quadrant_table is not None, "Tumour quadrant table not found in main body"
    categories = {row[1] for row in quadrant_table}
    assert "Other" in categories
    assert "Axillary tail" not in categories
    assert "Diffuse/Mixed" not in categories
    assert {"Central", "Upper outer", "Upper inner", "Lower outer", "Lower inner"} <= categories
    assert len(quadrant_table) <= 6


def test_excel_audit_preserves_raw_quadrant_categories() -> None:
    results = _build_results()

    class _FakeEntry:
        def __init__(self):
            n = 116
            quadrants = (
                ["Central"] * 25 + ["Upper outer"] * 25 + ["Upper inner"] * 10 +
                ["Lower outer"] * 15 + ["Lower inner"] * 10 + ["Axillary tail"] * 8 +
                ["Diffuse/Mixed"] * 23
            )
            self.df = pd.DataFrame({
                "Positive/ Negative": (["Yes"] * 70 + ["No"] * 46),
                "Tumour quadrant": quadrants,
            })
            self.meta = {
                "filename": "section-split.xlsx",
                "domain_profile": "breast_pathology",
                "assignment": {"outcome": "Positive/ Negative"},
                "classifications": [
                    {"column": "Positive/ Negative", "detected_type": "nominal"},
                    {"column": "Tumour quadrant", "detected_type": "nominal"},
                ],
                "data_version": 0,
            }

    entry = _FakeEntry()
    blob = export.to_xlsx(entry, results, {"outcome": "Positive/ Negative"})
    wb = load_workbook(io.BytesIO(blob))
    ws = wb["cleaned_processed_dataset"]
    all_text = " ".join(
        str(v) for row in ws.iter_rows(min_row=2, values_only=True) for v in row if v is not None
    )
    assert "Axillary tail" in all_text
    assert "Diffuse/Mixed" in all_text


def test_descriptive_interpretations_have_multiple_sentences() -> None:
    results = _build_results()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs if p.text]

    def _sentence_count(text: str) -> int:
        return len([s for s in text.split(".") if s.strip()])

    quadrant_interp = next((p for p in paragraphs if "tumour locations were" in p.lower()), None)
    assert quadrant_interp is not None
    assert _sentence_count(quadrant_interp) >= 2
    assert "for descriptive clarity" in quadrant_interp

    size_interp = next((p for p in paragraphs if "tumours measured" in p.lower()), None)
    assert size_interp is not None
    assert _sentence_count(size_interp) >= 2
    assert "T2-sized lesions" in size_interp

    nodal_interp = next((p for p in paragraphs if "most frequent nodal category" in p.lower()), None)
    assert nodal_interp is not None
    assert _sentence_count(nodal_interp) >= 2
    assert "Node-positive categories" in nodal_interp

    receptor_interp = next((p for p in paragraphs if "ER positivity was observed" in p), None)
    assert receptor_interp is not None
    assert "PR positivity" in receptor_interp


def test_marker_score_pattern_sorted_and_interpreted() -> None:
    results = _build_results()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))
    score_table = None
    for table in _descriptive_tables(doc):
        rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
        categories = [row[1] for row in rows]
        if {"4+2", "5+2", "6+1", "3+3"} <= set(categories):
            score_table = categories
            break
    assert score_table is not None, "p27 staining score table not found"
    assert score_table == sorted(score_table, key=lambda c: tuple(int(p) for p in c.split("+")))
    text = _docx_text(blob)
    assert "staining score pattern was" in text.lower()


def test_association_interpretations_mention_direction() -> None:
    results = _build_results()
    blob = chapter_v_export.generate_docx(results)
    text = _docx_text(blob)

    assert "Histological type showed a statistically significant association" in text
    assert "proportionately higher in the Negative group" in text
    assert "more commonly Positive" in text
    assert "This finding should be interpreted cautiously because some expected cell counts were below 5." in text

    assert "ER showed a statistically significant association" in text
    assert "ER-positive cases were more commonly Positive" in text or "ER-negative cases were proportionately higher" in text

    assert "Molecular subtype showed a statistically significant association" in text
    assert "Triple negative cases were proportionately higher in the Negative group" in text

    assert "PR showed a statistically significant association" in text
    assert "AR showed a statistically significant association" in text


def test_nonsignificant_node_derived_tables_excluded_from_main_body() -> None:
    results = _build_results()
    word_text = _docx_text(chapter_v_export.generate_docx(results))
    pdf_blob = chapter_v_export.generate_pdf(results)
    assert pdf_blob.startswith(b"%PDF")
    for variable in ("positive_nodes", "total_nodes", "node_ratio"):
        assert f"Comparison of {variable}" not in word_text
        assert f"p27 expression status by {variable}" not in word_text
    # Raw results retain the tests even though they are excluded from the main body.
    blueprint = results["thesis_analysis_blueprint"]
    raw_titles = " ".join(
        table.get("title", "")
        for section in blueprint["analysis_sections"]
        for table in section.get("tables", [])
    )
    for variable in ("positive_nodes", "total_nodes", "node_ratio"):
        assert variable in raw_titles


def test_no_duplicate_titles_in_word_export() -> None:
    results = _build_results()
    blob = chapter_v_export.generate_docx(results)
    doc = Document(io.BytesIO(blob))
    titles = []
    for p in doc.paragraphs:
        if p.text.startswith("Table ") and p.runs and p.runs[0].bold:
            titles.append(p.text)
    assert len(titles) == len(set(titles)), f"Duplicate table captions found: {titles}"


def test_generic_non_breast_fixture_still_exports() -> None:
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(20, 2),
        classifications=[
            {"column": "Score", "detected_type": "scale"},
            {"column": "Treatment", "detected_type": "nominal"},
        ],
        assignment={"outcome": "Score", "group": "Treatment"},
        table_one={
            "headers": ["Variable", "Type", "Overall"],
            "rows": [{"variable": "Score", "type": "mean ± SD", "cells": ["12.4 ± 3.1"]}],
        },
        tests=[],
        graphs=[],
        significant_findings=[],
        methods_text="Welch's t-test was used for the two-group comparison.",
        results_narrative="No significant finding was detected.",
        session={"study_type": "two-group comparison", "domain_profile": "generic"},
    )
    blob = chapter_v_export.generate_docx({
        "thesis_analysis_blueprint": blueprint,
        "export_metadata": {"result_id": "generic-section-split"},
    })
    text = _docx_text(blob)
    assert "CHAPTER V" in text
    assert "p27" not in text
    assert "Tumour quadrant" not in text


def main() -> None:
    test_section_ii_splits_into_focused_tables()
    print("  [ok] Section II splits into focused tables")
    test_section_iii_splits_into_focused_tables()
    print("  [ok] Section III splits into focused tables")
    test_marker_components_split_localization_and_score()
    print("  [ok] Marker component section splits into localization/score tables")
    test_word_export_no_giant_descriptive_table_and_no_raw_machine_phrases()
    print("  [ok] Word export has no giant descriptive table or raw machine phrases")
    test_tumour_quadrant_collapses_other_in_main_body()
    print("  [ok] Tumour quadrant collapses rare categories into Other")
    test_excel_audit_preserves_raw_quadrant_categories()
    print("  [ok] Excel audit workbook preserves raw quadrant categories")
    test_descriptive_interpretations_have_multiple_sentences()
    print("  [ok] Descriptive interpretations carry multiple deterministic sentences")
    test_marker_score_pattern_sorted_and_interpreted()
    print("  [ok] p27 staining score pattern sorted logically with interpretation")
    test_association_interpretations_mention_direction()
    print("  [ok] Association interpretations mention direction/pattern")
    test_nonsignificant_node_derived_tables_excluded_from_main_body()
    print("  [ok] Non-significant node-derived association tables excluded from main body")
    test_no_duplicate_titles_in_word_export()
    print("  [ok] No duplicate table titles in Word export")
    test_generic_non_breast_fixture_still_exports()
    print("  [ok] Generic non-breast fixture still exports")
    print("\nSigma Chapter V section-split verification passed.")


if __name__ == "__main__":
    main()
