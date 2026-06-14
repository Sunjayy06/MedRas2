"""Verify every Sigma export is bound to the latest completed analysis."""

import io

import pandas as pd
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook

from app.services import dataset_store
from main import app


PROFILE = "breast_pathology"


def _fixture() -> pd.DataFrame:
    n = 48
    return pd.DataFrame({
        "PR": ["Positive"] * 23 + ["Postive"] + ["Negative"] * 24,
        "Age": list(range(30, 30 + n)),
        "ER": ["Positive", "Negative"] * (n // 2),
        "pT": ["T1c", "T2", "T3", "T4b"] * (n // 4),
        "Nodal status": ["N0", "N1a", "N2a", "N3a"] * (n // 4),
    })


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    return "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )


def _xlsx_text(blob: bytes) -> str:
    wb = load_workbook(io.BytesIO(blob), data_only=True)
    return "\n".join(
        str(value)
        for ws in wb.worksheets
        for row in ws.iter_rows(values_only=True)
        for value in row
        if value is not None
    )


def verify_current_export_state() -> None:
    buf = io.BytesIO()
    _fixture().to_excel(buf, index=False)
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/stats/upload",
            files={"file": ("breast.xlsx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            data={"profile": "generic"},
        ).json()
        job_id = uploaded["job_id"]
        generic_classified = client.post(
            "/api/stats/classify",
            json={"job_id": job_id, "overrides": [], "profile": "generic"},
        )
        assert generic_classified.status_code == 200
        classified = client.post(
            "/api/stats/classify",
            json={"job_id": job_id, "overrides": [], "profile": PROFILE},
        )
        assert classified.status_code == 200
        assert classified.json()["domain_profile"] == PROFILE
        entry = dataset_store.get(job_id)
        assert entry.meta["domain_profile"] == PROFILE
        assert {"T1c", "T4b"} <= set(entry.df["pT"].dropna().astype(str))
        assert {"N0", "N3a"} <= set(entry.df["Nodal status"].dropna().astype(str))
        assert not any(
            "auto-extracted numeric values" in str(note).lower()
            for column, note in (entry.meta.get("cleanup_notes") or {}).items()
            if column in {"pT", "Nodal status"}
        )
        dupes = client.post(
            "/api/stats/detect-category-dupes",
            json={"job_id": job_id, "profile": PROFILE},
        ).json()["columns"]
        pr_merge = next(
            item for item in dupes["PR"]["borderline"]
            if {"Positive", "Postive"} <= set(item["members"])
        )
        merged = client.post(
            "/api/stats/apply-category-merge",
            json={
                "job_id": job_id,
                "profile": PROFILE,
                "merges": [{"column": "PR", "canonical": "Positive", "members": pr_merge["members"]}],
            },
        )
        assert merged.status_code == 200
        client.post("/api/stats/classify", json={"job_id": job_id, "overrides": [], "profile": PROFILE})
        client.post(
            "/api/stats/assign",
            json={"job_id": job_id, "outcome": "PR", "group": None, "covariates": ["Age", "ER", "pT", "Nodal status"]},
        )
        plan = client.get(f"/api/stats/generate-plan/{job_id}").json()["plan"]
        run = client.post(
            "/api/stats/run-analysis",
            json={
                "job_id": job_id,
                "confirmed_test_ids": [item["id"] for item in plan["tests"]],
                "confirmed_graph_ids": [item["id"] for item in plan["graphs"]],
            },
        )
        assert run.status_code == 200, run.text
        first = run.json()
        first_id = first["result_id"]
        first_meta = first["results"]["export_metadata"]
        assert first_meta["analysis_version"] == 1
        assert first_meta["domain_profile"] == PROFILE

        # Simulate stale historical notes lingering in metadata; export must not show them.
        entry = dataset_store.get(job_id)
        entry.meta["cleanup_notes"] = {
            "pT": "Auto-extracted numeric values from text; converted to numeric and treated as scale.",
            "Nodal status": "Auto-extracted numeric values from text; converted to numeric and treated as scale.",
        }

        endpoints = {
            "word": f"/api/stats/export/{job_id}/word?result_id={first_id}",
            "pdf": f"/api/stats/export/{job_id}/pdf?result_id={first_id}",
            "excel": f"/api/stats/export/{job_id}/excel?result_id={first_id}",
            "chapter_word": f"/api/stats/export/{job_id}/chapter_v_word?result_id={first_id}",
            "chapter_pdf": f"/api/stats/export/{job_id}/chapter_v_pdf?result_id={first_id}",
        }
        exported = {name: client.get(url) for name, url in endpoints.items()}
        for response in exported.values():
            assert response.status_code == 200, response.text
            assert response.headers["x-medras-result-id"] == first_id
            assert response.headers["x-medras-analysis-version"] == "1"
            assert response.headers["x-medras-domain-profile"] == PROFILE

        for text in (
            _docx_text(exported["word"].content),
            _docx_text(exported["chapter_word"].content),
            _xlsx_text(exported["excel"].content),
        ):
            assert first_id in text
            assert PROFILE in text
            assert "T1c" in text and "T4b" in text
            assert "N0" in text and "N3a" in text
            assert "Auto-extracted numeric values from text" not in text
            assert "No primary inferential test ran successfully" not in text

        rerun = client.post(
            "/api/stats/run-analysis",
            json={
                "job_id": job_id,
                "confirmed_test_ids": [item["id"] for item in plan["tests"]],
                "confirmed_graph_ids": [item["id"] for item in plan["graphs"]],
            },
        )
        assert rerun.status_code == 200
        second = rerun.json()
        assert second["result_id"] != first_id
        assert second["results"]["export_metadata"]["analysis_version"] == 2
        assert second["results"]["export_metadata"]["generated_at"] != first_meta["generated_at"]
        stale = client.get(f"/api/stats/export/{job_id}/pdf?result_id={first_id}")
        assert stale.status_code == 409
        current = client.get(f"/api/stats/export/{job_id}/pdf?result_id={second['result_id']}")
        assert current.status_code == 200
        assert current.headers["x-medras-analysis-version"] == "2"


if __name__ == "__main__":
    verify_current_export_state()
    print("Sigma current export-state verification passed.")
