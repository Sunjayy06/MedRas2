"""Verify Sigma's canonical association-testing and reporting policy."""

from __future__ import annotations

import io

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from app.services import chapter_v_export, export, results
from app.services import plan as sigma_plan
from app.services.thesis_blueprint import build_thesis_analysis_blueprint
from test_fixtures.verify_sigma_chapter_v_export import FakeEntry


def _docx_text(blob: bytes) -> str:
    doc = Document(io.BytesIO(blob))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def test_contingency_selection_policy() -> None:
    sparse_2x2 = pd.DataFrame({
        "Predictor": ["A"] * 8 + ["B"] * 4,
        "Outcome": ["Yes"] * 8 + ["Yes"] + ["No"] * 3,
    })
    sparse = results.run_chi_or_fisher("Predictor", "Outcome", {}, sparse_2x2)
    assert sparse["actual_test_used"] == "Fisher's exact test"
    assert sparse["dof"] is None

    non_sparse_2x2 = pd.DataFrame({
        "Predictor": ["A"] * 20 + ["B"] * 20,
        "Outcome": (["Yes"] * 10 + ["No"] * 10) * 2,
    })
    adequate = results.run_chi_or_fisher("Predictor", "Outcome", {}, non_sparse_2x2)
    assert adequate["actual_test_used"] == "Chi-square test"
    assert adequate["dof"] == 1

    sparse_rxc = pd.DataFrame({
        "Predictor": ["A"] * 8 + ["B"] * 4 + ["C"] * 2,
        "Outcome": ["Yes"] * 8 + ["Yes"] + ["No"] * 3 + ["Yes", "No"],
    })
    larger = results.run_chi_or_fisher("Predictor", "Outcome", {}, sparse_rxc)
    assert larger["actual_test_used"] == "Chi-square test"
    assert larger["dof"] == 2
    assert "expected cell counts were below 5" in larger["note"]

    secondary_sparse = results._run_chi_or_fisher(sparse_2x2, "Predictor", "Outcome")
    secondary_adequate = results._run_chi_or_fisher(non_sparse_2x2, "Predictor", "Outcome")
    assert secondary_sparse["actual_test_used"] == "Fisher's exact test"
    assert secondary_adequate["actual_test_used"] == "Chi-square test"


def _canonical_tests() -> list[dict]:
    return [
        {
            "id": "predictor_a",
            "title": "Predictor A vs Outcome: Chi-square test",
            "analysis_family": "bivariate",
            "test_type": "chi_square",
            "actual_test_used": "Chi-square test",
            "statistic": 6.25,
            "dof": 1,
            "p": 0.012,
            "p_corrected": 0.072,
            "cramers_v": 0.31,
            "note": "-",
        },
        {
            "id": "predictor_b",
            "title": "Predictor B vs Outcome: Chi-square test",
            "analysis_family": "bivariate",
            "test_type": "chi_square",
            "actual_test_used": "Chi-square test",
            "statistic": 0.42,
            "dof": 1,
            "p": 0.517,
            "p_corrected": 0.710,
            "cramers_v": 0.08,
            "note": "-",
        },
        {
            "id": "predictor_c",
            "title": "Predictor C vs Outcome: Fisher's exact test",
            "analysis_family": "bivariate",
            "test_type": "fisher_exact",
            "actual_test_used": "Fisher's exact test",
            "statistic": 4.2,
            "p": 0.002,
            "p_corrected": 0.012,
            "cramers_v": 0.39,
            "note": "Fisher's exact test was used because some expected cell counts were below 5.",
        },
    ]


def _payload() -> dict:
    tests = _canonical_tests()
    associations = results._tested_associations(tests, "Outcome")
    significant = results._significant_findings(tests)
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(100, 4),
        classifications=[
            {"column": "Outcome", "detected_type": "nominal"},
            {"column": "Predictor A", "detected_type": "nominal"},
            {"column": "Predictor B", "detected_type": "nominal"},
            {"column": "Predictor C", "detected_type": "nominal"},
        ],
        assignment={"outcome": "Outcome"},
        plan={"suggestions": [
            {"warning": "Add only after confirming predictors"},
            {"warning": "Run separately under the Correlation objective"},
        ]},
        tests=tests,
        tested_associations=associations,
        significant_findings=significant,
        methods_text=(
            "Multiple-testing adjustment used the Benjamini-Hochberg FDR method across "
            "three inferential tests."
        ),
        session={"study_type": "association", "outcome": "Outcome"},
    )
    return {
        "tests": tests,
        "tested_associations": associations,
        "significant_findings": significant,
        "thesis_analysis_blueprint": blueprint,
    }


def test_canonical_summary_and_adjusted_status() -> None:
    payload = _payload()
    canonical = {row["source_result_id"]: row for row in payload["tested_associations"]}
    projected = {
        row["source_result_id"]: row
        for row in payload["thesis_analysis_blueprint"]["tested_associations"]
    }
    assert canonical.keys() == projected.keys()
    for result_id, row in canonical.items():
        for key in ("test_applied", "test_statistic", "p_value", "adjusted_p_value", "effect_size"):
            assert projected[result_id][key] == row[key]
    assert canonical["predictor_a"]["significance_status"] == (
        "Nominally significant before adjustment, not significant after multiple-testing correction."
    )
    assert canonical["predictor_b"]["significance_status"] == "Not significant after multiple-testing correction"
    assert canonical["predictor_c"]["significance_status"] == "Significant after multiple-testing correction"
    assert canonical["predictor_a"]["notes_warnings"] == (
        "Nominal before adjustment; not significant after correction."
    )


def test_section_vi_and_excel_report_all_associations() -> None:
    payload = _payload()
    docx_text = _docx_text(chapter_v_export.generate_docx(payload))
    assert "Summary of tested associations" in docx_text
    assert "Predictor A" in docx_text and "Predictor B" in docx_text and "Predictor C" in docx_text
    assert "Nominally significant before adjustment" in docx_text
    assert "Significant Findings Highlight" in docx_text
    assert "Notes" in docx_text
    assert "Add only after confirming predictors" not in docx_text
    assert "Run separately under the Correlation objective" not in docx_text
    # Issue 4 (thesis-readability polish, batch 2): objective-routing
    # suggestions are not clinically actionable for a thesis reader and
    # must not appear in the main Word/PDF report at all (moved to the
    # Excel audit log instead — see test_objective_routing_warnings_move_to_excel_audit_log_only).
    assert "A multivariable model was not added because predictor selection was not confirmed." not in docx_text
    assert "Correlation analysis was not included in the selected analysis objective." not in docx_text

    workbook = load_workbook(io.BytesIO(export.to_xlsx(FakeEntry(), payload, {"outcome": "Outcome"})))
    assert "tested_associations" in workbook.sheetnames
    rows = list(workbook["tested_associations"].iter_rows(min_row=2, values_only=True))
    assert len(rows) == 3
    excel_by_predictor = {row[0]: row for row in rows}
    for association in payload["tested_associations"]:
        row = excel_by_predictor[association["predictor"]]
        assert row[1] == association["test_applied"]
        assert row[2] == association["test_statistic"]
        assert row[3] == association["p_value"]
        assert row[4] == association["adjusted_p_value"]
        assert row[5] == association["effect_size"]


def test_nominal_egfr_and_marker_component_reporting() -> None:
    tests = [
        {
            "id": "egfr",
            "title": "EGFR vs p27 expression status: Chi-square test",
            "analysis_family": "bivariate",
            "test_type": "chi_square",
            "actual_test_used": "Chi-square test",
            "statistic": 4.1,
            "dof": 1,
            "p": 0.043,
            "p_corrected": 0.081,
            "correction_method": "Benjamini-Hochberg FDR",
            "cramers_v": 0.19,
            "note": "-",
        },
        {
            "id": "marker_component",
            "title": "Interpretation-site vs p27 expression status: Chi-square test with sparse-cell warning",
            "analysis_family": "bivariate",
            "test_type": "chi_square",
            "actual_test_used": "Chi-square test",
            "statistic": 5.2,
            "dof": 2,
            "p": 0.021,
            "p_corrected": 0.042,
            "cramers_v": 0.24,
            "note": "This finding should be interpreted cautiously because some expected cell counts were below 5.",
        },
    ]
    associations = results._tested_associations(tests, "p27 expression status")
    methods_text = (
        "Categorical associations were tested using chi-square or Fisher's exact test as appropriate. "
        "Multiple-testing adjustment used the Benjamini-Hochberg FDR method across 24 inferential tests; "
        "adjusted p-values were used to determine significance after correction. "
        "All tests are two-sided with a significance threshold of 0.05."
    )
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(100, 3),
        classifications=[
            {"column": "p27 expression status", "detected_type": "nominal"},
            {"column": "EGFR", "detected_type": "nominal"},
            {"column": "Interpretation-site", "detected_type": "nominal"},
        ],
        assignment={"outcome": "p27 expression status"},
        tests=tests,
        tested_associations=associations,
        methods_text=methods_text,
        session={
            "study_type": "association",
            "main_marker": "p27",
            "main_outcome_concept": "p27 expression status",
        },
    )
    assert [row["predictor"] for row in blueprint["tested_associations"]] == ["EGFR"]
    assert any("Marker-component variables were summarized descriptively" in warning for warning in blueprint["warnings"])
    egfr_interpretation = next(
        table["interpretation"]
        for section in blueprint["analysis_sections"] if section["section_id"] == "bivariate_associations"
        for table in section["tables"] if "EGFR" in table["title"]
    )
    assert (
        "EGFR was nominally significant before adjustment, but this did not remain significant "
        "after Benjamini-Hochberg FDR correction."
    ) in egfr_interpretation

    # Issue 2: the methods note must explain why the FDR test count (24)
    # exceeds the number of clinically interpreted rows shown below it.
    assert "across 24 inferential tests, including detailed marker-component tests." in blueprint["methods_text"]
    assert (
        "marker-component variables were summarized descriptively and excluded from "
        "clinical association interpretation" in blueprint["methods_text"]
    )

    text = _docx_text(chapter_v_export.generate_docx({"thesis_analysis_blueprint": blueprint, "tests": tests}))
    assert "Chi-square test with sparse-cell Grade" not in text
    assert "Chi-square test with sparse-cell The distribution" not in text
    assert "Interpretation-site" not in text.split("Section VI - Summary of Tested Associations", 1)[1]
    assert "Marker-component variables were summarized descriptively" in text
    assert "across 24 inferential tests, including detailed marker-component tests." in text


def test_run_plan_projects_each_canonical_result_once() -> None:
    df = pd.DataFrame({
        "Outcome": ["Yes"] * 30 + ["No"] * 30,
        "Predictor A": ["Positive"] * 25 + ["Negative"] * 5 + ["Positive"] * 10 + ["Negative"] * 20,
        "Predictor B": (["Positive", "Negative"] * 15) * 2,
        "Predictor C": ["Positive"] * 28 + ["Negative"] * 2 + ["Positive"] * 12 + ["Negative"] * 18,
    })
    classifications = [
        {"column": column, "detected_type": "nominal"}
        for column in df.columns
    ]
    plan_tests = []
    for predictor in ("Predictor A", "Predictor B", "Predictor C"):
        test_id = predictor.lower().replace(" ", "_")
        plan_tests.append({
            "id": test_id,
            "title": f"{predictor} vs Outcome",
            "analysis_family": "bivariate",
            "_phase_b": {
                "function": "run_chi_or_fisher",
                "test_type": "chi_square",
                "args": {"col1": predictor, "col2": "Outcome"},
            },
        })
    output = results.run_plan(
        df,
        classifications,
        {"outcome": "Outcome", "group": None, "covariates": []},
        {"tests": plan_tests, "graphs": [], "debug": {"eligible_predictor_count": 3, "bivariate_test_count": 3}},
        session={"study_type": "association", "domain_profile": "generic"},
    )
    assert output["correction_info"]["method"] == "Bonferroni"
    assert "Multiple-testing adjustment used the Bonferroni method" in output["methods_md"]
    canonical = {test["id"]: test for test in output["tests"]}
    summaries = {row["source_result_id"]: row for row in output["tested_associations"]}
    projected = {
        row["source_result_id"]: row
        for row in output["thesis_analysis_blueprint"]["tested_associations"]
    }
    assert canonical.keys() == summaries.keys() == projected.keys()
    for result_id, test in canonical.items():
        summary = summaries[result_id]
        assert summary["test_applied"] == results._test_used(test)
        assert summary["test_statistic"] == results._statistic_text(test)
        assert summary["effect_size"] == results._effect_text(test)
        assert summary["p_numeric"] == test["p"]
        assert summary["p_corrected_numeric"] == test["p_corrected"]
        assert projected[result_id]["p_value"] == summary["p_value"]
        assert projected[result_id]["adjusted_p_value"] == summary["adjusted_p_value"]
    bivariate = next(
        section for section in output["thesis_analysis_blueprint"]["analysis_sections"]
        if section["section_id"] == "bivariate_associations"
    )
    table_ids = {
        source_id
        for table in bivariate["tables"]
        for source_id in table.get("source_test_ids") or []
    }
    figure_ids = {figure.get("source_result_id") for figure in bivariate["figures"]}
    assert canonical.keys() <= table_ids
    assert canonical.keys() <= figure_ids


def test_duplicate_predictor_warning_is_specific_or_dropped() -> None:
    """Issue 5: a duplicate-predictor-labels warning must name the actual
    predictor and the raw labels that were merged, instead of a generic
    'category grouping should be reviewed' sentence that doesn't say which
    variable is involved. When the predictor cannot be identified, the
    warning must not appear in the main report at all."""
    plan_with_specific_suggestion = {
        "suggestions": [{
            "id": "predictor_duplicate_labels_Ki67",
            "title": "Review likely duplicate labels in predictor Ki67",
            "requires_confirmation": True,
            "blocking": False,
            "warning": (
                "Likely duplicate predictor labels were detected: >=14 / >= 14%. "
                "Analysis can continue, but split categories may affect estimates and should be reviewed."
            ),
        }],
    }
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(50, 2),
        classifications=[{"column": "Positive/ Negative", "detected_type": "nominal"}],
        assignment={"outcome": "Positive/ Negative"},
        tests=[],
        plan=plan_with_specific_suggestion,
        session={"study_type": "association"},
    )
    assert any("Duplicate raw Ki67 category labels" in warning for warning in blueprint["warnings"]), blueprint["warnings"]
    assert any(">=14 / >= 14%" in warning for warning in blueprint["warnings"]), blueprint["warnings"]
    assert not any("should be reviewed" in warning for warning in blueprint["warnings"]), blueprint["warnings"]

    text = _docx_text(chapter_v_export.generate_docx({"thesis_analysis_blueprint": blueprint, "tests": []}))
    assert "Duplicate raw Ki67 category labels" in text
    assert "Some predictor labels appeared duplicated after cleaning" not in text

    # A suggestion id that doesn't match the expected pattern can't be
    # attributed to a specific predictor — it must be dropped, not shown
    # as a generic, unactionable warning.
    plan_without_id = {
        "suggestions": [{
            "id": "unrelated_suggestion_id",
            "warning": (
                "Likely duplicate predictor labels were detected: A / B. "
                "Analysis can continue, but split categories may affect estimates and should be reviewed."
            ),
        }],
    }
    dropped_blueprint = build_thesis_analysis_blueprint(
        df_shape=(50, 2),
        classifications=[{"column": "Positive/ Negative", "detected_type": "nominal"}],
        assignment={"outcome": "Positive/ Negative"},
        tests=[],
        plan=plan_without_id,
        session={"study_type": "association"},
    )
    assert not any("duplicate" in warning.lower() for warning in dropped_blueprint["warnings"]), dropped_blueprint["warnings"]


def test_thesis_conservative_exact_mode() -> None:
    """Issue 2: an opt-in, default-OFF thesis-conservative mode should use
    Fisher's exact test for every 2x2 clinical categorical association
    (even with adequate expected counts), while leaving larger RxC tables
    and marker-component descriptive columns on the existing chi-square
    routing. Verify the toggle, the routing, and Word/PDF/Excel consistency."""
    import app.core.config as config_module

    assert config_module.settings.sigma_thesis_conservative_exact is False, (
        "SIGMA_THESIS_CONSERVATIVE_EXACT must default to OFF"
    )

    adequate_2x2 = pd.DataFrame({
        "Predictor": ["A"] * 20 + ["B"] * 20,
        "Outcome": (["Yes"] * 10 + ["No"] * 10) * 2,
    })
    default_result = results.run_chi_or_fisher("Predictor", "Outcome", {}, adequate_2x2)
    assert default_result["actual_test_used"] == "Chi-square test"

    conservative_result = results.run_chi_or_fisher(
        "Predictor", "Outcome", {"thesis_conservative_exact": True}, adequate_2x2
    )
    assert conservative_result["actual_test_used"] == "Fisher's exact test"
    assert "thesis-conservative mode" in conservative_result["note"]

    rxc = pd.DataFrame({
        "Predictor": ["A"] * 15 + ["B"] * 15 + ["C"] * 15,
        "Outcome": (["Yes"] * 8 + ["No"] * 7) * 3,
    })
    rxc_conservative = results.run_chi_or_fisher(
        "Predictor", "Outcome", {"thesis_conservative_exact": True}, rxc
    )
    assert rxc_conservative["actual_test_used"] == "Chi-square test", (
        "Conservative mode must not force Fisher's exact onto larger RxC tables"
    )

    marker_component_df = pd.DataFrame({
        "Interpretation-site": ["Nuclear"] * 20 + ["Cytoplasmic"] * 20,
        "Outcome": (["Yes"] * 10 + ["No"] * 10) * 2,
    })
    marker_conservative = results.run_chi_or_fisher(
        "Interpretation-site", "Outcome", {"thesis_conservative_exact": True}, marker_component_df
    )
    assert marker_conservative["actual_test_used"] == "Chi-square test", (
        "Conservative mode must not be forced onto marker-component descriptive columns"
    )

    # End-to-end: run_plan with the flag set must produce internally
    # consistent p-values/test labels across the test result, Word/PDF,
    # and Excel tested-associations output.
    df = pd.DataFrame({
        "Positive/ Negative": ["Yes"] * 20 + ["No"] * 20,
        "ER": (["Positive"] * 10 + ["Negative"] * 10) * 2,
    })
    classes = [
        {"column": "Positive/ Negative", "detected_type": "nominal"},
        {"column": "ER", "detected_type": "nominal"},
    ]
    session = {
        "study_type": "association",
        "study_type_confirmed": True,
        "analysis_predictors": ["ER"],
        "domain_profile": "breast_pathology",
        "thesis_conservative_exact": True,
    }
    assignment = {"outcome": "Positive/ Negative", "group": None, "covariates": []}
    sigma_plan_dict = sigma_plan.generate_plan(df, classes, assignment, {"columns": []}, session)
    output = results.run_plan(df, classes, assignment, sigma_plan_dict, session=session)
    er_test = next(test for test in output["tests"] if test.get("test_type") == "fisher_exact")
    assert "In thesis-conservative mode" in output["methods_md"]

    blueprint = output["thesis_analysis_blueprint"]
    docx_text = _docx_text(chapter_v_export.generate_docx({
        "thesis_analysis_blueprint": blueprint, "tests": output["tests"],
    }))
    pdf_text = chapter_v_export.generate_pdf({
        "thesis_analysis_blueprint": blueprint, "tests": output["tests"],
    })
    assert isinstance(pdf_text, bytes) and len(pdf_text) > 0

    er_row = next(row for row in output["tested_associations"] if row["predictor"] == "ER")
    assert er_row["test_applied"] == "Fisher's exact test"
    assert "Fisher's exact test" in docx_text

    workbook = load_workbook(io.BytesIO(export.to_xlsx(FakeEntry(), output, assignment)))
    excel_rows = list(workbook["tested_associations"].iter_rows(min_row=2, values_only=True))
    excel_er_row = next(row for row in excel_rows if row[0] == "ER")
    assert excel_er_row[1] == "Fisher's exact test" == er_row["test_applied"]
    assert excel_er_row[3] == er_row["p_value"] == output["tested_associations"][0]["p_value"]
    assert er_test["p"] is not None


def test_outcome_duplicate_and_generic_technical_warnings_are_specific_or_dropped() -> None:
    """Issue 5 (batch 2): the outcome-label counterpart of the predictor
    duplicate-labels warning must also be made specific (naming the
    outcome and the raw labels) rather than shown verbatim with internal
    phrasing like 'test routing may be invalid'. Generic internal
    QA/implementation-detail warnings (phase_b/objective_routing/etc.)
    must never reach the main Word/PDF report."""
    plan_with_outcome_dup = {
        "suggestions": [{
            "id": "outcome_duplicate_labels",
            "title": "Resolve likely duplicate labels in outcome Positive/ Negative",
            "requires_confirmation": True,
            "blocking": True,
            "warning": (
                "Likely duplicate outcome labels were detected: Postive / Positive. "
                "Merge or explicitly resolve them before running inferential analyses; "
                "otherwise group counts and test routing may be invalid."
            ),
        }],
    }
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(50, 2),
        classifications=[{"column": "Positive/ Negative", "detected_type": "nominal"}],
        assignment={"outcome": "Positive/ Negative"},
        tests=[],
        plan=plan_with_outcome_dup,
        session={"study_type": "association"},
    )
    assert any(
        "Duplicate raw Positive/ Negative category labels (Postive / Positive) were merged before analysis." == warning
        for warning in blueprint["warnings"]
    ), blueprint["warnings"]
    assert not any("test routing may be invalid" in warning for warning in blueprint["warnings"])

    text = _docx_text(chapter_v_export.generate_docx({"thesis_analysis_blueprint": blueprint, "tests": []}))
    assert "Duplicate raw Positive/ Negative category labels" in text
    assert "test routing may be invalid" not in text

    plan_with_internal_warning = {
        "suggestions": [{
            "id": "internal_thing",
            "title": "Internal QA Check",
            "warning": "Internal QA: phase_b trigger entry routing mismatch detected for objective_routing logic.",
        }],
    }
    dropped_blueprint = build_thesis_analysis_blueprint(
        df_shape=(50, 2),
        classifications=[{"column": "Positive/ Negative", "detected_type": "nominal"}],
        assignment={"outcome": "Positive/ Negative"},
        tests=[],
        plan=plan_with_internal_warning,
        session={"study_type": "association"},
    )
    assert dropped_blueprint["warnings"] == []


def test_objective_routing_warnings_move_to_excel_audit_log_only() -> None:
    """Issue 4 (thesis-readability polish, batch 2): 'A multivariable model
    was not added...' and 'Correlation analysis was not included...' are
    objective-routing/implementation-detail notes, not clinically
    actionable for a thesis reader. They must be dropped from the main
    Word/PDF report (test_section_vi_and_excel_report_all_associations
    covers that) while the raw suggestion text is preserved in the Excel
    'Data Cleaning Log' audit sheet."""
    payload = _payload()
    docx_text = _docx_text(chapter_v_export.generate_docx(payload))
    assert "A multivariable model was not added" not in docx_text
    assert "Correlation analysis was not included" not in docx_text

    class _EntryWithPlan:
        def __init__(self):
            self.df = FakeEntry().df
            self.meta = dict(FakeEntry().meta)
            self.meta["plan"] = {"suggestions": [
                {"warning": "Add only after confirming predictors and checking event counts and separation risk."},
                {"warning": "Run separately under the Correlation objective; it is not part of the main outcome association results."},
            ]}

    workbook = load_workbook(io.BytesIO(export.to_xlsx(_EntryWithPlan(), payload, {"outcome": "Outcome"})))
    log_rows = list(workbook["Data Cleaning Log"].iter_rows(min_row=2, values_only=True))
    log_text = " ".join(str(cell) for row in log_rows for cell in row if cell)
    assert "Add only after confirming predictors" in log_text
    assert "Run separately under the Correlation objective" in log_text


def test_percentage_denominator_note_appears_in_word_pdf_excel() -> None:
    """Issue 3 (thesis-readability polish): detailed association tables use
    within-predictor-category percentages while the Significant Findings
    Highlight uses within-p27-expression-group percentages. A clear
    footnote explaining this must appear near Section V/VI in Word and
    PDF, and on the Excel Cover sheet, whenever tested associations exist."""
    note = (
        "Percentages in detailed association tables are calculated within predictor categories "
        "unless otherwise stated. Percentages in the Significant Findings Highlight describe "
        "marker/category distribution within p27 expression groups."
    )
    tests = [
        {
            "id": "er", "title": "ER vs Positive/ Negative: Chi-square test",
            "analysis_family": "bivariate", "test_type": "chi_square",
            "actual_test_used": "Chi-square test", "statistic": 12.0, "dof": 1,
            "p": 0.0005, "p_corrected": 0.0015, "cramers_v": 0.55, "note": "-",
        },
    ]
    associations = results._tested_associations(tests, "Positive/ Negative")
    blueprint = build_thesis_analysis_blueprint(
        df_shape=(100, 2),
        classifications=[
            {"column": "Positive/ Negative", "detected_type": "nominal"},
            {"column": "ER", "detected_type": "nominal"},
        ],
        assignment={"outcome": "Positive/ Negative"},
        tests=tests,
        tested_associations=associations,
        session={"study_type": "association"},
    )
    payload = {"tests": tests, "tested_associations": associations, "thesis_analysis_blueprint": blueprint}
    docx_text = _docx_text(chapter_v_export.generate_docx(payload))
    assert note in docx_text

    pdf_bytes = chapter_v_export.generate_pdf(payload)
    assert isinstance(pdf_bytes, bytes) and len(pdf_bytes) > 0

    workbook = load_workbook(io.BytesIO(export.to_xlsx(FakeEntry(), payload, {"outcome": "Positive/ Negative"})))
    cover_rows = list(workbook["Cover"].iter_rows(values_only=True))
    assert any(note in str(cell) for row in cover_rows for cell in row if cell)


def main() -> None:
    test_contingency_selection_policy()
    test_canonical_summary_and_adjusted_status()
    test_section_vi_and_excel_report_all_associations()
    test_nominal_egfr_and_marker_component_reporting()
    test_run_plan_projects_each_canonical_result_once()
    test_duplicate_predictor_warning_is_specific_or_dropped()
    test_thesis_conservative_exact_mode()
    test_outcome_duplicate_and_generic_technical_warnings_are_specific_or_dropped()
    test_objective_routing_warnings_move_to_excel_audit_log_only()
    test_percentage_denominator_note_appears_in_word_pdf_excel()
    print("Sigma association reporting policy verification passed.")


if __name__ == "__main__":
    main()
