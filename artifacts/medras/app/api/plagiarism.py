"""Plagiarism & AI-reduction module — backend routes.

Provides the following endpoints under ``/api/plagiarism``:

  * ``POST /check``        — score originality + AI likelihood from raw text
  * ``POST /check-file``   — same, but accepts a PDF/DOCX/TXT upload
  * ``POST /reduce``       — rewrite text to read more human / less templated
  * ``POST /analyze-file`` — extract text from a PDF/DOCX/TXT, detect IMRaD
                             sections and protected technical terms, return
                             the breakdown plus the extracted text so the
                             frontend can run /check or /reduce afterwards
                             without re-uploading the file

All endpoints return JSON. The actual analysis is delegated to
``app.services.plagiarism_analyzer`` (LLM calls) and
``app.services.text_analyzer`` (regex-based section + term detection).
"""

from __future__ import annotations

import asyncio
import io
import json
import re
from typing import Any, Dict, List, Literal, Optional

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.core.limiter import limiter
from app.core.logging import get_logger
from app.services import (
    citation_suggester,
    plagiarism_analyzer,
    plagiarism_jobs,
    text_analyzer,
)

log = get_logger(__name__)

router = APIRouter(prefix="/plagiarism", tags=["plagiarism"])


# ---------------------------------------------------------------------------
# Request models & limits
# ---------------------------------------------------------------------------

# Hard cap on how much text we'll send to the LLM in a single call.
# A typical academic paragraph is ~150 words ≈ 1000 chars; 30 KB is roughly
# 5000 words. Larger inputs are truncated client-side OR analysed
# section-by-section using /analyze-file first.
MAX_TEXT_CHARS = 30_000

# Hard cap on uploaded file size. The user explicitly asked for 100 MB. We
# read the upload in 64 KB chunks and bail as soon as the running total
# exceeds the cap, so an attacker can't lie about Content-Length.
MAX_UPLOAD_BYTES = 100 * 1024 * 1024  # 100 MB

# How much extracted text we'll accept downstream of an upload. Even if a
# 500-page PDF extracts to 2 MB of plain text, we won't analyse all of it —
# section detection runs on the full body (cheap), but the LLM-bound text
# is capped at ANALYZE_TEXT_CHARS.
ANALYZE_TEXT_CHARS = 1_000_000  # 1 MB of UTF-8 text — plenty for a thesis


class CheckRequest(BaseModel):
    text: str = Field(..., min_length=1)
    provider: Literal["openai", "gemini", "auto"] = "auto"


class PipelineSectionIn(BaseModel):
    """One section to feed into the 3-stage rewrite pipeline.

    The label drives References-skip detection (sections labelled
    "References", "Bibliography", "Works Cited", "Literature Cited" are
    kept verbatim and not rewritten).
    """
    label: str = Field(..., min_length=1, max_length=200)
    text: str = Field(..., min_length=1)


class ReduceRequest(BaseModel):
    text: str = Field(..., min_length=1)
    provider: Literal["openai", "gemini", "auto"] = "auto"
    # Strings that MUST appear unchanged in the rewrite. Usually populated
    # from the output of /analyze-file's protected_terms list.
    protected_terms: Optional[List[str]] = None
    # If supplied, /reduce runs the 3-stage GPT-4o → Gemini → GPT-4o
    # pipeline per section instead of the legacy single-shot rewrite.
    # ``text`` is still required (used for size validation and as the
    # legacy fallback if the pipeline raises).
    sections: Optional[List[PipelineSectionIn]] = None
    # Force the pipeline path even without explicit sections (treats the
    # full text as one body section). Useful for the paste-text flow
    # when the user wants the higher-quality multi-stage rewrite.
    pipeline: bool = False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _read_with_cap(file: UploadFile, cap: int) -> bytes:
    """Read an UploadFile in chunks, raising 413 as soon as ``cap`` is hit.

    ``await file.read()`` would happily slurp a 200 MB body into memory
    before any size check runs. Reading in 64 KB chunks lets us reject
    abusive uploads early with a clear error.
    """
    chunks: list[bytes] = []
    total = 0
    chunk_size = 64 * 1024
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        total += len(chunk)
        if total > cap:
            raise HTTPException(
                status_code=413,
                detail=f"File exceeds {cap // (1024 * 1024)} MB. Got at least {total} bytes.",
            )
        chunks.append(chunk)
    return b"".join(chunks)


def _safe_extract(filename: str, content: bytes) -> str:
    """Wrap extract_text_from_upload so each known failure mode becomes
    a clean HTTPException with a user-friendly message.

    The catch-all branch sanitises the exception text so an unexpected
    provider/library error can never include an API key.
    """
    try:
        return plagiarism_analyzer.extract_text_from_upload(filename or "", content)
    except plagiarism_analyzer.UploadExtractionError as exc:
        raise HTTPException(status_code=exc.http_status, detail=str(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        log.exception("file extraction failed")
        raise HTTPException(
            status_code=400,
            detail=f"Could not read file: {plagiarism_analyzer.sanitize_error_message(exc)}",
        ) from exc


def _normalise_provider(value: str | None) -> Literal["openai", "gemini", "auto"]:
    if value not in ("openai", "gemini", "auto"):
        return "auto"
    return value  # type: ignore[return-value]


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@router.post("/check")
@limiter.limit("20/minute")
async def check_text(request: Request, payload: CheckRequest) -> dict:
    """Score originality + AI likelihood for pasted or pre-extracted text."""
    text = payload.text.strip()
    if len(text) > MAX_TEXT_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"Text is too long ({len(text):,} chars). Maximum is {MAX_TEXT_CHARS:,}.",
        )
    try:
        result = plagiarism_analyzer.check_originality(text, provider=payload.provider)
    except RuntimeError as exc:
        # Missing API key — surface as 503 so the UI can show a helpful note.
        raise HTTPException(status_code=503, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        log.exception("plagiarism check failed")
        raise HTTPException(status_code=502, detail=f"Analysis failed: {plagiarism_analyzer.sanitize_error_message(exc)}") from exc
    return result


@router.post("/check-file")
@limiter.limit("10/minute")
async def check_file(
    request: Request,
    file: UploadFile = File(...),
    provider: str = Form("auto"),
) -> dict:
    """Same as /check but accepts a PDF/DOCX/TXT file upload (≤100 MB, ≤200 pages)."""
    provider_choice = _normalise_provider(provider)
    content = await _read_with_cap(file, MAX_UPLOAD_BYTES)
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    filename = file.filename
    text = _safe_extract(filename or "", content)
    # Free the raw upload bytes from memory as soon as extraction is done —
    # we never persist uploads to disk and we no longer need them after this.
    del content

    text = (text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="No readable text found in the file.")
    truncated = False
    if len(text) > MAX_TEXT_CHARS:
        # Trim rather than reject — the UI tells the user we analysed the
        # first N characters. Truncating gives a useful answer instead of
        # an error for a long PDF.
        text = text[:MAX_TEXT_CHARS]
        truncated = True

    try:
        result = plagiarism_analyzer.check_originality(text, provider=provider_choice)
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        log.exception("plagiarism file check failed")
        raise HTTPException(
            status_code=502,
            detail=f"Analysis failed: {plagiarism_analyzer.sanitize_error_message(exc)}",
        ) from exc

    analysed_chars = len(text)
    del text  # text is now embedded in `result`; drop the local reference
    result["filename"] = filename
    result["analysed_chars"] = analysed_chars
    result["truncated"] = truncated
    return result


@router.post("/analyze-file")
@limiter.limit("10/minute")
async def analyze_file(
    request: Request,
    file: UploadFile = File(...),
) -> dict:
    """Extract text from an upload and return the IMRaD breakdown.

    No LLM is called here — this endpoint is fast and cheap, and is the
    "step 1" the UI runs after a user picks a file. The response includes:

      * total word & character counts
      * detected sections with word counts and short previews
      * detected protected terms (drug names, p-values, citations, etc.)
      * the extracted plain text (capped at ANALYZE_TEXT_CHARS) so the UI
        can pass it back to /check or /reduce without re-uploading
    """
    content = await _read_with_cap(file, MAX_UPLOAD_BYTES)
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    filename = file.filename or ""
    name = filename.lower()
    if not (name.endswith(".pdf") or name.endswith(".docx") or name.endswith(".txt") or name.endswith(".md")):
        raise HTTPException(
            status_code=415,
            detail="Unsupported file type. Please upload a .pdf, .docx, .txt or .md file.",
        )

    size_bytes = len(content)
    text = _safe_extract(filename, content)
    # Drop the raw upload from memory as soon as we've turned it into text.
    del content

    text = (text or "").strip()
    if not text:
        raise HTTPException(
            status_code=400,
            detail="No readable text found in the file.",
        )

    truncated = False
    if len(text) > ANALYZE_TEXT_CHARS:
        text = text[:ANALYZE_TEXT_CHARS]
        truncated = True

    try:
        breakdown = text_analyzer.analyze_document(text)
    except Exception as exc:  # noqa: BLE001
        log.exception("plagiarism analyze-file detection failed")
        raise HTTPException(
            status_code=500,
            detail=f"Document analysis failed: {plagiarism_analyzer.sanitize_error_message(exc)}",
        ) from exc

    return {
        "filename": filename,
        "size_bytes": size_bytes,
        "extracted_chars": len(text),
        "truncated": truncated,
        "extracted_text": text,
        **breakdown,
    }


@router.post("/reduce")
@limiter.limit("10/minute")
async def reduce_text(request: Request, payload: ReduceRequest) -> dict:
    """Rewrite text to read more human and less templated.

    Two paths:

    1. **Pipeline path** — taken when ``sections`` is supplied OR
       ``pipeline=true``. Runs each section through 3 LLM stages:
       paraphrase (gpt-4o) → humanise (gemini-2.5-flash) → polish
       (gpt-4o). Sections labelled References / Bibliography / Works
       Cited / Literature Cited are skipped entirely and kept verbatim
       in the combined output. The ``provider`` field is ignored on this
       path — each stage has a fixed primary with the other provider as
       its automatic fallback.

    2. **Legacy single-shot path** — original behaviour. Used when
       neither ``sections`` nor ``pipeline`` is set.

    In both paths, ``protected_terms`` substrings are passed to the LLM
    as hard "do not change" constraints AND verified post-hoc.
    """
    text = payload.text.strip()
    if len(text) > MAX_TEXT_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"Text is too long ({len(text):,} chars). Maximum is {MAX_TEXT_CHARS:,}.",
        )

    use_pipeline = bool(payload.sections) or payload.pipeline
    if use_pipeline:
        sections_payload: list[dict[str, str]]
        if payload.sections:
            total = sum(len(s.text) for s in payload.sections)
            if total > MAX_TEXT_CHARS:
                raise HTTPException(
                    status_code=413,
                    detail=f"Combined sections are too long ({total:,} chars). Maximum is {MAX_TEXT_CHARS:,}.",
                )
            sections_payload = [{"label": s.label, "text": s.text} for s in payload.sections]
        else:
            sections_payload = [{"label": "Body", "text": text}]
        try:
            result = plagiarism_analyzer.rewrite_pipeline(
                sections_payload,
                protected_terms=payload.protected_terms,
            )
        except plagiarism_analyzer.ProviderQuotaExhausted as exc:
            # Both AI providers are out of quota. Surface a clear,
            # actionable 503 the UI can show instead of a 502 stack trace.
            raise HTTPException(
                status_code=503,
                detail=(
                    "Both AI providers are out of quota right now. "
                    "OpenAI returned insufficient_quota (billing) and Gemini "
                    "hit its free-tier daily request limit. Please try again "
                    "tomorrow, or top up one of the provider accounts."
                ),
            ) from exc
        except RuntimeError as exc:
            raise HTTPException(status_code=503, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
        except Exception as exc:  # noqa: BLE001
            log.exception("plagiarism rewrite_pipeline failed")
            raise HTTPException(status_code=502, detail=f"Rewrite pipeline failed: {plagiarism_analyzer.sanitize_error_message(exc)}") from exc
        return result

    try:
        result = plagiarism_analyzer.reduce_plagiarism(
            text,
            provider=payload.provider,
            protected_terms=payload.protected_terms,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=plagiarism_analyzer.sanitize_error_message(exc)) from exc
    except Exception as exc:  # noqa: BLE001
        log.exception("plagiarism reduce failed")
        raise HTTPException(status_code=502, detail=f"Rewrite failed: {plagiarism_analyzer.sanitize_error_message(exc)}") from exc
    return result


# ---------------------------------------------------------------------------
# Streaming pipeline (NDJSON) — used by the dedicated reduce-results page so
# the user sees a section-by-section progress bar instead of staring at a
# spinner for 30-90s.
# ---------------------------------------------------------------------------

class ReduceStreamRequest(BaseModel):
    """Same shape as ReduceRequest minus the legacy ``provider`` field.

    The streaming endpoint always uses the 3-stage pipeline; it would be
    nonsense to stream a single-shot legacy call.
    """
    text: Optional[str] = None
    sections: Optional[List[PipelineSectionIn]] = None
    protected_terms: Optional[List[str]] = None


@router.post("/reduce-stream")
@limiter.limit("10/minute")
async def reduce_stream(request: Request, payload: ReduceStreamRequest):
    """Run the 3-stage rewrite pipeline and stream NDJSON progress events.

    Each line of the response body is one JSON event:

      ``{"type": "init", "total_sections": N, "sections": [...]}``
      ``{"type": "section_start", "index": i, "label": "..."}``
      ``{"type": "stage_done", "index": i, "stage": "a"|"b"|"c", "model": "..."}``
      ``{"type": "section_done", "section": {...}}``
      ``{"type": "complete", "result": {...full pipeline response...}}``
      ``{"type": "error", "status": int, "message": "..."}``

    The frontend reads the response with ``fetch`` + a streaming reader and
    advances a progress bar on each ``stage_done`` event.

    Validation mirrors ``/reduce``: provide either ``sections`` or ``text``.
    """
    sections_payload: list[dict[str, str]]
    if payload.sections:
        total = sum(len(s.text) for s in payload.sections)
        if total > MAX_TEXT_CHARS:
            raise HTTPException(
                status_code=413,
                detail=f"Combined sections are too long ({total:,} chars). Maximum is {MAX_TEXT_CHARS:,}.",
            )
        sections_payload = [{"label": s.label, "text": s.text} for s in payload.sections]
    elif payload.text:
        text = payload.text.strip()
        if not text:
            raise HTTPException(status_code=400, detail="Text is empty.")
        if len(text) > MAX_TEXT_CHARS:
            raise HTTPException(
                status_code=413,
                detail=f"Text is too long ({len(text):,} chars). Maximum is {MAX_TEXT_CHARS:,}.",
            )
        sections_payload = [{"label": "Body", "text": text}]
    else:
        raise HTTPException(status_code=400, detail="Provide either 'text' or 'sections'.")

    protected_terms = payload.protected_terms or []

    async def event_generator():
        import threading
        loop = asyncio.get_running_loop()
        # Bounded queue so a runaway worker can't grow memory unboundedly
        # if the consumer falls behind. With per-section events this is
        # plenty of headroom (init + N×(start + 3 stage + done) + complete).
        queue: asyncio.Queue = asyncio.Queue(maxsize=256)
        cancel_event = threading.Event()
        SENTINEL = object()

        def progress_cb(event: dict) -> None:
            # Called from the worker thread; hop onto the loop safely.
            # If the queue is full, drop the event rather than block the
            # worker — the progress UI is best-effort, not delivery-critical.
            def _put():
                try:
                    queue.put_nowait(event)
                except asyncio.QueueFull:
                    log.warning("reduce-stream queue full, dropping event %s", event.get("type"))
            loop.call_soon_threadsafe(_put)

        async def runner():
            try:
                await asyncio.to_thread(
                    plagiarism_analyzer.rewrite_pipeline,
                    sections_payload,
                    protected_terms=protected_terms,
                    progress_cb=progress_cb,
                    cancel_event=cancel_event,
                )
            except plagiarism_analyzer.PipelineCancelled:
                # Client disconnected — pipeline cooperatively aborted.
                # Don't push an error; the consumer is gone anyway.
                log.info("reduce-stream pipeline aborted (client disconnect)")
            except plagiarism_analyzer.ProviderQuotaExhausted as exc:
                # Log the raw provider message for ops; never forward it to
                # the client. Only the fixed user-safe message goes out.
                log.warning("reduce-stream both providers exhausted: %s",
                            plagiarism_analyzer.sanitize_error_message(exc))
                loop.call_soon_threadsafe(queue.put_nowait, {
                    "type": "error",
                    "status": 503,
                    "message": (
                        "Both AI providers are out of quota right now. "
                        "OpenAI returned insufficient_quota (billing) and Gemini "
                        "hit its free-tier daily request limit. Please try again "
                        "tomorrow, or top up one of the provider accounts."
                    ),
                })
            except ValueError as exc:
                loop.call_soon_threadsafe(queue.put_nowait, {
                    "type": "error", "status": 400, "message": plagiarism_analyzer.sanitize_error_message(exc),
                })
            except RuntimeError as exc:
                loop.call_soon_threadsafe(queue.put_nowait, {
                    "type": "error", "status": 503, "message": plagiarism_analyzer.sanitize_error_message(exc),
                })
            except Exception as exc:  # noqa: BLE001
                log.exception("reduce-stream pipeline crashed")
                loop.call_soon_threadsafe(queue.put_nowait, {
                    "type": "error", "status": 502, "message": f"Rewrite pipeline failed: {plagiarism_analyzer.sanitize_error_message(exc)}",
                })
            finally:
                loop.call_soon_threadsafe(queue.put_nowait, SENTINEL)

        worker_task = asyncio.create_task(runner())

        async def _disconnect_watcher():
            # Poll for client disconnect every 2s. When detected, signal
            # the worker to bail at its next stage boundary so we stop
            # burning LLM credits on a request nobody is listening to.
            try:
                while not cancel_event.is_set() and not worker_task.done():
                    if await request.is_disconnected():
                        log.info("reduce-stream client disconnected; signalling cancel")
                        cancel_event.set()
                        return
                    await asyncio.sleep(2.0)
            except asyncio.CancelledError:
                pass

        watcher_task = asyncio.create_task(_disconnect_watcher())
        try:
            while True:
                try:
                    event = await asyncio.wait_for(queue.get(), timeout=30.0)
                except asyncio.TimeoutError:
                    # Heartbeat — empty newline keeps proxies from buffering
                    # during the long Stage-A LLM call. NDJSON parsers
                    # ignore blank lines.
                    yield "\n"
                    continue
                if event is SENTINEL:
                    break
                yield json.dumps(event, ensure_ascii=False) + "\n"
        finally:
            cancel_event.set()  # ensure worker bails on its next checkpoint
            watcher_task.cancel()
            # We can't kill the worker thread, but the cancel_event will
            # cause it to exit at the next stage boundary (within one LLM
            # call's worth of time, ~5-30s). The asyncio task will then
            # complete on its own; we don't await it because that would
            # hold the response open.
            if not worker_task.done():
                worker_task.cancel()

    return StreamingResponse(
        event_generator(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",  # disable proxy buffering
        },
    )


# ---------------------------------------------------------------------------
# Job-based async pipeline (POST + poll, no long-lived connections)
#
# This replaces /reduce-stream for the frontend. The streaming endpoint
# above is left in place for any direct/legacy consumer; both share the
# same underlying analyzer functions.
#
# Lifecycle:
#   1. POST /jobs              → returns {job_id, total_sections}
#   2. GET  /jobs/{job_id}     → poll every ~5s; returns full snapshot
#   3. POST /jobs/{job_id}/retry  → re-queues just the failed sections
#   4. DELETE /jobs/{job_id}   → cancel a running job or evict a finished one
# ---------------------------------------------------------------------------


class JobReportIn(BaseModel):
    """Optional plagiarism-report metadata supplied by the new intake flow.

    Drives per-section rewrite intensity so we don't burn tokens
    re-paraphrasing sections that already came back clean from the
    user's plagiarism checker (Turnitin / Drillbit / etc.).
    """
    software: Optional[str] = Field(default=None, max_length=80)
    # flagged_map keys are normalised section names (lower-case);
    # values look like {"similarity_percent": 34.0, "flagged": True}.
    flagged_map: Dict[str, Dict[str, Any]] = Field(default_factory=dict)


class JobCreateRequest(BaseModel):
    """Body for POST /jobs.

    Either ``sections`` (preferred — IMRaD breakdown from /analyze-file)
    or a single ``text`` blob can be supplied. ``text`` is wrapped in a
    one-section list so the pipeline always sees a list.
    """
    sections: Optional[List[PipelineSectionIn]] = None
    text: Optional[str] = None
    protected_terms: List[str] = Field(default_factory=list)
    title: str = Field(default="Rewritten document", max_length=200)
    filename: Optional[str] = Field(default=None, max_length=200)
    report: Optional[JobReportIn] = None


@router.post("/jobs")
@limiter.limit("10/minute")
async def create_job(request: Request, payload: JobCreateRequest) -> dict:
    """Register a new background rewrite job and return immediately.

    Does NOT wait for any sections to complete. The browser polls
    GET /jobs/{job_id} every few seconds to learn how it's going.
    """
    if payload.sections:
        sections = [{"label": s.label, "text": s.text} for s in payload.sections]
    elif payload.text and payload.text.strip():
        sections = [{"label": "Body", "text": payload.text}]
    else:
        raise HTTPException(
            status_code=400,
            detail="Provide either 'sections' (preferred) or 'text'.",
        )

    # Cap input size — prevents a single huge upload from monopolising
    # the in-memory budget. The 200-page PDF cap upstream already bounds
    # this, but a hand-crafted JSON request could bypass that path.
    total_chars = sum(len(s.get("text") or "") for s in sections)
    if total_chars > ANALYZE_TEXT_CHARS:
        raise HTTPException(
            status_code=413,
            detail=(
                f"Combined section text is {total_chars:,} characters, exceeding "
                f"the {ANALYZE_TEXT_CHARS:,}-character per-job cap. Please split "
                "into smaller documents."
            ),
        )

    try:
        state = plagiarism_jobs.job_manager.create_job(
            sections=sections,
            protected_terms=payload.protected_terms,
            title=payload.title,
            filename=payload.filename,
            report=payload.report.model_dump() if payload.report else None,
        )
    except plagiarism_jobs.CapacityError as exc:
        # 429 = Too Many Requests — the request is well-formed but the
        # server is at its concurrent-job ceiling.
        raise HTTPException(status_code=429, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail=plagiarism_analyzer.sanitize_error_message(exc),
        ) from exc

    return {
        "job_id": state.job_id,
        "total_sections": len(state.sections),
        "status": state.status,
    }


@router.get("/jobs/{job_id}")
@limiter.limit("120/minute")  # 2/sec is generous vs. the 5s spec
async def get_job(request: Request, job_id: str) -> dict:
    """Return the current snapshot of a job. Safe to poll every ~5s."""
    j = plagiarism_jobs.job_manager.get_job(job_id)
    if not j:
        raise HTTPException(
            status_code=404,
            detail="Job not found or expired (jobs are kept for 30 minutes).",
        )
    return plagiarism_jobs.serialize_job(j)


@router.post("/jobs/{job_id}/retry")
@limiter.limit("10/minute")
async def retry_job(request: Request, job_id: str) -> dict:
    """Re-queue only the failed/timed-out sections of a finished job."""
    try:
        j = plagiarism_jobs.job_manager.retry_failed(job_id)
    except RuntimeError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    if j is None:
        raise HTTPException(status_code=404, detail="Job not found or expired.")
    return plagiarism_jobs.serialize_job(j)


@router.delete("/jobs/{job_id}")
@limiter.limit("30/minute")
async def cancel_job(request: Request, job_id: str) -> dict:
    """Signal cancellation. The worker stops at the next section boundary."""
    ok = plagiarism_jobs.job_manager.cancel_job(job_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Job not found.")
    return {"ok": True}


# ---------------------------------------------------------------------------
# /parse-report — Path A intake helper
# ---------------------------------------------------------------------------
# Receives a plagiarism-checker report (PDF / DOCX / TXT) plus the
# software name, extracts text, and returns a {section: similarity_%}
# map the intake page can persist into sessionStorage and pass straight
# back into POST /jobs as the ``report`` field.
# ---------------------------------------------------------------------------


@router.post("/parse-report")
@limiter.limit("20/minute")
async def parse_report(
    request: Request,
    file: UploadFile = File(..., description="Plagiarism report PDF/DOCX/TXT"),
    software: str = Form(default="Other"),
) -> dict:
    """Extract a flagged-sections map from an uploaded plagiarism report."""
    from app.services import report_parser

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The uploaded report is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"Report exceeds the {MAX_UPLOAD_BYTES // (1024 * 1024)} MB upload cap.",
        )

    filename = file.filename or "report"
    suffix = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    # TXT goes straight to UTF-8 decode; PDF / DOCX reuse the same
    # extractor the original-document upload uses, so we get identical
    # error semantics (password-protected PDFs, oversized files, etc.).
    if suffix in ("txt", "text"):
        try:
            text = content.decode("utf-8", errors="replace")
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(
                status_code=400,
                detail=f"Could not read the report as text: {exc}",
            ) from exc
    else:
        try:
            text = _safe_extract_text(filename, content)
        except HTTPException:
            raise
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(
                status_code=400,
                detail=plagiarism_analyzer.sanitize_error_message(exc),
            ) from exc

    flagged_map = report_parser.parse_report_text(text or "")
    summary = report_parser.summarise_report(flagged_map)

    return {
        "filename": filename,
        "software": (software or "Other").strip()[:80] or "Other",
        "extracted_chars": len(text or ""),
        "flagged_map": flagged_map,
        "summary": summary,
        # Helpful diagnostic so the UI can warn the user when the
        # report parser couldn't find anything (e.g. they uploaded the
        # original by mistake).
        "parsed_section_count": len(flagged_map),
    }


# ---------------------------------------------------------------------------
# DOCX export — Times New Roman 12pt academic formatting
# ---------------------------------------------------------------------------

class ExportSectionIn(BaseModel):
    label: str = Field(..., min_length=1, max_length=200)
    text: str = ""
    skipped: bool = False
    skip_reason: Optional[str] = None


class ExportDocxRequest(BaseModel):
    title: str = Field(default="Rewritten document", max_length=200)
    sections: List[ExportSectionIn] = Field(..., min_length=1)
    notes: Optional[str] = None
    filename: Optional[str] = None  # for the download filename only


def _safe_filename(name: str) -> str:
    """Strip path separators and weird chars from a user-provided basename."""
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", name or "").strip("._")
    return base or "rewritten"


@router.post("/export-docx")
@limiter.limit("20/minute")
def export_docx(request: Request, payload: ExportDocxRequest):
    """Render the rewrite result as a DOCX file in academic format.

    Formatting:
      * Times New Roman, 12pt body
      * 1.5 line spacing, justified
      * 1 inch (2.54 cm) margins
      * Section heading per ``label`` (TNR 12pt, bold, before/after spacing)
      * Title page-style heading at the top
      * Page numbers in the footer (right-aligned)
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="DOCX export is unavailable: python-docx is not installed.",
        ) from exc

    doc = Document()

    # 1 inch (2.54 cm) margins on every section.
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Force the default style to Times New Roman 12pt — applies to body and
    # any paragraph that doesn't override it.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # Ensure east-asian / cs fallbacks also map to TNR (Word quirk).
    rpr = style.element.rPr
    if rpr is not None:
        for tag in ("w:eastAsia", "w:cs", "w:hAnsi", "w:ascii"):
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                rfonts.set(qn(tag), "Times New Roman")

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(payload.title or "Rewritten document")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)
    title_run.bold = True

    if payload.notes:
        notes_p = doc.add_paragraph()
        notes_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notes_run = notes_p.add_run(payload.notes)
        notes_run.font.name = "Times New Roman"
        notes_run.font.size = Pt(11)
        notes_run.italic = True

    doc.add_paragraph()  # spacer

    # Body sections
    for sec in payload.sections:
        # Heading — manual paragraph so we keep TNR 12pt (Word's built-in
        # heading styles use a different font and size).
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        h_run = h.add_run(sec.label)
        h_run.font.name = "Times New Roman"
        h_run.font.size = Pt(12)
        h_run.bold = True

        text = (sec.text or "").strip()
        if not text:
            empty_p = doc.add_paragraph()
            empty_run = empty_p.add_run("[empty section]")
            empty_run.font.name = "Times New Roman"
            empty_run.font.size = Pt(12)
            empty_run.italic = True
            continue

        # Split on blank lines into paragraphs; preserve internal newlines
        # within a paragraph as soft line breaks.
        paragraphs = re.split(r"\n\s*\n+", text)
        for para_text in paragraphs:
            para_text = para_text.strip("\n")
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pf.space_after = Pt(6)
            pf.first_line_indent = Cm(1.27)  # ~0.5 inch
            lines = para_text.split("\n")
            for li, line in enumerate(lines):
                if li > 0:
                    p.add_run().add_break()
                run = p.add_run(line)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    # Page numbers in the footer (right-aligned). Uses raw OOXML because
    # python-docx doesn't expose page-number fields directly.
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = fp.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    base = _safe_filename(payload.filename or payload.title or "rewritten")
    if base.lower().endswith(".docx"):
        base = base[:-5]
    download_name = f"{base}_rewritten.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{download_name}"',
            "Cache-Control": "no-store",
        },
    )


# ===========================================================================
# /suggest-citations — RAG-grounded citation suggestions
# ===========================================================================

class SuggestCitationsRequest(BaseModel):
    """Find real, verified published papers that could be cited to support
    the factual claims in ``text``. Used by the rewrite-results page after
    a section has been paraphrased."""
    text: str = Field(..., min_length=20, max_length=20_000,
                      description="The passage (typically one rewritten section) to analyse.")
    topic_hint: Optional[str] = Field(default=None, max_length=500,
                                      description="Optional topic to bias database routing (e.g. 'type 2 diabetes pharmacotherapy').")
    max_claims: int = Field(default=5, ge=1, le=8,
                            description="Maximum number of claims to extract from the passage.")


@router.post("/suggest-citations")
@limiter.limit("10/minute")
async def suggest_citations_route(request: Request, payload: SuggestCitationsRequest) -> Dict[str, Any]:
    """Return RAG-grounded citation suggestions for the given passage.

    Pipeline (see ``app/services/citation_suggester.py``):
    1. Gemini extracts up to ``max_claims`` citation-worthy claims +
       focused search queries (with anti-hallucination quote check).
    2. ``rag_router`` picks the appropriate databases.
    3. ``rag_retriever`` fans out concurrently and returns real
       deduplicated records (PubMed, Europe PMC, Crossref, OpenAlex, …).

    The response carries the original retriever shape per suggestion so
    the front-end can render title, authors, year, journal, and a DOI
    link without any further lookups. ``suggestions`` may be empty for an
    individual claim when no live database returned a match — never a
    placeholder.
    """
    try:
        result = await citation_suggester.suggest_citations(
            text=payload.text,
            topic_hint=payload.topic_hint,
            max_claims=payload.max_claims,
        )
    except Exception as exc:                                  # noqa: BLE001
        log.exception("suggest_citations failed")
        raise HTTPException(status_code=503,
                            detail=f"Citation suggestion service is unavailable: {exc}") from exc
    return result
