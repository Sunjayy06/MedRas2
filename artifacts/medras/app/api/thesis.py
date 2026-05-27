"""Thesis Writing Module API.

Endpoints
---------
* ``GET  /api/thesis/spine``                 — chapter spine + default rules
* ``POST /api/thesis/parse-guidelines``      — multipart upload of uni rules PDF
* ``POST /api/thesis/references/verify-dois``— bulk DOI verification
* ``POST /api/thesis/references/search``     — distilled RAG search
* ``GET  /api/thesis/proforma-template``     — blank standard Indian MD/MS proforma DOCX download
* ``POST /api/thesis/draft-section``         — RAG-grounded fresh section draft
* ``POST /api/thesis/improve-section``       — sentence-level inline-diff suggestions
* ``POST /api/thesis/draft-abstract``        — structured abstract from researcher's own data
* ``POST /api/thesis/extract-style-sample``  — extract body prose from uploaded sample for style matching
* ``POST /api/thesis/compliance-check``      — pre-flight checks on full state
* ``POST /api/thesis/extract-text``          — extract text from uploaded stats / data file
"""
from __future__ import annotations

import asyncio
import logging
from typing import Any, Dict, List

from fastapi import APIRouter, File, HTTPException, Request, Response, UploadFile

from app.core.limiter import limiter
from app.services import (
    thesis_compliance, thesis_export, thesis_formats, thesis_guidelines_parser,
    thesis_reference_library, thesis_section_writer,
)
from app.services.proposal_generator import GeneratorError

log = logging.getLogger(__name__)
router = APIRouter(prefix="/thesis", tags=["thesis"])

MAX_UPLOAD_BYTES = 30 * 1024 * 1024


# ---------------------------------------------------------------------------
# Spine + defaults
# ---------------------------------------------------------------------------

@router.get("/spine")
async def get_spine(
    mode: str = "thesis",
    article_type: str = "",
    design: str = "",
    tier: str = "",
    citation_style: str = "",
) -> Dict[str, Any]:
    """Return the right chapter spine for the writer.

    * ``mode=thesis`` (default) → canonical Indian MD/DNB/PhD spine.
    * ``mode=article`` → the spine for the matching reporting checklist
      (CARE / CONSORT / STROBE / PRISMA / MOOSE / COREQ / IMRaD /
      narrative), with per-section word budgets sized from the journal
      tier (``t1``..``t4``, default ``t3``).
    """
    if (mode or "").strip().lower() == "article":
        spine = thesis_formats.get_article_spine(article_type, design)
        tier_targets = thesis_formats.get_tier_targets(tier)
        spine = thesis_formats.apply_tier_to_spine(spine, tier_targets)
        rules = dict(thesis_formats.DEFAULT_RULES)
        rules["citation_style"] = (
            (citation_style or "").strip().lower()
            or tier_targets.get("default_citation_style", "vancouver")
        )
        rules["min_references"] = tier_targets.get("ref_min") or rules["min_references"]
        if tier_targets.get("ref_max"):
            rules["max_references"] = tier_targets["ref_max"]
        rules["abstract_words"] = tier_targets.get("abstract_words")
        rules["abstract_structured"] = tier_targets.get("abstract_structured")
        rules["max_pages"] = None  # articles aren't page-capped, they're word-capped
        return {
            "spine":     spine,
            "rules":     rules,
            "tier":      tier_targets,
            "checklist": thesis_formats.resolve_checklist(article_type, design),
            "mode":      "article",
            "version":   "v2-article-tier-aware",
        }
    return {
        "spine":   thesis_formats.CHAPTER_SPINE,
        "rules":   thesis_formats.DEFAULT_RULES,
        "mode":    "thesis",
        "version": "v1-indian-md-dnb-phd",
    }


# ---------------------------------------------------------------------------
# Guidelines parser
# ---------------------------------------------------------------------------

@router.post("/parse-guidelines")
@limiter.limit("10/minute")
async def parse_guidelines(request: Request, file: UploadFile = File(...)) -> Dict[str, Any]:
    """Upload a university thesis-guidelines PDF / DOCX / TXT and get back
    the rules autofilled from it (plus any rules that fell back to defaults).
    """
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (>30 MB).")
    try:
        return await asyncio.to_thread(
            thesis_guidelines_parser.parse_guidelines, file.filename or "", data
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))


# ---------------------------------------------------------------------------
# Reference library
# ---------------------------------------------------------------------------

@router.post("/references/verify-dois")
@limiter.limit("20/minute")
async def verify_dois(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: ``{dois: ["10.x/y", ...]}``  →  per-DOI verified record or
    ``{verified: false, error: ...}``."""
    dois: List[str] = payload.get("dois") or []
    if not isinstance(dois, list):
        raise HTTPException(status_code=400, detail="`dois` must be a list.")
    dois = [str(d).strip() for d in dois if str(d).strip()][:60]
    out = await thesis_reference_library.verify_dois(dois)
    # Attach a one-line distilled summary for verified records
    for r in out:
        if r.get("verified"):
            r["summary"] = thesis_reference_library.summarise(r)
    return {"records": out}


@router.post("/references/extract-dois")
@limiter.limit("20/minute")
async def extract_dois(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: ``{text: "..."}``  →  ``{dois: [...]}``. Used after the user
    pastes text or uploads a PDF whose text the client has extracted."""
    text = (payload.get("text") or "")
    if not isinstance(text, str):
        raise HTTPException(status_code=400, detail="`text` must be a string.")
    return {"dois": thesis_reference_library.extract_dois(text[:200_000])}


@router.post("/references/search")
@limiter.limit("20/minute")
async def references_search(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: ``{topic, domain_hint?, limit?}``."""
    topic = (payload.get("topic") or "").strip()
    if not topic:
        raise HTTPException(status_code=400, detail="`topic` is required.")
    domain_hint = payload.get("domain_hint")
    limit = max(5, min(int(payload.get("limit") or 20), 40))
    res = await thesis_reference_library.search(topic, domain_hint=domain_hint, limit=limit)
    # Attach distilled summaries
    for r in res.get("records", []):
        r["summary"] = thesis_reference_library.summarise(r)
    return res


@router.post("/references/import-list")
@limiter.limit("10/minute")
async def import_reference_list(
    request: Request, file: UploadFile = File(...)
) -> Dict[str, Any]:
    """Upload a DOCX / PDF / TXT file containing a reference list.

    Parses numbered or bulleted Vancouver-style references, DOI-only lines,
    or mixed text. Entries with resolvable DOIs are Crossref-verified;
    others are heuristically parsed and marked ``verified: false``.

    Returns ``{entries: [...], parsed: N, verified: N}``.
    """
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (>30 MB).")

    fname = (file.filename or "").lower()

    def _extract_text() -> str:
        if fname.endswith(".docx"):
            try:
                from docx import Document  # type: ignore
                import io
                doc = Document(io.BytesIO(data))
                return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            except Exception as exc:
                raise ValueError(f"Could not read DOCX: {exc}") from exc
        elif fname.endswith(".pdf"):
            try:
                from pypdf import PdfReader  # type: ignore
                import io
                reader = PdfReader(io.BytesIO(data))
                pages = [(page.extract_text() or "").strip() for page in reader.pages]
                return "\n\n".join(p for p in pages if p)
            except Exception as exc:
                raise ValueError(f"Could not read PDF: {exc}") from exc
        else:
            return data.decode("utf-8", errors="replace")

    try:
        text = await asyncio.to_thread(_extract_text)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract any text from the file.",
        )

    try:
        entries = await thesis_reference_library.parse_reference_list(text[:200_000])
    except Exception as exc:  # noqa: BLE001
        log.exception("import_reference_list: parse failed")
        raise HTTPException(status_code=500, detail=f"Parsing failed: {exc}")

    # Attach distilled summaries to verified entries
    for e in entries:
        if e.get("verified"):
            e["summary"] = thesis_reference_library.summarise(e)

    verified_count = sum(1 for e in entries if e.get("verified"))
    return {
        "entries": entries,
        "parsed": len(entries),
        "verified": verified_count,
    }


# ---------------------------------------------------------------------------
# Section writer
# ---------------------------------------------------------------------------

@router.post("/draft-section")
@limiter.limit("8/minute")
async def draft_section(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: ``{chapter_id, topic, citation_style?, locked_numbers?,
    extra_context?, domain_hint?, mode?, style_choice?, style_sample?,
    ref_library?}``."""
    ref_lib_raw = payload.get("ref_library")
    # Never ground the AI on retracted papers; silently exclude them.
    if isinstance(ref_lib_raw, list):
        ref_library = [r for r in ref_lib_raw if not r.get("retracted")][:200]
    else:
        ref_library = None
    try:
        word_limit_raw = payload.get("word_limit")
        word_limit = int(word_limit_raw) if word_limit_raw and str(word_limit_raw).isdigit() else None
        subsection_hint = payload.get("subsection_hint")
        if not isinstance(subsection_hint, dict):
            subsection_hint = None
        result = await thesis_section_writer.draft_section(
            chapter_id=payload.get("chapter_id") or "",
            topic=payload.get("topic") or "",
            citation_style=(payload.get("citation_style") or "vancouver"),
            locked_numbers=payload.get("locked_numbers") or {},
            extra_context=payload.get("extra_context"),
            domain_hint=payload.get("domain_hint"),
            mode=payload.get("mode") or "thesis",
            style_choice=payload.get("style_choice") or "indian_formal",
            style_sample=payload.get("style_sample"),
            ref_library=ref_library,
            word_limit=word_limit,
            subsection_hint=subsection_hint,
            rol_writing_format=payload.get("rol_writing_format"),
        )
    except GeneratorError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return result


@router.post("/plan-subsections")
@limiter.limit("10/minute")
async def plan_subsections_endpoint(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: ``{chapter_id, topic, aim?, objectives?, study_type?, extra_context?, mode?}``."""
    try:
        result = await thesis_section_writer.plan_subsections(
            chapter_id=payload.get("chapter_id") or "",
            topic=payload.get("topic") or "",
            aim=payload.get("aim"),
            objectives=payload.get("objectives"),
            study_type=payload.get("study_type"),
            extra_context=payload.get("extra_context"),
            mode=payload.get("mode") or "thesis",
        )
    except GeneratorError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return result


@router.post("/improve-section")
@limiter.limit("12/minute")
async def improve_section(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: ``{chapter_id, topic, current_text, citation_style?,
    locked_numbers?, domain_hint?, mode?, style_choice?, style_sample?,
    ref_library?}``."""
    ref_lib_raw = payload.get("ref_library")
    if isinstance(ref_lib_raw, list):
        ref_library = [r for r in ref_lib_raw if not r.get("retracted")][:200]
    else:
        ref_library = None
    try:
        result = await thesis_section_writer.improve_section(
            chapter_id=payload.get("chapter_id") or "",
            topic=payload.get("topic") or "",
            current_text=payload.get("current_text") or "",
            citation_style=(payload.get("citation_style") or "vancouver"),
            locked_numbers=payload.get("locked_numbers") or {},
            domain_hint=payload.get("domain_hint"),
            mode=payload.get("mode") or "thesis",
            style_choice=payload.get("style_choice") or "indian_formal",
            style_sample=payload.get("style_sample"),
            ref_library=ref_library,
            polish_instruction=payload.get("polish_instruction"),
        )
    except GeneratorError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return result


@router.post("/draft-abstract")
@limiter.limit("8/minute")
async def draft_abstract_endpoint(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Draft a structured abstract from the researcher's own chapter content.

    Body: ``{topic, extra_context?, locked_numbers?, word_limit?,
    mode?, style_choice?, style_sample?}``.
    """
    try:
        result = await thesis_section_writer.draft_abstract(
            topic=payload.get("topic") or "",
            extra_context=payload.get("extra_context"),
            locked_numbers=payload.get("locked_numbers") or {},
            word_limit=int(payload.get("word_limit") or 280),
            mode=payload.get("mode") or "thesis",
            style_choice=payload.get("style_choice") or "indian_formal",
            style_sample=payload.get("style_sample"),
            domain_hint=payload.get("domain_hint"),
        )
    except GeneratorError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return result


@router.post("/import-section")
@limiter.limit("20/minute")
async def import_section(request: Request, file: UploadFile = File(...)) -> Dict[str, Any]:
    """Extract prose and tables from a researcher-uploaded DOCX / PDF / TXT file.

    Returns ``{prose: str, tables: [{rows: [[str]], caption: str}]}``.

    * **DOCX** — paragraphs become prose; Word tables are extracted as structured
      row/cell arrays so the client can render keep/discard review UI.
    * **PDF** — text extracted per-page via pypdf (no table structure — tables=[]).
    * **TXT** — raw text treated as prose.
    """
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (>30 MB).")

    fname = (file.filename or "").lower()

    def _extract() -> Dict[str, Any]:
        import re
        prose_parts: List[str] = []
        tables: List[Dict[str, Any]] = []

        if fname.endswith(".docx"):
            try:
                from docx import Document  # type: ignore
                from docx.oxml.ns import qn  # type: ignore
                import io
                doc = Document(io.BytesIO(data))
                # Walk body children in document order so paragraphs and tables
                # are interleaved exactly as authored.
                for child in doc.element.body:
                    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                    if tag == "p":
                        para_text = "".join(
                            t.text or "" for t in child.iter(qn("w:t"))
                        ).strip()
                        if para_text:
                            prose_parts.append(para_text)
                    elif tag == "tbl":
                        rows: List[List[str]] = []
                        for tr in child.iter(qn("w:tr")):
                            cells: List[str] = []
                            for tc in tr.iter(qn("w:tc")):
                                # Skip vertically-merged continuation cells
                                # (vMerge with no "val" attr = continuation; has val = start)
                                v_merge = tc.find(
                                    ".//" + qn("w:vMerge")
                                )
                                if v_merge is not None and v_merge.get(
                                    qn("w:val")
                                ) is None:
                                    # continuation cell — insert empty placeholder
                                    cells.append("")
                                    continue
                                cell_text = "".join(
                                    t.text or "" for t in tc.iter(qn("w:t"))
                                ).strip()
                                cells.append(cell_text)
                            if cells:
                                rows.append(cells)
                        if rows:
                            tables.append({"rows": rows, "caption": ""})
            except Exception as exc:
                raise ValueError(f"Could not read DOCX: {exc}") from exc

        elif fname.endswith(".pdf"):
            try:
                from pypdf import PdfReader  # type: ignore
                import io
                reader = PdfReader(io.BytesIO(data))
                for page in reader.pages:
                    text = (page.extract_text() or "").strip()
                    if text:
                        prose_parts.append(text)
                # PDF: no structural table extraction — tables stays []
            except Exception as exc:
                raise ValueError(f"Could not read PDF: {exc}") from exc

        else:
            # Plain text or unknown
            prose_parts.append(data.decode("utf-8", errors="replace"))

        prose = "\n\n".join(prose_parts)

        # Strip back matter (References / Bibliography section)
        ref_match = re.search(
            r"\b(REFERENCES|BIBLIOGRAPHY|WORKS CITED)\b",
            prose, flags=re.IGNORECASE,
        )
        if ref_match:
            prose = prose[: ref_match.start()]

        return {
            "prose": prose.strip()[:12_000],
            "tables": tables,
        }

    try:
        result = await asyncio.to_thread(_extract)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    if not result["prose"].strip() and not result["tables"]:
        raise HTTPException(
            status_code=400,
            detail=(
                "Could not extract content from the file. "
                "Try a DOCX or TXT version of the document."
            ),
        )

    return result


@router.get("/proforma-template")
async def proforma_template() -> Response:
    """Return a blank standard Indian MD / MS proforma as a downloadable DOCX.

    The template covers all 14 conventional sections (patient identification,
    chief complaints, history, examination, investigations, diagnosis, treatment,
    outcome, follow-up) in the format used across NBEMS MD/MS programmes.
    """
    from app.services.proforma_template import build_proforma_docx

    buf = build_proforma_docx()
    return Response(
        content=buf.getvalue(),
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": 'attachment; filename="proforma_template.docx"'
        },
    )


@router.post("/earlier-studies")
@limiter.limit("5/minute")
async def earlier_studies_endpoint(
    request: Request, payload: Dict[str, Any]
) -> Dict[str, Any]:
    """Generate the Earlier Studies section for Review of Literature.

    Retrieves up to 20 records, selects up to 15 most relevant, and returns:
    ``{text, table_html, paragraphs, sources}``.
    """
    ref_lib_raw = payload.get("ref_library")
    ref_library = None
    if isinstance(ref_lib_raw, list):
        ref_library = [r for r in ref_lib_raw if not r.get("retracted")][:200]
    try:
        result = await thesis_section_writer.generate_earlier_studies(
            topic=payload.get("topic") or "",
            extra_context=payload.get("extra_context"),
            ref_library=ref_library,
            domain_hint=payload.get("domain_hint"),
        )
    except GeneratorError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return result


@router.post("/extract-style-sample")
@limiter.limit("10/minute")
async def extract_style_sample(
    request: Request, file: UploadFile = File(...)
) -> Dict[str, Any]:
    """Upload a DOCX / PDF / TXT sample of the researcher's own writing.

    Returns ``{text}`` — up to ~2000 words of body prose with front matter
    (title, authors) and back matter (references) stripped. The client stores
    this in sessionStorage and passes it to draft-section as ``style_sample``
    when ``style_choice == "uploaded"``.
    """
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (>30 MB).")

    fname = (file.filename or "").lower()

    def _extract() -> str:
        raw = ""
        if fname.endswith(".docx"):
            try:
                from docx import Document  # type: ignore
                import io
                doc = Document(io.BytesIO(data))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                raw = "\n\n".join(paragraphs)
            except Exception as exc:
                raise ValueError(f"Could not read DOCX: {exc}") from exc
        elif fname.endswith(".pdf"):
            try:
                from pypdf import PdfReader  # type: ignore
                import io
                reader = PdfReader(io.BytesIO(data))
                pages = [
                    (page.extract_text() or "").strip()
                    for page in reader.pages
                    if (page.extract_text() or "").strip()
                ]
                raw = "\n\n".join(pages)
            except Exception as exc:
                raise ValueError(f"Could not read PDF: {exc}") from exc
        else:
            raw = data.decode("utf-8", errors="replace")

        # Strip front matter (first 300 chars likely title/authors/affiliations)
        if len(raw) > 300:
            raw = raw[300:]

        # Strip back matter — everything after REFERENCES / BIBLIOGRAPHY heading
        import re
        ref_match = re.search(
            r"\b(REFERENCES|BIBLIOGRAPHY|WORKS CITED)\b",
            raw, flags=re.IGNORECASE,
        )
        if ref_match:
            raw = raw[:ref_match.start()]

        # Return up to 6000 characters (~2000 words) of body prose
        return raw.strip()[:6000]

    try:
        text = await asyncio.to_thread(_extract)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    if len(text.strip()) < 100:
        raise HTTPException(
            status_code=400,
            detail="Could not extract enough prose from the file. "
                   "Try a DOCX or TXT version of the document.",
        )

    return {"text": text}


# ---------------------------------------------------------------------------
# Compliance
# ---------------------------------------------------------------------------

@router.post("/compliance-check")
async def compliance_check(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Body: full thesis state JSON. Returns ``{items, summary}``."""
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    return thesis_compliance.check(payload)


# ---------------------------------------------------------------------------
# Export
# ---------------------------------------------------------------------------

def _export_filename(payload: Dict[str, Any], ext: str) -> str:
    state = payload.get("state") or {}
    tm = payload.get("title_meta") or state.get("title_meta") or {}
    setup = state.get("setup") or {}
    name = (tm.get("study_title") or setup.get("topic") or "thesis").strip()
    import re as _re
    safe = _re.sub(r"[^A-Za-z0-9._-]+", "_", name)[:60] or "thesis"
    return f"{safe}.{ext}"


@router.post("/export/docx")
@limiter.limit("12/minute")
async def export_docx(request: Request, payload: Dict[str, Any]) -> Response:
    """Body: ``{state, title_meta?, consent?, assets?}``. Returns DOCX bytes."""
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    try:
        data = await asyncio.to_thread(thesis_export.build_docx, payload)
    except Exception as exc:                                      # noqa: BLE001
        log.exception("thesis docx export failed")
        raise HTTPException(status_code=500, detail=f"Word export failed: {exc}")
    fname = _export_filename(payload, "docx")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/export/pdf")
@limiter.limit("12/minute")
async def export_pdf(request: Request, payload: Dict[str, Any]) -> Response:
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    try:
        data = await asyncio.to_thread(thesis_export.build_pdf, payload)
    except Exception as exc:                                      # noqa: BLE001
        log.exception("thesis pdf export failed")
        raise HTTPException(status_code=500, detail=f"PDF export failed: {exc}")
    fname = _export_filename(payload, "pdf")
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/export/zip")
@limiter.limit("8/minute")
async def export_zip(request: Request, payload: Dict[str, Any]) -> Response:
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    try:
        data = await asyncio.to_thread(thesis_export.build_zip, payload)
    except Exception as exc:                                      # noqa: BLE001
        log.exception("thesis zip export failed")
        raise HTTPException(status_code=500, detail=f"Bundle export failed: {exc}")
    fname = _export_filename(payload, "zip")
    return Response(
        content=data,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/export/plaintext")
@limiter.limit("20/minute")
async def export_plaintext(request: Request, payload: Dict[str, Any]) -> Dict[str, str]:
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    try:
        return {"text": thesis_export.build_plaintext(payload)}
    except Exception as exc:                                      # noqa: BLE001
        log.exception("thesis plaintext export failed")
        raise HTTPException(status_code=500, detail=f"Plaintext export failed: {exc}")


# ---------------------------------------------------------------------------
# Thesis → Article conversion
# ---------------------------------------------------------------------------

@router.post("/convert-to-article")
@limiter.limit("4/minute")
async def convert_to_article(request: Request, payload: Dict[str, Any]) -> Dict[str, Any]:
    """Condense thesis chapters to IMRaD journal-article format.

    Body::

        {
          "topic":          str,
          "journal_family": "plos"|"bmc"|"bmj"|"frontiers"|"tier1"|"regional",
          "article_type":   "original_research"|"brief_report"|"short_communication",
          "chapters": {
              "introduction":    {"text": str},
              "literature_review": {"text": str},
              "aims":            {"text": str},
              "methods":         {"text": str},
              "results":         {"text": str},
              "discussion":      {"text": str},
              "conclusion":      {"text": str},
          },
          "locked_numbers": {label: value, …},
        }

    Returns the same shape as ``thesis_section_writer.condense_for_article``
    plus ``journal_family``, ``article_type``, and ``topic``.
    """
    topic = (payload.get("topic") or "").strip()
    if not topic:
        raise HTTPException(status_code=400, detail="`topic` is required.")

    journal_family = (payload.get("journal_family") or "plos").strip().lower()
    article_type   = (payload.get("article_type")   or "original_research").strip().lower()
    chapters       = payload.get("chapters") or {}
    locked_numbers = payload.get("locked_numbers") or {}
    ref_lib_raw    = payload.get("ref_library")
    ref_library    = ref_lib_raw[:200] if isinstance(ref_lib_raw, list) else None

    def _ch(cid: str) -> str:
        val = chapters.get(cid) or {}
        if isinstance(val, dict):
            return (val.get("text") or "").strip()
        return str(val).strip()

    intro_blob = "\n\n".join(filter(None, [
        _ch("introduction"), _ch("literature_review"), _ch("aims"),
    ]))
    methods_blob = _ch("methods")
    results_blob = _ch("results")
    disc_blob    = "\n\n".join(filter(None, [_ch("discussion"), _ch("conclusion")]))

    missing = []
    if not intro_blob:   missing.append("Introduction")
    if not methods_blob: missing.append("Methods")
    if not results_blob: missing.append("Results")
    if not disc_blob:    missing.append("Discussion")

    if missing:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Missing content in required chapters: {', '.join(missing)}. "
                "Write these chapters in the Editor first, then convert to article."
            ),
        )

    try:
        result = await thesis_section_writer.condense_for_article(
            topic=topic,
            journal_family=journal_family,
            article_type=article_type,
            introduction_text=intro_blob,
            methods_text=methods_blob,
            results_text=results_blob,
            discussion_text=disc_blob,
            locked_numbers=locked_numbers,
            ref_library=ref_library,
        )
    except GeneratorError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    return {
        **result,
        "journal_family": journal_family,
        "article_type":   article_type,
        "topic":          topic,
    }


@router.post("/export/article-docx")
@limiter.limit("12/minute")
async def export_article_docx(request: Request, payload: Dict[str, Any]) -> Response:
    """Export a condensed article as DOCX.

    Body: ``{title, authors?, sections: {abstract, introduction, methods,
    results, discussion}, references?, metadata: {credit?,
    data_availability?, competing_interests?}}``.
    """
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    try:
        data = await asyncio.to_thread(thesis_export.build_article_docx, payload)
    except Exception as exc:                                      # noqa: BLE001
        log.exception("article docx export failed")
        raise HTTPException(status_code=500, detail=f"Word export failed: {exc}")
    import re as _re
    name = _re.sub(r"[^A-Za-z0-9._-]+", "_",
                   (payload.get("title") or "article").strip())[:60] or "article"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{name}.docx"'},
    )


@router.post("/export/article-pdf")
@limiter.limit("12/minute")
async def export_article_pdf(request: Request, payload: Dict[str, Any]) -> Response:
    """Export a condensed article as PDF."""
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="Payload must be a JSON object.")
    try:
        data = await asyncio.to_thread(thesis_export.build_article_pdf, payload)
    except Exception as exc:                                      # noqa: BLE001
        log.exception("article pdf export failed")
        raise HTTPException(status_code=500, detail=f"PDF export failed: {exc}")
    import re as _re
    name = _re.sub(r"[^A-Za-z0-9._-]+", "_",
                   (payload.get("title") or "article").strip())[:60] or "article"
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{name}.pdf"'},
    )
