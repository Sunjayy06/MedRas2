"""Folio — Document Management Module API.

POST /api/folio/import-docx         — Extract structured content from DOCX upload.
POST /api/folio/format-references   — Reformat reference list in a citation style.
POST /api/folio/parse-feedback      — Parse guide feedback into structured operations.
POST /api/folio/export-docx         — Reconstruct DOCX from Folio document model.
"""
from __future__ import annotations

import asyncio
import io
import json
import os
import re
import uuid
from typing import Any, Dict, List

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import Response

from app.core.logging import get_logger

log = get_logger(__name__)
router = APIRouter(prefix="/folio", tags=["folio"])

# ── Helpers ───────────────────────────────────────────────────────────────────

def _uid() -> str:
    return str(uuid.uuid4())[:8]

_REF_LINE_RE   = re.compile(r"^\s*(\d+)[.\)]\s+.{10,}")
_TABLE_CAP_RE  = re.compile(r"^(table)\s+\d+", re.I)
_FIGURE_CAP_RE = re.compile(r"^(figure|fig\.?)\s+\d+", re.I)
_REF_SECTION_RE = re.compile(r"\b(references|bibliography|reference list|works cited)\b", re.I)


def _extract_docx(content: bytes) -> Dict[str, Any]:
    doc = DocxDocument(io.BytesIO(content))

    sections: List[Dict]  = []
    references: List[Dict] = []
    tables_out: List[Dict] = []
    figures_out: List[Dict] = []

    current_section: Dict | None = None
    in_ref_section  = False
    ref_num         = 0
    tbl_num         = 0
    fig_num         = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = (para.style.name or "").lower()

        # ── Headings ──────────────────────────────────────────────────────────
        if style.startswith("heading"):
            try:
                level = int(re.search(r"\d+", style).group())
            except Exception:
                level = 1
            in_ref_section = bool(_REF_SECTION_RE.search(text))
            current_section = {
                "id": _uid(), "level": level,
                "title": text, "paragraphs": [],
            }
            sections.append(current_section)
            continue

        # ── Reference list items ──────────────────────────────────────────────
        if in_ref_section and _REF_LINE_RE.match(text):
            ref_num += 1
            references.append({
                "id": _uid(), "number": ref_num,
                "text": text, "raw": text,
            })
            continue

        # ── Figure captions ───────────────────────────────────────────────────
        if _FIGURE_CAP_RE.match(text):
            fig_num += 1
            figures_out.append({"id": _uid(), "number": fig_num, "caption": text})
            continue

        # ── Body paragraph ────────────────────────────────────────────────────
        pid = _uid()
        if current_section is None:
            current_section = {
                "id": _uid(), "level": 0,
                "title": "Preamble", "paragraphs": [],
            }
            sections.insert(0, current_section)
        current_section["paragraphs"].append({"id": pid, "text": text})

    # ── Tables ────────────────────────────────────────────────────────────────
    for tbl in doc.tables:
        tbl_num += 1
        rows = []
        caption = f"Table {tbl_num}"
        for i, row in enumerate(tbl.rows[:8]):
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
            if i == 0:
                caption = " | ".join(c for c in cells if c)[:80] or caption
        tables_out.append({
            "id": _uid(), "number": tbl_num,
            "caption": caption, "rows": rows,
        })

    title = (doc.core_properties.title or "").strip()
    if not title and sections:
        title = sections[0].get("title", "Untitled Document")

    return {
        "title":      title or "Untitled Document",
        "sections":   sections,
        "references": references,
        "tables":     tables_out,
        "figures":    figures_out,
    }


# ── Citation formatters ───────────────────────────────────────────────────────

def _fmt_vancouver(ref: Dict, n: int) -> str:
    raw = ref.get("raw", ref.get("text", ""))
    # Strip any leading number from the raw text and re-prefix
    clean = re.sub(r"^\d+[.\)]\s*", "", raw).strip()
    return f"{n}. {clean}"


def _fmt_apa(ref: Dict, _n: int) -> str:
    raw = ref.get("raw", ref.get("text", ""))
    return re.sub(r"^\d+[.\)]\s*", "", raw).strip()


def _fmt_harvard(ref: Dict, _n: int) -> str:
    return _fmt_apa(ref, _n)


_STYLE_FMT = {
    "vancouver": _fmt_vancouver,
    "apa":       _fmt_apa,
    "harvard":   _fmt_harvard,
}


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/import-docx")
async def import_docx(file: UploadFile = File(...)) -> Dict[str, Any]:
    fn = file.filename or ""
    if not fn.lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files are supported.")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(413, "File exceeds 20 MB limit.")
    if len(content) < 100:
        raise HTTPException(400, "File appears empty or corrupt.")
    try:
        model = _extract_docx(content)
    except Exception as exc:
        log.exception("DOCX extraction failed")
        raise HTTPException(422, f"Could not read document: {exc}") from exc
    return {"ok": True, "document": model}


@router.post("/format-references")
async def format_references(body: Dict[str, Any]) -> Dict[str, Any]:
    style = str(body.get("style", "vancouver")).lower()
    if style not in _STYLE_FMT:
        raise HTTPException(400, f"Unknown style '{style}'. Use: vancouver, apa, harvard.")
    fmt    = _STYLE_FMT[style]
    refs   = body.get("references", [])
    result = [{"id": r.get("id", ""), "text": fmt(r, i + 1)} for i, r in enumerate(refs)]
    return {"ok": True, "formatted": result}


@router.post("/parse-feedback")
async def parse_feedback(body: Dict[str, Any]) -> Dict[str, Any]:
    feedback_text = str(body.get("feedback", "")).strip()
    if not feedback_text:
        raise HTTPException(400, "feedback field is required.")

    doc      = body.get("document", {})
    sections = doc.get("sections", [])
    titles   = [s.get("title", "") for s in sections]
    ref_count = len(doc.get("references", []))

    prompt = f"""You are parsing a research supervisor's written feedback on a thesis or academic document.

Document has {ref_count} references and sections: {json.dumps(titles)}

Supervisor's feedback:
\"\"\"
{feedback_text}
\"\"\"

Return ONLY valid JSON — no markdown, no explanation — with this exact shape:
{{
  "operations": [
    {{
      "type": "delete_refs",
      "description": "Delete references 45 to 75",
      "params": {{"from": 45, "to": 75}}
    }}
  ]
}}

Allowed operation types and their params:
- delete_refs:    {{ "from": int, "to": int }}
- move_section:   {{ "section": str, "position": "before"|"after", "target": str }}
- edit_text:      {{ "instruction": str }}
- change_chart:   {{ "from": str, "to": str }}
- add_section:    {{ "title": str, "after": str }}
- delete_section: {{ "title": str }}

Extract every distinct instruction from the feedback. Return an empty operations list if nothing actionable is found."""

    from app.services.llm_client import get_gemini_client, gemini_is_configured
    if not gemini_is_configured():
        raise HTTPException(503, "AI service not configured.")

    def _call_gemini() -> dict:
        from google.genai import types as gtypes
        client = get_gemini_client()
        resp = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=gtypes.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=1024,
            ),
        )
        return json.loads(resp.text or "{}")

    try:
        result = await asyncio.to_thread(_call_gemini)
        return {"ok": True, **result}
    except json.JSONDecodeError as exc:
        raise HTTPException(502, "AI returned malformed JSON.") from exc
    except Exception as exc:
        log.exception("parse-feedback failed")
        raise HTTPException(502, f"AI service error: {exc}") from exc


@router.post("/export-docx")
async def export_docx_endpoint(body: Dict[str, Any]) -> Response:
    model     = body.get("document", {})
    title     = model.get("title", "Untitled Document")
    sections  = model.get("sections", [])
    refs      = model.get("references", [])
    style     = str(body.get("style", "vancouver")).lower()
    fmt       = _STYLE_FMT.get(style, _fmt_vancouver)

    try:
        doc = DocxDocument()
        sec = doc.sections[0]
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(sec, attr, Inches(1))

        def _para(text: str, bold: bool = False, size: int = 12, indent: bool = False):
            p   = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold      = bold
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing = Pt(18)  # 1.5 × 12
            if indent:
                p.paragraph_format.first_line_indent = Pt(36)
            return p

        # Cover title
        _para(title, bold=True, size=14)

        # Sections
        for sec_data in sections:
            level = sec_data.get("level", 1)
            _para(sec_data.get("title", ""), bold=True, size=12 if level > 1 else 13)
            for para in sec_data.get("paragraphs", []):
                _para(para.get("text", ""), indent=True)

        # References
        if refs:
            _para("References", bold=True, size=13)
            for i, r in enumerate(refs, 1):
                _para(fmt(r, i))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe = re.sub(r"[^\w\s-]", "", title)[:40].strip().replace(" ", "_") or "document"
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe}_folio.docx"'},
        )
    except Exception as exc:
        log.exception("export-docx failed")
        raise HTTPException(500, f"Export failed: {exc}") from exc
