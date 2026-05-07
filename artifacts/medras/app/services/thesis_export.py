"""Thesis export — DOCX / PDF / ZIP / plaintext renderers for the
Thesis Writing Module.

Reuses every paragraph / heading / footer / consent helper from
``proposal_export`` so the visual style (Times New Roman 12 / 14 bold,
1.5 line spacing, 1-inch margins, page-numbered footer, IEC title page,
multi-language consent forms) stays identical across modules.

The thesis-specific bits handled here are:

1. Walking the canonical 16-chapter NBEMS-style spine in order
   (front → body → back), and treating three spine entries as built-in
   structural pages rather than free-text bodies:

       * ``title_page``  → the existing _docx/_pdf_title_page helper
       * ``consent``     → the existing _consent_packs renderer
       * ``references``  → the existing _bibliography_entries renderer

2. Embedding researcher-supplied figures (pictures, certificates,
   appendix images) under the chapter they belong to. Each asset is
   validated for MIME, byte-cap, and total-cap; bad assets are skipped
   so a single malformed upload can never crash an export.

Server-stateless: the entire payload is sent from the client on each
export call.
"""
from __future__ import annotations

import base64
import binascii
import io
import logging
import re
import zipfile
from typing import Any, Dict, List, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from docx import Document

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage, PageBreak, Paragraph, SimpleDocTemplate, Spacer,
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.styles import ParagraphStyle

from app.services import proposal_export as PE
from app.services.thesis_formats import CHAPTER_SPINE

log = logging.getLogger(__name__)

# Hard caps for embedded figures — protect server memory.
_MAX_ASSET_BYTES   = 5 * 1024 * 1024     # 5 MB per image
_MAX_TOTAL_ASSETS  = 30 * 1024 * 1024    # 30 MB across all images
_MAX_ASSET_COUNT   = 60
_ALLOWED_MIME      = {"image/png", "image/jpeg", "image/jpg", "image/webp"}

# Spine entries that are rendered structurally (not from chapter free text).
_STRUCTURAL_CHAPTERS = {"title_page", "consent", "references"}


# ---------------------------------------------------------------------------
# Payload adapters
# ---------------------------------------------------------------------------

def _intake_from_state(state: Dict[str, Any]) -> Dict[str, Any]:
    """Build a proposal-style ``intake`` dict from the thesis state so the
    reused proposal_export helpers (consent template, title page) can
    pull topic / role / format from the same shape they expect."""
    setup = state.get("setup") or {}
    title_meta = state.get("title_meta") or {}
    fmt_id = (setup.get("format_id") or "phd-syn").strip().lower()
    return {
        "topic":  (title_meta.get("study_title") or setup.get("topic") or "").strip(),
        "role":   (setup.get("degree") or "Thesis").strip(),
        "format": {"id": fmt_id, "label": setup.get("format_label") or fmt_id},
    }


def _normalised_assets(assets: Dict[str, Any]) -> Dict[str, List[Dict[str, Any]]]:
    """Validate, decode and group assets by chapter_id.

    Returns ``{chapter_id: [{caption, mime, bytes}, ...]}``. Silently
    skips invalid items (any exception during a single asset's decode) so
    a single malformed upload can never abort the export.
    """
    if not isinstance(assets, dict):
        return {}
    # Pre-decode base64 string-length cap. Base64 inflates by ~4/3, so a
    # 5 MB binary cap means a ~6.8 MB base64 string. We round up to 7 MB
    # to allow padding / data: prefix overhead.
    max_b64_chars = int(_MAX_ASSET_BYTES * 4 / 3) + 1024
    out: Dict[str, List[Dict[str, Any]]] = {}
    total = 0
    count = 0
    for bucket in ("pictures", "certificates", "annexures"):
        items = assets.get(bucket) or []
        if not isinstance(items, list):
            continue
        for item in items:
            if count >= _MAX_ASSET_COUNT or total >= _MAX_TOTAL_ASSETS:
                log.warning("thesis_export: asset cap reached, dropping remaining")
                return out
            if not isinstance(item, dict):
                continue
            try:
                mime = str(item.get("mime") or "").strip().lower()
                if mime not in _ALLOWED_MIME:
                    continue
                chapter_id = str(item.get("chapter_id") or "annexures").strip() or "annexures"
                b64_raw = item.get("b64") or item.get("data") or ""
                if not isinstance(b64_raw, (str, bytes)):
                    continue
                if isinstance(b64_raw, bytes):
                    b64_raw = b64_raw.decode("ascii", "ignore")
                if b64_raw.startswith("data:"):
                    _, _, b64_raw = b64_raw.partition(",")
                # Pre-decode size guard — refuse anything obviously too big
                # before paying the decode cost.
                if not b64_raw or len(b64_raw) > max_b64_chars:
                    continue
                raw = base64.b64decode(b64_raw, validate=False)
            except (binascii.Error, ValueError, TypeError, AttributeError) as exc:
                log.info("thesis_export: skip malformed asset: %s", exc)
                continue
            except Exception as exc:                                # noqa: BLE001
                log.warning("thesis_export: unexpected asset error: %s", exc)
                continue
            if not raw or len(raw) > _MAX_ASSET_BYTES:
                continue
            total += len(raw)
            count += 1
            out.setdefault(chapter_id, []).append({
                "caption": str(item.get("caption") or "").strip(),
                "mime":    mime,
                "bytes":   raw,
            })
    return out


def _safe_filename(name: str, ext: str) -> str:
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", (name or "thesis").strip())[:60]
    return (base or "thesis") + "." + ext


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _docx_render_chapter_assets(doc, items: List[Dict[str, Any]]) -> None:
    """Embed chapter assets with a small italic caption underneath."""
    for idx, item in enumerate(items, 1):
        try:
            buf = io.BytesIO(item["bytes"])
            doc.add_picture(buf, width=Inches(5.5))
            # Centre the picture paragraph.
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:                                  # noqa: BLE001
            log.warning("thesis_export: docx picture embed failed: %s", exc)
            continue
        caption = item.get("caption") or f"Figure {idx}"
        cap_par = doc.add_paragraph()
        cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_par.paragraph_format.space_after = Pt(8)
        run = cap_par.add_run(caption)
        run.italic = True
        PE._set_run_font(run, size=11)


def build_docx(payload: Dict[str, Any]) -> bytes:
    """Render the full thesis (16-chapter spine) to DOCX bytes."""
    state      = payload.get("state")      or {}
    title_meta = payload.get("title_meta") or state.get("title_meta") or {}
    chapters   = state.get("chapters")     or {}
    references = state.get("references")   or []
    consent    = payload.get("consent")    or state.get("consent") or {}
    assets     = _normalised_assets(payload.get("assets") or state.get("assets") or {})

    languages = consent.get("languages") or []
    delivery  = (consent.get("delivery") or "attached").strip().lower()
    intake    = _intake_from_state(state)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)
        PE._add_page_number_footer(sec)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # === Title page ===
    PE._docx_title_page(doc, intake, title_meta)
    doc.add_page_break()

    # === Table of contents ===
    PE._add_heading(doc, "Table of Contents", level=1)
    PE._add_toc_field(doc)
    doc.add_page_break()

    # === Walk the spine in order ===
    for ch in CHAPTER_SPINE:
        cid = ch["id"]
        label = ch["label"]

        if cid == "title_page":
            continue  # already rendered above

        if cid == "consent":
            if delivery == "separate":
                continue
            for lang_label, text in PE._consent_packs(intake, title_meta, languages):
                doc.add_page_break()
                PE._add_heading(doc, f"Informed Consent Form — {lang_label}", level=1)
                PE._add_para(doc, text, size=12,
                             align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
            continue

        if cid == "references":
            entries = PE._bibliography_entries(references)
            if not entries:
                continue
            doc.add_page_break()
            PE._add_heading(doc, label, level=1)
            for entry in entries:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.5
                pf.space_after = Pt(4)
                run = p.add_run(entry)
                PE._set_run_font(run, size=12)
            continue

        # Free-text body chapters
        body = ((chapters.get(cid) or {}).get("text") or "").strip()
        chapter_assets = assets.get(cid) or []
        if not body and not chapter_assets:
            continue
        doc.add_page_break()
        PE._add_heading(doc, label, level=1)
        if body:
            PE._render_section_text(doc, body)
        if chapter_assets:
            _docx_render_chapter_assets(doc, chapter_assets)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _pdf_render_chapter_assets(story: List[Any], styles, items: List[Dict[str, Any]]) -> None:
    cap_style = ParagraphStyle("Cap", parent=styles["tsmall"],
                               fontName="Times-Italic", fontSize=11, leading=14)
    for idx, item in enumerate(items, 1):
        try:
            img = RLImage(io.BytesIO(item["bytes"]))
            # Scale to fit a 5.5-inch column while preserving aspect ratio.
            iw, ih = img.imageWidth, img.imageHeight
            target_w = 5.5 * inch
            scale = min(1.0, target_w / float(iw or 1))
            img.drawWidth = float(iw) * scale
            img.drawHeight = float(ih) * scale
            story.append(img)
        except Exception as exc:                                  # noqa: BLE001
            log.warning("thesis_export: pdf picture embed failed: %s", exc)
            continue
        story.append(Paragraph(PE._pdf_safe(item.get("caption") or f"Figure {idx}"),
                               cap_style))
        story.append(Spacer(1, 0.15 * inch))


def build_pdf(payload: Dict[str, Any]) -> bytes:
    state      = payload.get("state")      or {}
    title_meta = payload.get("title_meta") or state.get("title_meta") or {}
    chapters   = state.get("chapters")     or {}
    references = state.get("references")   or []
    consent    = payload.get("consent")    or state.get("consent") or {}
    assets     = _normalised_assets(payload.get("assets") or state.get("assets") or {})

    languages = consent.get("languages") or []
    delivery  = (consent.get("delivery") or "attached").strip().lower()
    intake    = _intake_from_state(state)

    buf = io.BytesIO()
    doc = PE._PropPdfDoc(
        buf, pagesize=LETTER,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
        title=(title_meta.get("study_title") or intake.get("topic") or "Thesis"),
    )
    styles = PE._pdf_styles()
    story: List[Any] = []

    PE._pdf_title_page(story, styles, intake, title_meta)
    story.append(PageBreak())

    story.append(Paragraph("Table of Contents", styles["h1"]))
    toc = TableOfContents()
    toc.levelStyles = [ParagraphStyle(
        name="TOCLevel0", fontName="Times-Roman", fontSize=12,
        leading=18, leftIndent=0, firstLineIndent=0, spaceAfter=2,
    )]
    story.append(toc)
    story.append(PageBreak())

    for ch in CHAPTER_SPINE:
        cid = ch["id"]; label = ch["label"]
        if cid == "title_page":
            continue
        if cid == "consent":
            if delivery == "separate":
                continue
            for lang_label, text in PE._consent_packs(intake, title_meta, languages):
                story.append(PageBreak())
                story.append(Paragraph(PE._pdf_safe(
                    f"Informed Consent Form — {lang_label}"), styles["h1"]))
                for line in text.split("\n"):
                    story.append(Paragraph(PE._pdf_safe(line) or "&nbsp;",
                                           styles["pre"]))
            continue
        if cid == "references":
            entries = PE._bibliography_entries(references)
            if not entries:
                continue
            story.append(PageBreak())
            story.append(Paragraph(label, styles["h1"]))
            for entry in entries:
                story.append(Paragraph(PE._pdf_safe(entry), styles["pre"]))
            continue

        body = ((chapters.get(cid) or {}).get("text") or "").strip()
        chapter_assets = assets.get(cid) or []
        if not body and not chapter_assets:
            continue
        story.append(PageBreak())
        story.append(Paragraph(label, styles["h1"]))
        if body:
            for chunk in PE._pdf_paragraph_chunks(body):
                story.append(Paragraph(PE._pdf_safe(chunk), styles["body"]))
        if chapter_assets:
            _pdf_render_chapter_assets(story, styles, chapter_assets)

    doc.multiBuild(story, onFirstPage=PE._pdf_footer, onLaterPages=PE._pdf_footer)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# ZIP + plaintext
# ---------------------------------------------------------------------------

def _consent_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Build the smaller payload that proposal_export.build_consent_{docx,pdf}
    expects: just intake + title_meta + consent_languages."""
    state = payload.get("state") or {}
    consent = payload.get("consent") or state.get("consent") or {}
    return {
        "intake":             _intake_from_state(state),
        "title_meta":         payload.get("title_meta") or state.get("title_meta") or {},
        "consent_languages":  consent.get("languages") or [],
    }


def build_zip(payload: Dict[str, Any]) -> bytes:
    state = payload.get("state") or {}
    consent = payload.get("consent") or state.get("consent") or {}
    delivery = (consent.get("delivery") or "attached").strip().lower()
    docx_bytes = build_docx(payload)
    pdf_bytes  = build_pdf(payload)
    title = ((payload.get("title_meta") or state.get("title_meta") or {})
             .get("study_title") or "thesis")
    safe = _safe_filename(title, "")[:60].rstrip(".") or "thesis"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{safe}.docx", docx_bytes)
        zf.writestr(f"{safe}.pdf",  pdf_bytes)
        if delivery in {"separate", "both"}:
            cp = _consent_payload(payload)
            zf.writestr(f"{safe}__consent.docx", PE.build_consent_docx(cp))
            zf.writestr(f"{safe}__consent.pdf",  PE.build_consent_pdf(cp))
    return buf.getvalue()


def build_plaintext(payload: Dict[str, Any]) -> str:
    """Plain-text bundle of every chapter + bibliography. Used by the
    "Send to Plagiarism Checker" handoff."""
    state      = payload.get("state")      or {}
    chapters   = state.get("chapters")     or {}
    references = state.get("references")   or []
    parts: List[str] = []
    for ch in CHAPTER_SPINE:
        cid = ch["id"]
        if cid in _STRUCTURAL_CHAPTERS:
            continue
        body = ((chapters.get(cid) or {}).get("text") or "").strip()
        if not body:
            continue
        parts.append(f"=== {ch['label']} ===\n\n{body}\n")
    entries = PE._bibliography_entries(references)
    if entries:
        parts.append("=== References ===\n\n" + "\n".join(entries))
    return "\n\n".join(parts)
