"""Render thesis Chapter V from Sigma's thesis_analysis_blueprint.

This module deliberately consumes the blueprint contract instead of the raw
statistical result text.  Detailed statistical subtables remain available to
other exports, but the main thesis chapter stays doctor-readable.
"""

from __future__ import annotations

import base64
import html
import io
import re
from copy import deepcopy
from datetime import date as _date
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


FORBIDDEN_MAIN_TABLE_TITLES = {
    "observed counts",
    "expected counts",
    "row percentages",
    "column percentages",
    "test summary",
}


_BINARY_MARKER_VARS = {"er", "pr", "ar", "her2", "her2neu", "egfr"}
_PRESENCE_MARKER_VARS = {"lvi", "ene", "necrosis", "dcis", "txinfiltratingl", "tumourinfiltratinglymphocytes"}
_T4B_DISPLAY_PATTERN = re.compile(
    r"Ulceration\s+and\s*/\s*or\s+ipsilateral satellite nodules\s+and\s*/\s*or\s+(?:edema|oedema)[^;|:]*",
    flags=re.IGNORECASE,
)

_NODE_VARIABLE_NAME_RE = re.compile(r"(?:node|nodal|lymph)", re.IGNORECASE)
_DATE_TEXT_RE = re.compile(r"^\s*(\d{4})[-/](\d{1,2})[-/](\d{1,2})(?:\s|$)")
_DERIVED_NODE_LABELS = {
    "positivenodes": "Positive nodes",
    "totalnodes": "Total nodes",
    "noderatio": "Node ratio",
}


def _recover_node_fraction(value: Any, variable: Any = "") -> str:
    """Recover Excel-corrupted node fractions stored as dates (e.g. ``1/17``
    saved as ``2026-01-17``) for display. Mirrors export.py's
    ``_excel_recover_node_fraction`` so Word/PDF/Excel show the same cleaned
    fraction instead of a raw date or a colon-truncated label fragment."""
    if not _NODE_VARIABLE_NAME_RE.search(_text(variable)):
        return ""
    match = _DATE_TEXT_RE.match(_text(value))
    if not match:
        return ""
    year, month, day = (int(part) for part in match.groups())
    try:
        _date(year, month, day)
    except ValueError:
        return ""
    pos, total = month, day
    if total <= 0 or pos > total:
        return ""
    return f"{pos}/{total}"


def _clean_variable_label(value: Any) -> str:
    text = _text(value)
    text = _T4B_DISPLAY_PATTERN.sub("Skin involvement / T4b features", text)
    if "_" in text and " " not in text:
        text = text.replace("_", " ")
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
    key = _variable_key(text)
    if key in _DERIVED_NODE_LABELS:
        return _DERIVED_NODE_LABELS[key]
    # "No of nodes involved" is a column-name abbreviation for "Number of",
    # not a Yes/No category value — must not be confused with the
    # value-level Yes/No -> Positive/Negative substitution in _display_value.
    text = re.sub(r"\bNo\.?\s+of\b", "Number of", text, flags=re.IGNORECASE)
    replacements = {
        "Her2Neu": "HER2",
        "HER2neu": "HER2",
        "Her2neu": "HER2",
        "Ki67": "Ki-67",
        "pT": "Pathological T stage",
        "pt": "Pathological T stage",
    }
    for raw, clean in replacements.items():
        text = re.sub(rf"\b{re.escape(raw)}\b", clean, text)
    for raw, clean in {"Er": "ER", "Pr": "PR", "Ar": "AR", "Egfr": "EGFR"}.items():
        text = re.sub(rf"\b{raw}\b", clean, text)
    text = re.sub(r"\bTx\s+infiltrating\s+L\b", "Tumour-infiltrating lymphocytes", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTumou?r site\s*/\s*quadrant\b", "Tumour quadrant", text, flags=re.IGNORECASE)
    text = re.sub(r"\bTumou?r site\b", "Tumour quadrant", text, flags=re.IGNORECASE)
    if "treatment timing" not in text.lower():
        text = re.sub(
            r"\bUpfront\s*/?\s*(?:vs\.?|versus)?\s*post[\s-]*chemo(?:therapy)?\b",
            "Treatment timing / upfront versus post-chemotherapy status",
            text, flags=re.IGNORECASE,
        )
    return text


def _variable_key(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", _text(value).lower())


def _clinical_category_label(variable: Any, category: Any, label_ctx: Optional[Dict[str, Any]] = None) -> str:
    var_key = _variable_key(variable)
    text = _display_value(category, label_ctx) if label_ctx else _text(category)
    recovered_fraction = _recover_node_fraction(text, variable)
    if recovered_fraction:
        return recovered_fraction
    low = text.strip().lower()
    low_compact = re.sub(r"\s+", "", low)
    if "histologicaltype" in var_key or var_key in {"grade", "histologicalgrade"}:
        match = re.search(r"(?:type|grade)?\s*([123])(?:\.0)?\b", low, flags=re.IGNORECASE)
        if match:
            return f"Grade {match.group(1)}"
    if "ki67" in var_key or "ki67" in _text(variable).lower() or "ki-67" in _text(variable).lower():
        if low_compact in {">=14", ">=14%", "=>14", "=>14%"}:
            return ">=14%"
        if low_compact in {"<14", "<14%", "<=14", "<=14%"}:
            return "<14%"
    if "molecular" in var_key and "subtype" in var_key:
        if low_compact in {"her2neu", "her2", "her2enriched", "her2-enriched"}:
            return "HER2-enriched"
    if "nodal" in var_key and low_compact in {"no", "n0"}:
        return "N0"
    if "her2" in var_key:
        if low_compact in {"negative", "neg", "no", "0", "1", "1+", "low"}:
            return "Negative/low"
        if low_compact in {"2", "2+", "equivocal"}:
            return "Equivocal (2+)"
        if low_compact in {"3", "3+", "positive", "postive", "yes", "present"}:
            return "Positive (3+)"
    if "egfr" in var_key:
        if low in {"positive", "postive", "yes", "present", "patchy positive"}:
            return "Positive"
        if low in {"negative", "no", "absent"}:
            return "Negative"
    if var_key == "dcis":
        if low in {"positive", "postive", "yes", "present", "high grade", "low grade", "intermediate grade"}:
            return "Present"
        if low in {"negative", "no", "absent"}:
            return "Absent"
    if "treatmenttiming" in var_key or "upfrontpostchemo" in var_key:
        if low_compact == "upfront" or low in {"negative", "no"}:
            return "Upfront"
        if low_compact in {"postchemo", "postchemotherapy"} or low in {"positive", "postive", "yes"}:
            return "Post-chemotherapy"
    if var_key in _PRESENCE_MARKER_VARS:
        if low in {"positive", "postive", "yes", "present"}:
            return "Present"
        if low in {"negative", "no", "absent", "abse"}:
            return "Absent"
    if var_key in _BINARY_MARKER_VARS or any(marker in var_key for marker in _BINARY_MARKER_VARS):
        if low in {"positive", "postive", "yes", "present"}:
            return "Positive"
        if low in {"negative", "no", "absent"}:
            return "Negative"
    return _clean_variable_label(text)


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.4g}"
    if isinstance(value, (list, tuple)):
        return "; ".join(_text(item) for item in value if _text(item))
    if isinstance(value, dict):
        return "; ".join(f"{_label(k)}: {_text(v)}" for k, v in value.items() if _text(v))
    text = html.unescape(str(value)).strip()
    replacements = {
        "Â±": "±",
        "ą": "±",
        "â‰¤": "≤",
        "â‰¥": "≥",
        "â€”": "—",
        "â€“": "–",
        "â€¢": "•",
        "Cramer's V": "Cramér's V",
        "Cramer’s V": "Cramér's V",
    }
    for raw, clean in replacements.items():
        text = text.replace(raw, clean)
    text = text.replace("Postive", "Positive")
    text = re.sub(r">=\s*14%?", ">=14%", text)
    return text


def _text_preserve_raw_labels(value: Any) -> str:
    """Same cleanup as _text() (encoding fixes, "Postive" typo fix) but
    without the blanket ">=14%" category-value normalisation, which would
    otherwise collapse deliberately distinct raw duplicate labels (e.g.
    ">=14%", ">=14", ">= 14%" quoted together in a duplicate-labels
    warning) into one repeated-looking value."""
    if value is None:
        return ""
    text = html.unescape(str(value)).strip()
    replacements = {
        "Â±": "±", "ą": "±", "â‰¤": "≤", "â‰¥": "≥", "â€”": "—", "â€“": "–",
        "â€¢": "•", "Cramer's V": "Cramér's V", "Cramer’s V": "Cramér's V",
    }
    for raw, clean in replacements.items():
        text = text.replace(raw, clean)
    return text.replace("Postive", "Positive")


def _label(value: Any) -> str:
    text = _text(value).replace("_", " ").strip()
    if not text:
        return ""
    acronyms = {"id": "ID", "roc": "ROC", "auc": "AUC", "pdf": "PDF", "docx": "DOCX"}
    words = [acronyms.get(part.lower(), part.capitalize()) for part in text.split()]
    return " ".join(words)


def _readable_study_design(value: Any) -> str:
    text = _text(value).replace("_", " ").strip().lower()
    mapping = {
        "cross sectional association": "Cross-sectional association study",
        "cohort prognostic association": "Cohort/prognostic association study",
        "two group comparison": "Two-group comparison study",
        "descriptive prevalence": "Descriptive/prevalence study",
        "diagnostic accuracy": "Diagnostic accuracy study",
        "reliability agreement": "Reliability/agreement study",
        "regression prediction": "Regression/prediction study",
        "repeated measures": "Repeated-measures study",
        "survival": "Survival/time-to-event study",
        "correlation": "Correlation study",
    }
    return mapping.get(text, _label(text))


def _outcome_label_context(blueprint: Dict[str, Any]) -> Dict[str, Any]:
    display = _text(blueprint.get("primary_outcome"))
    raw_values = set()
    core_figure_variables = set()
    outcome_values = {"yes", "no", "positive", "negative"}
    for finding in blueprint.get("significant_findings") or []:
        variable = _text(finding.get("variable") if isinstance(finding, dict) else "")
        for sep in (" vs ", " by "):
            if sep in variable:
                core_figure_variables.add(variable.split(sep, 1)[0].strip())
                break
    for section in blueprint.get("analysis_sections") or []:
        for table in section.get("tables") or []:
            if str(table.get("table_type") or "").startswith("continuous_or_group"):
                for variable in _source_variables(table):
                    if (
                        variable and variable != display
                        and variable.lower() not in outcome_values
                        and not _is_node_derived_variable(variable)
                    ):
                        core_figure_variables.add(variable)
        if not isinstance(section, dict) or section.get("section_id") != "primary_outcome_distribution":
            continue
        for figure in section.get("figures") or []:
            for variable in figure.get("source_variables") or []:
                if _text(variable) and _text(variable) != display:
                    raw_values.add(_text(variable))
        for table in section.get("tables") or []:
            for variable in table.get("source_variables") or []:
                if _text(variable) and _text(variable) != display:
                    raw_values.add(_text(variable))
    return {
        "display": display,
        "raw_values": sorted(raw_values, key=len, reverse=True),
        "core_figure_variables": sorted(core_figure_variables),
        "status_like": any(token in display.lower() for token in ("expression", "status", "marker", "positive", "negative")),
    }


def _display_value(value: Any, label_ctx: Optional[Dict[str, Any]]) -> str:
    text = _text(value)
    if not label_ctx:
        return _format_statistical_display(_clean_variable_label(text))
    display = _text(label_ctx.get("display"))
    if not display:
        return _format_statistical_display(_clean_variable_label(text))
    for raw in label_ctx.get("raw_values") or []:
        text = text.replace(raw, display)
    if label_ctx.get("status_like"):
        # Only convert "Yes"/"No" when they are a self-contained category
        # token (the whole text, or immediately followed by the ": "/"; "
        # bit-separator used to join category summaries), never when they
        # are part of a longer label/phrase such as "No of nodes involved"
        # or prose like "No significant association was found".
        text = re.sub(r"\bYes(?=\s*[:;]|\s*$)", "Positive", text)
        text = re.sub(r"\bNo(?=\s*[:;]|\s*$)", "Negative", text)
    marker = re.sub(r"\b(?:expression\s+status|status|expression|marker)\b.*$", "", display, flags=re.IGNORECASE).strip()
    marker = marker or "Marker"
    if re.fullmatch(r"Interpretation\s*[- ]?\s*site", text, flags=re.IGNORECASE):
        return f"{marker} staining localization"
    if re.fullmatch(r"Staining\s+Result", text, flags=re.IGNORECASE):
        return f"{marker} staining score pattern"
    return _format_statistical_display(_clean_variable_label(text))


def _is_node_derived_variable(value: Any) -> bool:
    key = re.sub(r"[^a-z0-9]+", "", _text(value).lower())
    return key in {"positivenodes", "totalnodes", "noderatio"}


def _format_statistical_display(value: Any, header: Any = "") -> str:
    text = _text(value)
    test_names = {
        "welch_ttest": "Welch's t-test",
        "welch_t_test": "Welch's t-test",
        "mann_whitney": "Mann-Whitney U test",
        "chi_square": "Chi-square test",
        "fisher_exact": "Fisher's exact test",
    }
    if text.strip().lower() in test_names:
        text = test_names[text.strip().lower()]
    text = re.sub(r"\bwelch[_\s-]*t[\s-]*test\b", "Welch's t-test", text, flags=re.IGNORECASE)
    text = re.sub(r"\bwelch\s+ttest\b", "Welch's t-test", text, flags=re.IGNORECASE)

    def _round_match(match: re.Match, digits: int) -> str:
        number = float(match.group(1))
        rendered = f"{number:.{digits}f}".rstrip("0").rstrip(".")
        return match.group(0).replace(match.group(1), rendered)

    text = re.sub(
        r"(?i)df\s*=\s*(-?\d+(?:\.\d+)?)",
        lambda match: _round_match(match, 2),
        text,
    )
    text = re.sub(
        r"(?i)Cohen(?:'|\u2019)s\s+d\s*=\s*(-?\d+(?:\.\d+)?)",
        lambda match: _round_match(match, 3),
        text,
    )
    return text


def _sanitize_thesis_claims(value: Any) -> str:
    text = _text(value)
    replacements = (
        (r"\bproves?\b", "supports"),
        (r"\bcauses?\b", "was associated with"),
        (r"\bpredicts?\b", "was associated with"),
        (r"\bprognostic significance\b", "statistical association"),
        (r"\bindependent association\b", "association"),
    )
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text


def _clean_finding_text(value: Any) -> str:
    text = _text(value)
    text = text.replace("Chi-square test: Chi-square test", "Chi-square test")
    for token in ("; p =", "; adjusted p", "; Cram", "; Cohen", " p =", " adjusted p"):
        if token in text:
            text = text.split(token, 1)[0].strip()
    return _sanitize_thesis_claims(text)


def _is_node_derived_association_insignificant(table: Dict[str, Any]) -> bool:
    variables = table.get("source_variables") or []
    if not any(_is_node_derived_variable(var) for var in variables):
        return False
    table_type = str(table.get("table_type") or "")
    if "association" not in table_type and "comparison" not in table_type:
        return False
    headers, rows = _table_rows(table)
    return not _table_is_significant(headers, rows)


def _is_main_table(table: Dict[str, Any]) -> bool:
    if not isinstance(table, dict):
        return False
    if table.get("detailed_report_only"):
        return False
    if table.get("thesis_ready") is False:
        return False
    if str(table.get("priority") or "").lower() == "detailed_report_only":
        return False
    source_keys = {_variable_key(item) for item in table.get("source_variables") or []}
    table_key = _variable_key(table.get("title"))
    if ("txinfiltratingl" in source_keys or "txinfiltratingl" in table_key) and (
        table.get("optional") or str(table.get("priority") or "").lower() == "optional"
    ):
        return False
    if _is_node_derived_association_insignificant(table):
        return False
    placement = str(table.get("placement") or "main_thesis").lower()
    if placement and placement not in {"main_thesis", "thesis_preview", "chapter_v"}:
        return False
    title = _text(table.get("title")).lower()
    return title not in FORBIDDEN_MAIN_TABLE_TITLES


def _is_main_figure(figure: Dict[str, Any], include_optional: bool = False) -> bool:
    if not isinstance(figure, dict):
        return False
    if not figure.get("png_data_uri"):
        return False
    if figure.get("detailed_report_only"):
        return False
    if str(figure.get("priority") or "").lower() == "detailed_report_only":
        return False
    return bool(figure.get("thesis_ready") or str(figure.get("priority") or "").lower() == "thesis_ready_primary")


def _is_core_figure(figure: Dict[str, Any], label_ctx: Dict[str, Any], section_id: str = "") -> bool:
    if section_id == "primary_outcome_distribution":
        return True
    core = {str(item) for item in label_ctx.get("core_figure_variables") or []}
    if not core:
        return False
    fig_vars = {str(item) for item in figure.get("source_variables") or [] if item}
    title = _text(figure.get("title") or figure.get("caption"))
    display = _text(label_ctx.get("display"))
    non_outcome_vars = {var for var in fig_vars if var and var != display and var not in set(label_ctx.get("raw_values") or [])}
    return bool(non_outcome_vars.intersection(core) or any(variable and variable in title for variable in core))


def _association_figure_sort_key(figure: Dict[str, Any], label_ctx: Dict[str, Any]) -> Tuple[int, str]:
    title = _text(figure.get("title") or figure.get("caption")).lower()
    variables = " ".join(_text(item) for item in figure.get("source_variables") or []).lower()
    haystack = f"{title} {variables}"
    order = [
        ("age", 0),
        ("histological", 1),
        ("grade", 1),
        ("er", 2),
        ("pr", 3),
        ("molecular", 4),
        ("ar", 5),
    ]
    for token, rank in order:
        if re.search(rf"(^|[^a-z0-9]){re.escape(token)}([^a-z0-9]|$)", haystack):
            return rank, title
    core = {str(item).lower() for item in label_ctx.get("core_figure_variables") or []}
    if any(item and item in haystack for item in core):
        return 20, title
    return 100, title


def _normalise_figure_metadata(figure: Dict[str, Any], label_ctx: Dict[str, Any]) -> Dict[str, Any]:
    clone = deepcopy(figure)
    display = _text(label_ctx.get("display"))
    source_vars = [_display_value(item, label_ctx) for item in clone.get("source_variables") or []]
    predictor = next((item for item in source_vars if item and item != display), "")
    for key in ("title", "caption", "interpretation"):
        if clone.get(key):
            clone[key] = _display_value(clone[key], label_ctx)
    if display and predictor and (" by " in _text(clone.get("title")).lower() or " vs " in _text(clone.get("title")).lower()):
        clone["title"] = f"{display} by {predictor}"
        clone["caption"] = f"{display} by {predictor}."
    return clone


def _section_tables(section: Dict[str, Any]) -> List[Dict[str, Any]]:
    return [table for table in section.get("tables") or [] if _is_main_table(table)]


def _section_figures(section: Dict[str, Any], include_optional: bool = False) -> List[Dict[str, Any]]:
    return [figure for figure in section.get("figures") or [] if _is_main_figure(figure, include_optional)]


def _expand_export_tables(tables: List[Dict[str, Any]], label_ctx: Dict[str, Any]) -> List[Dict[str, Any]]:
    expanded: List[Dict[str, Any]] = []
    for table in tables:
        payload = _descriptive_export_table(table, label_ctx)
        if isinstance(payload, list):
            expanded.extend(payload)
        else:
            expanded.append(payload)
    return expanded


BASELINE_TERMS = ("age", "sex", "gender", "laterality", "side", "group", "cohort")

_GENERIC_PERCENTAGE_DENOMINATOR_NOTE = (
    "Percentages in detailed association tables are calculated within predictor categories "
    "unless otherwise stated."
)

_P27_PERCENTAGE_DENOMINATOR_NOTE = (
    "Percentages in detailed association tables are calculated within predictor categories "
    "unless otherwise stated. Percentages in the Significant Findings Highlight describe "
    "marker/category distribution within p27 expression groups."
)


def _p27_breast_context_from_blueprint(blueprint: Dict[str, Any]) -> bool:
    summary = blueprint.get("study_summary") or {}
    debug = blueprint.get("debug_metadata") or {}
    text_parts = [
        blueprint.get("primary_outcome"),
        summary.get("objective"),
        summary.get("domain_profile"),
        blueprint.get("study_design"),
        debug.get("canonical_outcome"),
        debug.get("mapped_outcome"),
        debug.get("confirmed_outcome_col"),
    ]
    text = " ".join(_text(part) for part in text_parts if _text(part)).lower()
    breast_context = (
        "breast_pathology" in text
        or "breast pathology" in text
        or any(term in text for term in ("breast", "mammary", "carcinoma"))
    )
    mentions_p27 = bool(re.search(r"\bp\s*[- ]?\s*27\b", text))
    has_positive_negative = False
    has_marker_component = False
    for section in blueprint.get("analysis_sections") or []:
        for table in section.get("tables") or []:
            table_type = _text(table.get("table_type")).lower()
            if "marker_component" in table_type:
                has_marker_component = True
            variables = " ".join(_text(var) for var in (table.get("source_variables") or []))
            if any(term in variables.lower() for term in ("staining", "localization", "localisation", "interpretation", "score")):
                has_marker_component = True
            rows_text = " ".join(_text(cell) for row in (table.get("rows") or []) for cell in (row if isinstance(row, (list, tuple)) else row.values() if isinstance(row, dict) else [row]))
            if "positive" in rows_text.lower() and "negative" in rows_text.lower():
                has_positive_negative = True
        for figure in section.get("figures") or []:
            variables = " ".join(_text(var) for var in (figure.get("source_variables") or []))
            if "positive" in variables.lower() and "negative" in variables.lower():
                has_positive_negative = True
    return breast_context and mentions_p27 and has_positive_negative and has_marker_component


def _percentage_denominator_note(blueprint: Dict[str, Any]) -> str:
    return (
        _P27_PERCENTAGE_DENOMINATOR_NOTE
        if _p27_breast_context_from_blueprint(blueprint)
        else _GENERIC_PERCENTAGE_DENOMINATOR_NOTE
    )

DESCRIPTIVE_TABLE_TYPES = {
    "descriptive_table",
    "domain_profile_descriptive_table",
    "marker_component_descriptive_table",
    "primary_outcome_distribution_table",
}


def _row_variable(row: Any) -> str:
    if isinstance(row, dict):
        return _text(row.get("variable"))
    if isinstance(row, (list, tuple)) and row:
        return _text(row[0])
    return ""


def _source_variables(table: Dict[str, Any]) -> List[str]:
    variables = [_text(var) for var in table.get("source_variables") or [] if _text(var)]
    if variables:
        return variables
    return [_row_variable(row) for row in table.get("rows") or [] if _row_variable(row)]


def _clean_table_title(title: Any) -> str:
    text = _text(title)
    return re.sub(r"^Table\s+\d+[A-Z]?\.\s*", "", text, flags=re.IGNORECASE).strip()


def _clean_interpretation(value: Any, label_ctx: Optional[Dict[str, Any]] = None) -> str:
    text = _display_value(value, label_ctx) if label_ctx else _text(value)
    if re.search(r"Molecular subtype-negative cases were more commonly Negative", text, flags=re.IGNORECASE):
        outcome = _text((label_ctx or {}).get("display")) or "the primary outcome"
        text = re.sub(
            r"Molecular subtype-negative cases were more commonly Negative\.?",
            f"The distribution of {outcome} varied across molecular subtype categories.",
            text,
            flags=re.IGNORECASE,
        )
    blocked = (
        "Domain-profile grouping is descriptive",
        "statistical tests remain generated from variable roles",
        "Use this figure to inspect outcome balance",
    )
    if any(item in text for item in blocked):
        return ""
    text = text.replace(
        "Marker or outcome-component variables are described here and omitted from final prognostic findings by default.",
        "Localization and staining score are presented descriptively as components of the immunohistochemical assessment.",
    )
    text = text.replace("Chi-square test with sparse-cell warning was", "was")
    text = text.replace("Chi-square test: Chi-square test", "Chi-square test")
    text = re.sub(
        r"Chi-square test with sparse-cell(?: warning)?\s+(?=[A-Z])",
        "A chi-square test was used, and sparse expected cell counts were noted. ",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r":\s*(?:Welch's t-test|Chi-square test|Fisher's exact test)\.(?=\s|$)",
        ".",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(.+?)\s+vs\s+(.+?):\s+was significantly associated with the outcome\.?$",
        r"\1 showed a statistically significant association with \2.",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(.+?)\s+vs\s+(.+?):\s+was not significantly associated with the outcome\.?$",
        r"\1 did not show a statistically significant association with \2.",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(.+?)\s+vs\s+(.+?)\s+was significantly associated with the outcome\.?$",
        r"\1 showed a statistically significant association with \2.",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(.+?)\s+vs\s+(.+?)\s+was not significantly associated with the outcome\.?$",
        r"\1 did not show a statistically significant association with \2.",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"^(.+?)\s+by\s+(.+?):\s+Welch's t-test was not significantly associated with the outcome\.?$",
        r"\1 was not significantly associated with \2.",
        text,
        flags=re.IGNORECASE,
    )
    sparse = bool(re.search(
        r"sparse|expected(?:\s+cell)?\s+counts?|because some(?:\.|\s)|minimum expected",
        text,
        flags=re.IGNORECASE,
    ))
    text = re.sub(r"\s*Interpret with caution:\s*-\s*", " ", text, flags=re.IGNORECASE)
    if sparse:
        caution = "This finding should be interpreted cautiously because some expected cell counts were below 5."
        has_clean_sparse_clause = "sparse expected cell counts were noted" in text.lower()
        text = re.sub(
            r"This finding should be interpreted cautiously because some(?: expected cell counts were below 5)?\.?",
            " ",
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(
            r"Sparse (?:cells|categories)(?: were)? detected(?:;?\s*interpret with caution)?\.?",
            " ",
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(r"\s*Interpret with caution[:.]?\s*", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*Warning[:.]?\s*", " ", text, flags=re.IGNORECASE)
        text = re.sub(
            r"\b(?:minimum|min)\s+expected(?:\s+cell)?\s+count(?:\s+was|\s*=|\s*:)?\s*[0-9.]+(?:\s*<\s*5)?\.?",
            " ",
            text,
            flags=re.IGNORECASE,
        )
        text = re.sub(r"\b(?:some\s+)?expected(?:\s+cell)?\s+counts?\s+(?:were\s+)?(?:below|<)\s*5[^.]*\.?", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"Chi-square test with sparse-cell(?:\s+Chi-square used:)?\s*some\.?", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+", " ", text).strip(" .")
        text = re.sub(r"\s*:\s*(?:Chi-square test|Fisher's exact test)\.?$", "", text, flags=re.IGNORECASE)
        stripped = text.rstrip(". ")
        if has_clean_sparse_clause:
            text = f"{stripped}." if stripped else ""
        else:
            text = (f"{stripped}. {caution}" if stripped else caution).strip()
    text = re.sub(r"^(.+?):\s+was ", r"\1 was ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return _sanitize_thesis_claims(text)


_CATEGORY_BIT_RE = re.compile(
    r"^(.*):\s*([0-9]+(?:\.[0-9]+)?)\s*\(([0-9]+(?:\.[0-9]+)?)%\)$"
)


def _parse_category_summary(summary: Any, label_ctx: Optional[Dict[str, Any]] = None) -> List[Tuple[str, str, str]]:
    text = _display_value(summary, label_ctx) if label_ctx else _text(summary)
    rows: List[Tuple[str, str, str]] = []
    # Split on the "; " bit separator first, then greedily match the LAST
    # colon in each bit as the label/count boundary. This is required
    # because a category label may itself contain colons (e.g. a corrupted
    # timestamp like "2026-01-17 00:00:00"); a non-greedy/charclass split
    # on the first colon would otherwise truncate the label to a fragment
    # like "00" and silently merge unrelated categories together.
    for bit in text.split(";"):
        bit = bit.strip()
        if not bit:
            continue
        match = _CATEGORY_BIT_RE.match(bit)
        if not match:
            continue
        label, count, pct = match.groups()
        rows.append((_text(label).strip(), str(int(float(count))) if float(count).is_integer() else count, f"{pct}%"))
    return rows


def _parse_continuous_summary(summary: Any, label_ctx: Optional[Dict[str, Any]] = None) -> Optional[Dict[str, str]]:
    text = _display_value(summary, label_ctx) if label_ctx else _text(summary)
    mean_sd = re.search(r"([0-9.+-]+)\s*(?:±|Â±|ą|\+/-)\s*([0-9.+-]+)", text)
    n_match = re.search(r"\(n\s*=\s*([0-9]+)\)", text, flags=re.IGNORECASE)
    missing = re.search(r"Missing:\s*([0-9]+)\s*\(([0-9.]+)%\)", text, flags=re.IGNORECASE)
    median = re.search(r"Median(?:\s*[:=])?\s*([0-9.+-]+)", text, flags=re.IGNORECASE)
    minimum = re.search(r"(?:Minimum|Min)(?:\s*[:=])?\s*([0-9.+-]+)", text, flags=re.IGNORECASE)
    maximum = re.search(r"(?:Maximum|Max)(?:\s*[:=])?\s*([0-9.+-]+)", text, flags=re.IGNORECASE)
    if not (mean_sd or n_match or median or minimum or maximum):
        return None
    mean_sd_text = f"{mean_sd.group(1)} ± {mean_sd.group(2)}" if mean_sd else "-"
    missing_text = f"{missing.group(1)} ({missing.group(2)}%)" if missing else "-"
    return {
        "n": n_match.group(1) if n_match else "-",
        "mean_sd": mean_sd_text,
        "median": median.group(1) if median else "-",
        "min": minimum.group(1) if minimum else "-",
        "max": maximum.group(1) if maximum else "-",
        "missing": missing_text,
    }


_QUADRANT_CANONICAL = (
    ("central", "Central"),
    ("upper outer", "Upper outer"),
    ("upper inner", "Upper inner"),
    ("lower outer", "Lower outer"),
    ("lower inner", "Lower inner"),
    ("uoq", "Upper outer"),
    ("uiq", "Upper inner"),
    ("loq", "Lower outer"),
    ("liq", "Lower inner"),
)


def _quadrant_category_label(category: Any) -> str:
    text = re.sub(r"\bquadrant\b", "", _text(category), flags=re.IGNORECASE).strip().lower()
    for token, label in _QUADRANT_CANONICAL:
        if token in text:
            return label
    return "Other"


def _is_tumour_quadrant_variable(variable: Any) -> bool:
    key = _variable_key(variable)
    return "quadrant" in key or (("tumour" in key or "tumor" in key) and "site" in key)


def _score_sort_key(category: Any) -> Tuple[int, int, int]:
    text = _text(category).strip()
    match = re.match(r"^\s*(\d+)\s*\+\s*(\d+)\s*$", text)
    if match:
        return (0, int(match.group(1)), int(match.group(2)))
    if text.lower() == "missing":
        return (2, 0, 0)
    return (1, 0, 0)


def _categories_look_like_score_pattern(categories: List[str]) -> bool:
    return any(re.match(r"^\s*\d+\s*\+\s*\d+\s*$", _text(category).strip()) for category in categories)


def _build_categorical_rows(
    categorical_order: List[Tuple[str, str]],
    categorical_counts: Dict[Tuple[str, str], float],
    concept: str = "",
) -> List[List[str]]:
    totals: Dict[str, float] = {}
    variable_order: List[str] = []
    variable_categories: Dict[str, List[str]] = {}
    for variable, category in categorical_order:
        totals[variable] = totals.get(variable, 0.0) + categorical_counts[(variable, category)]
        if variable not in variable_categories:
            variable_order.append(variable)
            variable_categories[variable] = []
        variable_categories[variable].append(category)

    # Score-pattern categories (e.g. "1+1".."5+3") must sort in their
    # logical numeric order rather than insertion order, regardless of
    # whether other variables sharing this table are score patterns —
    # multi-variable descriptive tables (e.g. Excel's unsplit "Table 1")
    # mix score-pattern rows with ordinary categorical rows, so the sort
    # is applied per-variable instead of gating on the whole table's
    # "concept" (which only the dedicated single-variable score tables
    # used in Word/PDF set to "score").
    for variable in variable_order:
        categories = variable_categories[variable]
        if concept == "score" or _categories_look_like_score_pattern(categories):
            variable_categories[variable] = sorted(categories, key=_score_sort_key)

    rows: List[List[str]] = []
    for variable in variable_order:
        total = totals.get(variable, 0.0)
        for category in variable_categories[variable]:
            count = categorical_counts[(variable, category)]
            pct = (count / total * 100.0) if total else 0.0
            count_text = str(int(count)) if float(count).is_integer() else f"{count:g}"
            rows.append([variable, category, count_text, f"{pct:.1f}%"])
    return rows


_GENERIC_DESCRIPTIVE_TITLES = {
    "baseline and study characteristics",
    "descriptive findings",
    "continuous descriptive findings",
    "categorical descriptive findings",
    "other pathology characteristics",
    "other immunophenotype characteristics",
}


def _specific_or_generic_title(base_title: str, variables: List[str], fallback: str) -> str:
    """When a descriptive table's title is one of the known generic
    catch-all titles (e.g. 'Baseline and study characteristics') and the
    table describes exactly one variable, replace it with a specific
    '<Variable> distribution' title (e.g. 'Age distribution',
    'Laterality distribution') instead of letting Table 1 and Table 2
    share the same vague base title. Already-specific titles (set by the
    domain-profile concept classifiers) are left untouched."""
    base = _text(base_title).strip()
    unique = list(dict.fromkeys(_text(var) for var in variables if _text(var)))
    if base.lower() in _GENERIC_DESCRIPTIVE_TITLES and len(unique) == 1:
        return f"{unique[0]} distribution"
    return base or fallback


def _descriptive_export_table(table: Dict[str, Any], label_ctx: Dict[str, Any]) -> Any:
    if str(table.get("table_type") or "") not in DESCRIPTIVE_TABLE_TYPES:
        return table
    concept = _text(table.get("concept"))
    categorical_counts: Dict[Tuple[str, str], float] = {}
    categorical_order: List[Tuple[str, str]] = []
    continuous_rows: List[List[str]] = []
    for row in table.get("rows") or []:
        variable = _row_variable(row)
        if not variable:
            continue
        variable = _display_value(variable, label_ctx)
        cells = row.get("cells") if isinstance(row, dict) else None
        kind = _text(row.get("type") if isinstance(row, dict) else "")
        summary = cells[0] if isinstance(cells, list) and cells else ""
        continuous = _parse_continuous_summary(summary, label_ctx)
        if continuous and ("mean" in kind.lower() or "sd" in kind.lower() or not _parse_category_summary(summary, label_ctx)):
            continuous_rows.append([
                variable,
                continuous["n"],
                continuous["mean_sd"],
                continuous["median"],
                continuous["min"],
                continuous["max"],
                continuous["missing"],
            ])
            continue
        categories = _parse_category_summary(summary, label_ctx)
        if categories:
            for category, count, pct in categories:
                clean_category = _clinical_category_label(variable, category, label_ctx)
                if concept == "tumour_quadrant" or _is_tumour_quadrant_variable(variable):
                    clean_category = _quadrant_category_label(clean_category)
                key = (variable, clean_category)
                if key not in categorical_counts:
                    categorical_order.append(key)
                    categorical_counts[key] = 0.0
                categorical_counts[key] += float(count)
        elif summary:
            continuous_rows.append([variable, "-", _display_value(summary, label_ctx), "-", "-", "-", "-"])
    if not categorical_counts and not continuous_rows:
        return table
    polish = table.get("_polish_interpretation")
    clone = deepcopy(table)
    if continuous_rows and categorical_counts:
        continuous = deepcopy(table)
        continuous["columns"] = ["Parameter", "n", "Mean ± SD", "Median", "Minimum", "Maximum", "Missing n (%)"]
        continuous["rows"] = continuous_rows
        continuous["title"] = _specific_or_generic_title(
            _clean_table_title(continuous.get("title")), [row[0] for row in continuous_rows],
            "Continuous descriptive findings",
        )
        continuous["interpretation"] = polish or _descriptive_table_interpretation(continuous, label_ctx)
        categorical = deepcopy(table)
        categorical["columns"] = ["Parameter", "Category", "n", "%"]
        categorical["rows"] = _build_categorical_rows(categorical_order, categorical_counts, concept)
        categorical["title"] = _specific_or_generic_title(
            _clean_table_title(categorical.get("title")), [key[0] for key in categorical_order],
            "Categorical descriptive findings",
        )
        categorical["interpretation"] = polish or _descriptive_table_interpretation(categorical, label_ctx)
        for item in (continuous, categorical):
            item.pop("headers", None)
        return [continuous, categorical]
    if continuous_rows and not categorical_counts:
        clone["columns"] = ["Parameter", "n", "Mean ± SD", "Median", "Minimum", "Maximum", "Missing n (%)"]
        clone["rows"] = continuous_rows
        title_variables = [row[0] for row in continuous_rows]
    else:
        clone["columns"] = ["Parameter", "Category", "n", "%"]
        clone["rows"] = _build_categorical_rows(categorical_order, categorical_counts, concept)
        title_variables = [key[0] for key in categorical_order]
    clone.pop("headers", None)
    clone["title"] = _specific_or_generic_title(
        _clean_table_title(clone.get("title")), title_variables, "Descriptive findings"
    )
    clone["interpretation"] = polish or _descriptive_table_interpretation(clone, label_ctx)
    return clone


def _rows_by_parameter(rows: List[List[str]]) -> "Dict[str, List[List[str]]]":
    grouped: Dict[str, List[List[str]]] = {}
    for row in rows:
        if not isinstance(row, list) or not row:
            continue
        grouped.setdefault(_text(row[0]), []).append(row)
    return grouped


def _safe_float(value: Any) -> float:
    try:
        return float(str(value).replace(",", "").replace("%", ""))
    except (TypeError, ValueError):
        return 0.0


def _lower_first(text: str) -> str:
    if not text:
        return text
    if text[0].isupper() and (len(text) == 1 or text[1].islower()):
        return text[0].lower() + text[1:]
    return text


def _mode_rows(rows: List[List[str]]) -> List[List[str]]:
    best_count = -1.0
    best: List[List[str]] = []
    for row in rows:
        if not isinstance(row, list) or len(row) < 4:
            continue
        try:
            count = float(str(row[2]).replace(",", ""))
        except ValueError:
            continue
        if count > best_count:
            best_count = count
            best = [row]
        elif count == best_count:
            best.append(row)
    return best


def _ranked_rows(rows: List[List[str]]) -> List[List[str]]:
    valid = [row for row in rows if isinstance(row, list) and len(row) >= 4]
    return sorted(valid, key=lambda row: _safe_float(row[2]), reverse=True)


def _join_category_labels(labels: List[str]) -> str:
    cleaned = [label for label in labels if label]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    if len(cleaned) == 2:
        return f"{cleaned[0]} and {cleaned[1]}"
    return ", ".join(cleaned[:-1]) + f", and {cleaned[-1]}"


def _category_positive_row(rows: List[List[str]]) -> Optional[List[str]]:
    for row in rows:
        if not isinstance(row, list) or len(row) < 4:
            continue
        label = _text(row[1]).strip().lower()
        if label.startswith("positive") or label == ">=14%":
            return row
    return None


def _size_category_to_tstage(category: Any) -> Optional[str]:
    text = _text(category).lower().replace("≤", "<=").replace("≥", ">=")
    has_gt2 = bool(re.search(r">\s*2\b", text))
    has_le5 = bool(re.search(r"<=\s*5\b", text))
    has_le2 = bool(re.search(r"<=\s*2\b", text)) and not has_gt2
    has_gt5 = bool(re.search(r">\s*5\b", text))
    if has_le2:
        return "T1"
    if has_gt2 and has_le5:
        return "T2"
    if has_gt5:
        return "T3"
    return None


def _missingness_sentence(rows: List[List[str]]) -> str:
    for row in rows:
        if not isinstance(row, list) or len(row) < 4:
            continue
        if _text(row[1]).strip().lower() != "missing":
            continue
        count = _safe_float(row[2])
        if count > 0:
            noun = "case" if count == 1 else "cases"
            return f"Missing values were recorded for {_text(row[2])} {noun} ({_text(row[3])})."
    return ""


def _continuous_missingness_sentence(variable: str, missing_text: str) -> str:
    match = re.match(r"([0-9.]+)\s*\(([0-9.]+)%\)", _text(missing_text))
    if not match:
        return ""
    var_lower = _lower_first(variable)
    count = _safe_float(match.group(1))
    if count <= 0:
        return f"{variable} had no missing values in this sample."
    noun = "case" if count == 1 else "cases"
    return f"Missing data were recorded for {match.group(1)} {noun} ({match.group(2)}%) of {var_lower}."


def _association_transition_sentence(subject: str, label_ctx: Optional[Dict[str, Any]]) -> str:
    outcome = _text((label_ctx or {}).get("display"))
    if not outcome:
        return ""
    if subject.strip().lower() == "ar expression status":
        return f"AR expression status was subsequently assessed for association with {outcome}."
    return f"These {subject} were subsequently assessed for association with {outcome}."


def _generic_categorical_sentence(rows: List[List[str]]) -> str:
    named_rows = [row for row in rows if isinstance(row, list) and len(row) >= 4 and _text(row[1]).strip().lower() != "missing"]
    if not named_rows:
        return ""
    grouped = _rows_by_parameter(named_rows)
    missing = _missingness_sentence(rows)
    if len(grouped) == 1:
        parameter, param_rows = next(iter(grouped.items()))
        ranked = _ranked_rows(param_rows)
        if not ranked:
            return ""
        mode = ranked[0]
        sentence = (
            f"This table summarises the distribution of {parameter} in the analysed sample. "
            f"{_text(mode[1])} was the most common category, accounting for {_text(mode[2])} cases ({_text(mode[3])})"
        )
        if len(ranked) > 1 and _safe_float(ranked[1][2]) > 0:
            second = ranked[1]
            sentence += f", followed by {_text(second[1])} with {_text(second[2])} cases ({_text(second[3])})."
        else:
            sentence += "."
        if missing:
            sentence += " " + missing
        return sentence
    parts = []
    for parameter, param_rows in grouped.items():
        ranked = _ranked_rows(param_rows)
        if not ranked:
            continue
        mode = ranked[0]
        parts.append(f"{parameter} was most often {_text(mode[1])} ({_text(mode[2])} cases, {_text(mode[3])})")
    if not parts:
        return ""
    sentence = "This table summarises baseline categorical characteristics in the analysed sample. " + "; ".join(parts) + "."
    if missing:
        sentence += " " + missing
    return sentence


def _continuous_table_sentence(variable: str, n: str, mean_sd: str, median: str, minimum: str, maximum: str, missing_text: str) -> str:
    var_lower = _lower_first(variable)
    sentence = (
        f"{variable} was summarised as a continuous variable in the analysed sample. "
        f"The mean {var_lower} was {_text(mean_sd)}, with a median of {_text(median)} "
        f"and a range from {_text(minimum)} to {_text(maximum)}."
    )
    missing_sentence = _continuous_missingness_sentence(variable, missing_text)
    if missing_sentence:
        sentence += " " + missing_sentence
    return sentence


def _generic_continuous_sentence(rows: List[List[str]]) -> str:
    valid = [row for row in rows if isinstance(row, list) and len(row) >= 7 and _text(row[0])]
    if not valid:
        return ""
    if len(valid) == 1:
        row = valid[0]
        return _continuous_table_sentence(_text(row[0]), row[1], row[2], row[3], row[4], row[5], row[6])
    names = _join_category_labels([_text(row[0]) for row in valid])
    sentence = f"This table summarises {names} as continuous variables in the analysed sample."
    for row in valid:
        var_lower = _lower_first(_text(row[0]))
        sentence += (
            f" The mean {var_lower} was {_text(row[2])}, with a median of {_text(row[3])} "
            f"and a range from {_text(row[4])} to {_text(row[5])}."
        )
    return sentence


def _interpretation_tumour_quadrant(rows: List[List[str]]) -> str:
    named_rows = [row for row in rows if isinstance(row, list) and len(row) >= 4 and _text(row[1]).strip().lower() != "other"]
    ranked = _ranked_rows(named_rows)
    if not ranked:
        return ""
    mode = _mode_rows(named_rows)
    sentence = "This table summarises the distribution of tumour quadrant in the analysed sample. "
    labels = [_text(row[1]).strip().lower() for row in mode]
    n_text = _text(mode[0][2])
    pct_text = _text(mode[0][3])
    if len(mode) > 1:
        sentence += (
            f"The most common tumour locations were {_join_category_labels(labels)} quadrants, "
            f"each accounting for {n_text} cases ({pct_text})."
        )
    else:
        sentence += (
            f"The most common tumour location was the {labels[0]} quadrant, "
            f"accounting for {n_text} cases ({pct_text})"
        )
        if len(ranked) > 1 and ranked[1] not in mode and _safe_float(ranked[1][2]) > 0:
            second = ranked[1]
            sentence += f", followed by the {_text(second[1]).strip().lower()} quadrant with {_text(second[2])} cases ({_text(second[3])})."
        else:
            sentence += "."
    if any(_text(row[1]).strip().lower() == "other" for row in rows):
        sentence += " Less frequent and mixed quadrant locations were grouped under Other for descriptive clarity."
    return sentence


def _interpretation_tumour_size(rows: List[List[str]]) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    mode = ranked[0]
    category = _text(mode[1])
    sentence = (
        "This table summarises the distribution of tumour size in the analysed sample. "
        f"Most tumours measured {category}, corresponding to {_text(mode[2])} cases ({_text(mode[3])})."
    )
    tstage = _size_category_to_tstage(category)
    if tstage:
        sentence += f" This indicates that {tstage}-sized lesions formed the largest size group in the analysed sample."
    if len(ranked) > 1:
        second = ranked[1]
        sentence += f" The next most frequent size group was {_text(second[1])}, seen in {_text(second[2])} cases ({_text(second[3])})."
    return sentence


def _interpretation_pathological_t_stage(rows: List[List[str]]) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    mode = ranked[0]
    label = _text(mode[1])
    sentence = (
        "This table summarises the distribution of pathological T stage in the analysed sample. "
        f"{label} was the most frequent category, seen in {_text(mode[2])} cases ({_text(mode[3])})."
    )
    if len(ranked) > 1:
        second = ranked[1]
        sentence += f" {_text(second[1])} was the next most frequent stage, with {_text(second[2])} cases ({_text(second[3])})."
    return sentence


def _interpretation_nodal_status(rows: List[List[str]]) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    mode = ranked[0]
    label = _text(mode[1])
    sentence = (
        "This table summarises the distribution of nodal status in the analysed sample. "
        f"{label} was the most frequent nodal category, seen in {_text(mode[2])} cases ({_text(mode[3])})."
    )
    if label.strip().upper() == "N0":
        sentence += " Node-positive categories together constituted the remaining clinically relevant nodal disease burden."
    elif len(ranked) > 1:
        second = ranked[1]
        sentence += f" {_text(second[1])} was the next most frequent category, with {_text(second[2])} cases ({_text(second[3])})."
    return sentence


def _interpretation_nodal_burden(table: Dict[str, Any], rows: List[List[str]]) -> str:
    if rows and isinstance(rows[0], list) and len(rows[0]) >= 7:
        sentence = _generic_continuous_sentence(rows)
        return sentence or "Nodal burden parameters are summarised descriptively above using the available counts and ratios."
    sentence = _generic_categorical_sentence(rows)
    return sentence or "Nodal burden categories are summarised descriptively above using the available counts."


def _interpretation_adverse_features(rows: List[List[str]]) -> str:
    grouped = _rows_by_parameter(rows)
    present: List[Tuple[str, str, str]] = []
    for parameter, param_rows in grouped.items():
        for row in param_rows:
            if _text(row[1]).strip().lower() == "present":
                present.append((parameter, _text(row[2]), _text(row[3])))
                break
    if not present:
        return ""
    present.sort(key=lambda item: _safe_float(item[1]), reverse=True)
    top_count = _safe_float(present[0][1])
    top_group = [item for item in present if _safe_float(item[1]) == top_count]
    rest = [item for item in present if _safe_float(item[1]) != top_count]
    if len(top_group) > 1:
        names = _join_category_labels([item[0] for item in top_group])
        clause = f"{names} were each present in {top_group[0][1]} cases ({top_group[0][2]})"
    else:
        clause = f"{top_group[0][0]} was present in {top_group[0][1]} cases ({top_group[0][2]})"
    if rest:
        rest_clauses = [f"{item[0]} was present in {item[1]} cases ({item[2]})" for item in rest]
        clause += ", while " + _join_category_labels(rest_clauses)
    sentence = f"This table summarises adverse pathological features in the analysed sample. {clause}."
    missing = _missingness_sentence(rows)
    if missing:
        sentence += " " + missing
    return sentence


def _interpretation_hormone_receptor(rows: List[List[str]], label_ctx: Optional[Dict[str, Any]] = None) -> str:
    grouped = _rows_by_parameter(rows)
    parts = []
    positive_pcts = []
    for idx, (parameter, param_rows) in enumerate(grouped.items()):
        positive = _category_positive_row(param_rows)
        if not positive:
            continue
        positive_pcts.append(_safe_float(positive[3]))
        if idx == 0:
            parts.append(f"{parameter} positivity was observed in {_text(positive[2])} cases ({_text(positive[3])})")
        else:
            parts.append(f"{parameter} positivity in {_text(positive[2])} cases ({_text(positive[3])})")
    if not parts:
        return ""
    sentence = "This table describes the hormone receptor profile of the analysed sample. "
    sentence += " and ".join(parts) + "."
    if positive_pcts and (sum(positive_pcts) / len(positive_pcts)) > 50:
        sentence += " This indicates that hormone-receptor-positive tumours formed the majority of the sample."
    transition = _association_transition_sentence("hormone receptor variables", label_ctx)
    if transition:
        sentence += " " + transition
    return sentence


def _interpretation_her2_proliferation(rows: List[List[str]], label_ctx: Optional[Dict[str, Any]] = None) -> str:
    grouped = _rows_by_parameter(rows)
    sentences = ["This table describes HER2 status and proliferation marker expression in the analysed sample."]
    for parameter, param_rows in grouped.items():
        positive = _category_positive_row(param_rows)
        if not positive:
            continue
        label = _text(positive[1])
        if "her2" in parameter.lower():
            match = re.search(r"\(([^)]+)\)", label)
            qualifier = f" ({match.group(1)})" if match else ""
            sentences.append(f"HER2-positive{qualifier} expression was present in {_text(positive[2])} cases ({_text(positive[3])}).")
        else:
            sentences.append(f"{parameter} positivity was present in {_text(positive[2])} cases ({_text(positive[3])}).")
    if len(sentences) == 1:
        return ""
    sentence = " ".join(sentences)
    transition = _association_transition_sentence("HER2 and proliferation markers", label_ctx)
    if transition:
        sentence += " " + transition
    return sentence


def _interpretation_ar_expression(rows: List[List[str]], label_ctx: Optional[Dict[str, Any]] = None) -> str:
    positive = _category_positive_row(rows)
    if not positive:
        return ""
    sentence = (
        "This table describes AR expression status in the analysed sample. "
        f"AR positivity was observed in {_text(positive[2])} cases ({_text(positive[3])})."
    )
    if _safe_float(positive[3]) > 50:
        sentence += " This indicates that AR-positive tumours formed the majority of the sample."
    transition = _association_transition_sentence("AR expression status", label_ctx)
    if transition:
        sentence += " " + transition
    return sentence


def _interpretation_molecular_subtype(rows: List[List[str]], label_ctx: Optional[Dict[str, Any]] = None) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    mode = _mode_rows(rows)
    labels = _join_category_labels([_text(row[1]) for row in mode])
    sentence = (
        "This table describes the distribution of molecular subtype in the analysed sample. "
        f"The most common molecular subtype was {labels}, accounting for {_text(mode[0][2])} cases ({_text(mode[0][3])})."
    )
    if len(ranked) > 1 and ranked[1] not in mode:
        second = ranked[1]
        sentence += f" {_text(second[1])} was the next most frequent subtype, with {_text(second[2])} cases ({_text(second[3])})."
    transition = _association_transition_sentence("molecular subtype groups", label_ctx)
    if transition:
        sentence += " " + transition
    return sentence


_MARKER_LOCALIZATION_NOTE = (
    "Localization is presented descriptively as a component of the immunohistochemical assessment, "
    "not as an independent prognostic feature."
)
_MARKER_SCORE_NOTE = (
    "Staining score is presented descriptively as a component of the immunohistochemical assessment, "
    "not as an independent prognostic feature."
)


def _interpretation_marker_localization(rows: List[List[str]]) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    mode = ranked[0]
    label = _text(mode[1])
    sentence = (
        "This table summarises the staining localization pattern of the marker assessment in the analysed sample. "
        f"{label} localization was the most common pattern, observed in {_text(mode[2])} cases ({_text(mode[3])})."
    )
    if len(ranked) > 1:
        second = ranked[1]
        sentence += f" {_text(second[1])} localization was the next most frequent pattern, seen in {_text(second[2])} cases ({_text(second[3])})."
    return f"{sentence} {_MARKER_LOCALIZATION_NOTE}"


def _interpretation_marker_score(rows: List[List[str]]) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    mode = ranked[0]
    label = _text(mode[1])
    sentence = (
        "This table summarises the staining score pattern of the marker assessment in the analysed sample. "
        f"The most common staining score pattern was {label}, observed in {_text(mode[2])} cases ({_text(mode[3])})."
    )
    if len(ranked) > 1:
        second = ranked[1]
        sentence += f" The next most frequent score pattern was {_text(second[1])}, seen in {_text(second[2])} cases ({_text(second[3])})."
    return f"{sentence} {_MARKER_SCORE_NOTE}"


def _interpretation_primary_outcome(rows: List[List[str]], label_ctx: Optional[Dict[str, Any]] = None) -> str:
    ranked = _ranked_rows(rows)
    if not ranked:
        return ""
    display = _text((label_ctx or {}).get("display")) or _text(ranked[0][0])
    marker = re.sub(r"\b(?:expression\s+status|status|expression|marker)\b.*$", "", display, flags=re.IGNORECASE).strip()
    marker = marker or display
    sentence = f"This table shows the distribution of {display}, which was the primary marker outcome in this analysis."
    positive = next((row for row in ranked if _text(row[1]).strip().lower().startswith("positive")), None)
    if positive and len(ranked) > 1:
        negative = next((row for row in ranked if _text(row[1]).strip().lower().startswith("negative")), None)
        if negative:
            sentence += (
                f" {marker} positivity was observed in {_text(positive[2])} cases ({_text(positive[3])}), "
                f"while {_text(negative[2])} cases ({_text(negative[3])}) were {marker}-negative."
            )
        else:
            other = [
                row for row in ranked
                if row is not positive
                and _text(row[1]).strip().lower() != "missing"
                and _safe_float(row[2]) > 0
            ]
            if other:
                other_clause = "; ".join(
                    f"{_text(row[2])} cases ({_text(row[3])}) were classified as {_text(row[1])}" for row in other
                )
                sentence += (
                    f" {marker} positivity was observed in {_text(positive[2])} cases ({_text(positive[3])}), "
                    f"while {other_clause}."
                )
    else:
        top = ranked[0]
        sentence += f" {_text(top[1])} was the most frequent category, observed in {_text(top[2])} cases ({_text(top[3])})."
    sentence += " This distribution forms the basis for the subsequent association analyses."
    return sentence


_CONCEPT_INTERPRETERS = {
    "tumour_quadrant": lambda table, rows, label_ctx: _interpretation_tumour_quadrant(rows),
    "tumour_size": lambda table, rows, label_ctx: _interpretation_tumour_size(rows),
    "pathological_t_stage": lambda table, rows, label_ctx: _interpretation_pathological_t_stage(rows),
    "nodal_status": lambda table, rows, label_ctx: _interpretation_nodal_status(rows),
    "nodal_burden": lambda table, rows, label_ctx: _interpretation_nodal_burden(table, rows),
    "adverse_features": lambda table, rows, label_ctx: _interpretation_adverse_features(rows),
    "hormone_receptor": lambda table, rows, label_ctx: _interpretation_hormone_receptor(rows, label_ctx),
    "her2_proliferation": lambda table, rows, label_ctx: _interpretation_her2_proliferation(rows, label_ctx),
    "ar_expression": lambda table, rows, label_ctx: _interpretation_ar_expression(rows, label_ctx),
    "molecular_subtype": lambda table, rows, label_ctx: _interpretation_molecular_subtype(rows, label_ctx),
    "localization": lambda table, rows, label_ctx: _interpretation_marker_localization(rows),
    "score": lambda table, rows, label_ctx: _interpretation_marker_score(rows),
    "primary_outcome": lambda table, rows, label_ctx: _interpretation_primary_outcome(rows, label_ctx),
}


def _descriptive_table_interpretation(table: Dict[str, Any], label_ctx: Optional[Dict[str, Any]] = None) -> str:
    rows = table.get("rows") or []
    if not rows:
        return ""
    is_continuous = isinstance(rows[0], list) and len(rows[0]) >= 7
    concept = _text(table.get("concept"))
    interpreter = _CONCEPT_INTERPRETERS.get(concept)
    if interpreter:
        sentence = interpreter(table, rows, label_ctx)
        if sentence:
            return sentence
    if is_continuous:
        sentence = _generic_continuous_sentence(rows)
        if sentence:
            return sentence
    else:
        sentence = _generic_categorical_sentence(rows)
        if sentence:
            return sentence
    first_param = _text(rows[0][0]) if isinstance(rows[0], list) and rows[0] else ""
    if first_param:
        return f"This table summarises {first_param} in the analysed sample."
    return "This table summarises the analysed sample."


def _filter_table_variables(table: Dict[str, Any], allowed: Set[str], seen: Set[str]) -> Optional[Dict[str, Any]]:
    clone = deepcopy(table)
    rows = []
    for row in clone.get("rows") or []:
        variable = _row_variable(row)
        if variable and variable in seen:
            continue
        if allowed and variable and variable not in allowed:
            continue
        rows.append(row)
        if variable:
            seen.add(variable)
    if not rows:
        return None
    clone["rows"] = rows
    clone["source_variables"] = [var for var in _source_variables(clone) if not allowed or var in allowed]
    return clone


def _dedupe_descriptive_sections(sections: List[Dict[str, Any]], label_ctx: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    domain_ids = {
        "clinical_study_characteristics",
        "immunophenotype_characteristics",
        "marker_outcome_components",
    }
    has_domain_sections = any(section.get("section_id") in domain_ids for section in sections)
    marker_vars: Set[str] = set()
    outcome_vars: Set[str] = set(label_ctx.get("raw_values") or []) if label_ctx else set()
    if label_ctx and _text(label_ctx.get("display")):
        outcome_vars.add(_text(label_ctx.get("display")))
    for section in sections:
        if section.get("section_id") != "marker_outcome_components":
            continue
        for table in section.get("tables") or []:
            marker_vars.update(_source_variables(table))
    seen: Set[str] = set()
    out: List[Dict[str, Any]] = []
    for section in sections:
        section_id = section.get("section_id")
        clone = deepcopy(section)
        allowed: Set[str] = set()
        if section_id == "baseline_characteristics" and has_domain_sections:
            for table in clone.get("tables") or []:
                for var in _source_variables(table):
                    if any(term in var.lower() for term in BASELINE_TERMS):
                        allowed.add(var)
        new_tables = []
        for table in clone.get("tables") or []:
            if section_id != "marker_outcome_components" and marker_vars:
                clone_table = deepcopy(table)
                clone_table["rows"] = [
                    row for row in clone_table.get("rows") or []
                    if _row_variable(row) not in marker_vars
                ]
                table = clone_table
            if section_id != "marker_outcome_components" and outcome_vars:
                clone_table = deepcopy(table)
                clone_table["rows"] = [
                    row for row in clone_table.get("rows") or []
                    if _row_variable(row) not in outcome_vars
                ]
                table = clone_table
            filtered = _filter_table_variables(table, allowed, seen)
            if filtered:
                new_tables.append(filtered)
        clone["tables"] = new_tables
        if new_tables or clone.get("figures") or _text(clone.get("interpretation")):
            out.append(clone)
    return out


def _primary_outcome_tables(blueprint: Dict[str, Any], label_ctx: Dict[str, Any]) -> List[Dict[str, Any]]:
    display = _text(label_ctx.get("display"))
    raw_values = set(label_ctx.get("raw_values") or [])
    candidates = raw_values | ({display} if display else set())
    if not candidates:
        return []
    for section in blueprint.get("analysis_sections") or []:
        for table in section.get("tables") or []:
            for row in table.get("rows") or []:
                variable = _row_variable(row)
                if variable not in candidates:
                    continue
                clone = deepcopy(table)
                clone["title"] = f"Primary outcome distribution: {display or variable}"
                clone["table_type"] = "primary_outcome_distribution_table"
                clone["concept"] = "primary_outcome"
                clone["rows"] = [row]
                clone["source_variables"] = [variable]
                clone["interpretation"] = f"This table reports the distribution of {display or variable}."
                clone["warnings"] = []
                return [clone]
    return []


def _parse_distribution_counts(table: Dict[str, Any], label_ctx: Dict[str, Any]) -> Tuple[List[str], List[float]]:
    labels: List[str] = []
    counts: List[float] = []
    _, rows = _table_rows(table)
    for row in rows:
        # Prefer columns after index 1 (skip variable name + type/summary label).
        # If that slice is empty or blank, fall back to scanning all cells so
        # tables using a "Summary" column (index 1) are also parsed correctly.
        cells_to_scan = row[2:] if len(row) > 2 else row
        if not any(_text(cell) for cell in cells_to_scan):
            cells_to_scan = row[1:] if len(row) > 1 else row
        for cell in cells_to_scan:
            text = _display_value(cell, label_ctx)
            for label, count in re.findall(r"([^:;]+):\s*([0-9]+(?:\.[0-9]+)?)", text):
                clean_label = _display_value(label.strip(), label_ctx)
                if clean_label and clean_label not in labels:
                    labels.append(clean_label)
                    counts.append(float(count))
    return labels, counts


def _bar_chart_data_uri(labels: List[str], counts: List[float], title: str) -> str:
    if not labels or not counts:
        return ""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return ""
    fig, ax = plt.subplots(figsize=(5.5, 3.4), dpi=160)
    colours = ["#4B78A8", "#F58518", "#54A24B", "#B279A2", "#E45756"]
    bars = ax.bar(labels, counts, color=colours[:len(labels)], edgecolor="white")
    ax.set_title(title)
    ax.set_ylabel("n")
    ax.spines[["top", "right"]].set_visible(False)
    for bar, value in zip(bars, counts):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(counts) * 0.02, f"{value:g}", ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format="png", bbox_inches="tight")
    plt.close(fig)
    return "data:image/png;base64," + base64.b64encode(out.getvalue()).decode("ascii")


def _primary_outcome_figure(blueprint: Dict[str, Any], table: Dict[str, Any], label_ctx: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    labels, counts = _parse_distribution_counts(table, label_ctx)
    display = _text(label_ctx.get("display") or blueprint.get("primary_outcome") or "primary outcome")
    png = _bar_chart_data_uri(labels, counts, f"Distribution of {display}")
    if not png:
        return None
    return {
        "figure_id": "primary_outcome_distribution_generated",
        "title": f"Distribution of {display}",
        "caption": f"Distribution of {display}.",
        "interpretation": f"The figure shows the distribution of {display} across the analysed sample.",
        "source_variables": [display],
        "png_data_uri": png,
        "thesis_ready": True,
        "detailed_report_only": False,
    }


def _compact_table_for_export(
    table: Dict[str, Any], keep_notes_column: bool = False
) -> Tuple[Dict[str, Any], List[str]]:
    clone = deepcopy(table)
    headers = [_text(header) for header in (clone.get("columns") or clone.get("headers") or [])]
    warnings: List[str] = list(clone.get("warnings") or [])
    warning_indexes = [] if keep_notes_column else [
        idx for idx, header in enumerate(headers)
        if header.strip().lower() in {"warning", "warnings", "notes/warnings"}
    ]
    if len(headers) > 7 and warning_indexes:
        remove = set(warning_indexes)
        clone["columns"] = [header for idx, header in enumerate(headers) if idx not in remove]
        new_rows = []
        for row in clone.get("rows") or []:
            if isinstance(row, dict):
                row_warning = row.get("warnings") or row.get("warning") or row.get("notes_warnings")
                if row_warning and _text(row_warning) not in {"-", "None"}:
                    warnings.append(_text(row_warning))
                new_rows.append(row)
            elif isinstance(row, (list, tuple)):
                values = list(row)
                for idx in warning_indexes:
                    if idx < len(values) and _text(values[idx]) not in {"", "-", "None"}:
                        warnings.append(_text(values[idx]))
                new_rows.append([cell for idx, cell in enumerate(values) if idx not in remove])
            else:
                new_rows.append(row)
        clone["rows"] = new_rows
    return clone, list(dict.fromkeys(warnings))


def _parse_count_cell(value: Any) -> Optional[float]:
    text = _text(value)
    match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
    if not match:
        return None
    try:
        return float(match.group(1))
    except ValueError:
        return None


def _format_count_pct(count: float, total: float) -> str:
    count_text = str(int(count)) if float(count).is_integer() else f"{count:g}"
    pct = (count / total * 100.0) if total else 0.0
    return f"{count_text} ({pct:.1f}%)"


def _table_is_significant(headers: List[str], rows: List[List[str]]) -> bool:
    p_idx = next((idx for idx, header in enumerate(headers) if header.strip().lower() == "p-value"), None)
    if p_idx is None or not rows:
        return False
    text = _text(rows[0][p_idx]) if p_idx < len(rows[0]) else ""
    match = re.search(r"([0-9]*\.?[0-9]+)", text)
    if not match:
        return False
    try:
        return float(match.group(1)) < 0.05
    except ValueError:
        return False


def _qualified_category_label(predictor_label: str, category: Any) -> str:
    text = _text(category).strip()
    low = text.lower()
    if low in {"positive", "negative"}:
        return f"{predictor_label}-{low}"
    if low in {"present", "absent", "yes", "no"}:
        return f"{predictor_label} {low}"
    return text


def _association_direction_sentence(
    predictor_label: str,
    headers: List[str],
    count_indexes: List[int],
    order: List[str],
    grouped: Dict[str, List[Any]],
    label_ctx: Dict[str, Any],
) -> str:
    if len(count_indexes) != 2 or len(order) < 2:
        return ""
    outcome_labels = []
    for idx in count_indexes:
        label = _display_value(headers[idx], label_ctx)
        label = re.sub(r"\s*n\s*\(%\)\s*$", "", label, flags=re.IGNORECASE).strip()
        outcome_labels.append(label)
    predictor_key = _variable_key(predictor_label)
    marker = re.sub(
        r"\s+(?:expression\s+)?status$",
        "",
        _text(label_ctx.get("display")),
        flags=re.IGNORECASE,
    ).strip() or "outcome"
    negative_pos = next(
        (pos for pos, label in enumerate(outcome_labels) if "negative" in label.lower()),
        None,
    )
    if predictor_key in {"histologicaltype", "histologicalgrade", "grade"} and negative_pos is not None:
        grade_rates: Dict[int, float] = {}
        for category in order:
            match = re.search(r"(?:grade\s*)?([123])(?:\.0)?$", _text(category), flags=re.IGNORECASE)
            counts = [float(grouped[category][idx] or 0.0) for idx in count_indexes]
            if match and sum(counts):
                grade_rates[int(match.group(1))] = counts[negative_pos] / sum(counts)
        if all(grade in grade_rates for grade in (1, 2, 3)) and grade_rates[3] > max(grade_rates[1], grade_rates[2]):
            return (
                f"Grade 3 tumours had a higher proportion of {marker}-negative cases "
                "compared with Grade 1 and Grade 2 tumours."
            )
    if predictor_key == "molecularsubtype":
        valid_categories = [
            category for category in order
            if _text(category).strip().lower() not in {"positive", "negative", "present", "absent", "yes", "no"}
        ]
        positive_pos = next(
            (pos for pos, label in enumerate(outcome_labels) if "positive" in label.lower()),
            0,
        )
        if valid_categories:
            largest = max(
                valid_categories,
                key=lambda category: float(grouped[category][count_indexes[positive_pos]] or 0.0),
            )
            return (
                f"The distribution of {_text(label_ctx.get('display')) or 'the primary outcome'} varied across "
                f"molecular subtype categories, with {_text(largest)} forming the largest {marker}-positive subgroup."
            )
        return (
            f"The distribution of {_text(label_ctx.get('display')) or 'the primary outcome'} varied across "
            "molecular subtype categories."
        )
    col_totals = [sum(float(grouped[category][idx] or 0.0) for category in order) for idx in count_indexes]
    if not all(col_totals):
        return ""
    dominant: List[Optional[str]] = []
    for col_pos, idx in enumerate(count_indexes):
        best_category, best_share = None, -1.0
        for category in order:
            share = float(grouped[category][idx] or 0.0) / col_totals[col_pos]
            if share > best_share:
                best_share = share
                best_category = category
        dominant.append(best_category)
    dominant_first, dominant_last = dominant[0], dominant[-1]
    if not dominant_first or not dominant_last or dominant_first == dominant_last:
        return ""
    dominant_first = _qualified_category_label(predictor_label, dominant_first)
    dominant_last = _qualified_category_label(predictor_label, dominant_last)
    # Qualify the outcome group labels with the marker name too (e.g.
    # "p27-positive group"), not bare "Positive"/"Negative" — otherwise the
    # sentence reads ambiguously, as if the predictor's own category were
    # being repeated as the group name.
    qualified_last_outcome = _qualified_category_label(marker, outcome_labels[-1])
    qualified_first_outcome = _qualified_category_label(marker, outcome_labels[0])
    return (
        f"{dominant_last} cases were proportionately higher in the {qualified_last_outcome} group, "
        f"while {dominant_first} cases contributed a larger share of the {qualified_first_outcome} group."
    )


def _splice_sentences(interpretation: str, extra_sentences: List[str]) -> str:
    extras = [sentence for sentence in extra_sentences if sentence]
    if not extras:
        return interpretation
    joined = " ".join(extras)
    for marker in (" This finding should be interpreted cautiously", " Interpret with caution:"):
        idx = interpretation.find(marker)
        if idx != -1:
            head, tail = interpretation[:idx], interpretation[idx:]
            return f"{head.strip()} {joined}{tail}"
    return f"{interpretation.strip()} {joined}".strip()


def _effect_size_sentence(headers: List[str], row: List[str]) -> str:
    idx = next((i for i, header in enumerate(headers) if header.strip().lower() == "effect size"), None)
    if idx is None or idx >= len(row):
        return ""
    text = _text(row[idx]).strip()
    if not text or text == "-":
        return ""
    match = re.match(r"^(.+?)\s*=\s*(.+)$", text)
    if not match:
        return ""
    label, value = match.group(1).strip(), match.group(2).strip()
    article = "an" if label[:1].lower() in "aeiou" else "a"
    return f"The association had {article} {label} of {value}."


def _test_applied_clause(headers: List[str], row: List[str], significant: bool) -> str:
    idx = next((i for i, header in enumerate(headers) if header.strip().lower() == "test applied"), None)
    if idx is None or idx >= len(row):
        return ""
    test_name = _text(row[idx]).strip()
    if not test_name or test_name == "-":
        return ""
    qualifier = "" if significant else "not "
    return f"This difference was {qualifier}statistically significant using {test_name}."


def _group_comparison_means_sentence(
    predictor_label: str, headers: List[str], rows: List[List[str]], label_ctx: Dict[str, Any]
) -> str:
    mean_idx = next((i for i, header in enumerate(headers) if "mean" in header.lower()), None)
    if mean_idx is None or len(rows) < 2:
        return ""
    outcome_display = _text(label_ctx.get("display"))
    marker = re.sub(r"\s+(?:expression\s+)?status$", "", outcome_display, flags=re.IGNORECASE).strip()
    is_age = _variable_key(predictor_label) == "age"
    parts = []
    for row in rows[:2]:
        if len(row) <= mean_idx:
            continue
        group_label = _display_value(row[0], label_ctx)
        group_key = _variable_key(group_label)
        if marker and group_key in {"positive", "negative", "p27positive", "p27negative"}:
            polarity = "positive" if "positive" in group_key else "negative"
            group_label = f"{marker}-{polarity}"
        value = _text(row[mean_idx])
        if is_age and not re.search(r"\byears?\b", value, flags=re.IGNORECASE):
            value = f"{value} years"
        parts.append(f"{value} in the {group_label} group")
    if len(parts) < 2:
        return ""
    return f"The mean {_lower_first(predictor_label)} was {parts[0]} and {parts[1]}."


def _association_table_for_export(table: Dict[str, Any], label_ctx: Dict[str, Any]) -> Dict[str, Any]:
    table_type = str(table.get("table_type") or "").lower()
    if "association" not in table_type and "comparison" not in table_type:
        return table
    headers, rows = _table_rows(table)
    if not rows or not headers:
        return table
    first_header = headers[0].strip().lower()
    if first_header not in {"predictor category", "category", "group"}:
        return table

    source_vars = _source_variables(table)
    display = _text(label_ctx.get("display"))
    predictor = next((var for var in source_vars if var and var != display and var not in set(label_ctx.get("raw_values") or [])), "")
    if not predictor:
        title = _text(table.get("title"))
        match = re.search(r"with\s+(.+)$", title, flags=re.IGNORECASE)
        predictor = match.group(1).strip() if match else ""
    predictor_label = _clean_variable_label(predictor)
    significant = _table_is_significant(headers, rows)

    if first_header == "group":
        clone = deepcopy(table)
        clone["columns"] = [_display_value(header, label_ctx) for header in headers]
        clean_rows = [
            [
                _format_statistical_display(_display_value(cell, label_ctx), headers[idx] if idx < len(headers) else "")
                for idx, cell in enumerate(row)
            ]
            for row in rows
        ]
        clone["rows"] = clean_rows
        means = _group_comparison_means_sentence(predictor_label, headers, clean_rows, label_ctx)
        effect_sentence = _effect_size_sentence(headers, clean_rows[0]) if clean_rows else ""
        test_clause = _test_applied_clause(headers, clean_rows[0] if clean_rows else [], significant)
        clone["interpretation"] = _splice_sentences(
            _text(clone.get("interpretation")), [means, effect_sentence, test_clause]
        )
        return clone

    count_indexes = [
        idx for idx, header in enumerate(headers)
        if idx > 0 and ("n (%)" in header.lower() or "count" in header.lower())
    ]
    if not count_indexes:
        return table

    grouped: Dict[str, List[Any]] = {}
    order: List[str] = []
    for row in rows:
        if not row:
            continue
        category = _clinical_category_label(predictor, row[0], label_ctx)
        if category not in grouped:
            order.append(category)
            grouped[category] = list(row)
            grouped[category][0] = category
            for idx in count_indexes:
                grouped[category][idx] = 0.0
        for idx in count_indexes:
            grouped[category][idx] = float(grouped[category][idx] or 0.0) + float(_parse_count_cell(row[idx]) or 0.0)

    direction = ""
    if significant:
        direction = _association_direction_sentence(predictor_label, headers, count_indexes, order, grouped, label_ctx)

    if len(order) == len(rows) and all(_text(row[0]) == order[idx] for idx, row in enumerate(rows)):
        clone = deepcopy(table)
        clone["columns"] = [_display_value(header, label_ctx) for header in headers]
        clean_rows = []
        for row in rows:
            clean_row = [
                _format_statistical_display(_display_value(cell, label_ctx), headers[idx] if idx < len(headers) else "")
                for idx, cell in enumerate(row)
            ]
            if clean_row:
                clean_row[0] = _clinical_category_label(predictor, row[0], label_ctx)
            clean_rows.append(clean_row)
        clone["rows"] = clean_rows
        effect_sentence = _effect_size_sentence(headers, clean_rows[0]) if clean_rows else ""
        clone["interpretation"] = _splice_sentences(_text(clone.get("interpretation")), [direction, effect_sentence])
        return clone

    out_rows: List[List[str]] = []
    for category in order:
        row = grouped[category]
        total = sum(float(row[idx] or 0.0) for idx in count_indexes)
        normalized = [
            _format_statistical_display(_display_value(cell, label_ctx), headers[idx] if idx < len(headers) else "")
            for idx, cell in enumerate(row)
        ]
        normalized[0] = category
        for idx in count_indexes:
            normalized[idx] = _format_count_pct(float(row[idx] or 0.0), total)
        out_rows.append(normalized)

    clone = deepcopy(table)
    clone["columns"] = [_display_value(header, label_ctx) for header in headers]
    clone["rows"] = out_rows
    effect_sentence = _effect_size_sentence(headers, out_rows[0]) if out_rows else ""
    clone["interpretation"] = _splice_sentences(_text(clone.get("interpretation")), [direction, effect_sentence])
    return clone


def _normalise_table_for_export(table: Dict[str, Any], label_ctx: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    ctx = label_ctx or {}
    payload = _association_table_for_export(table, ctx)
    if payload is table:
        payload = deepcopy(table)
        headers, rows = _table_rows(payload)
        payload["columns"] = [_display_value(header, ctx) for header in headers]
        payload["rows"] = [
            [
                _format_statistical_display(_display_value(cell, ctx), headers[idx] if idx < len(headers) else "")
                for idx, cell in enumerate(row)
            ]
            for row in rows
        ]
    return payload


def _table_rows(table: Dict[str, Any]) -> tuple[List[str], List[List[str]]]:
    headers = [_text(header) for header in (table.get("columns") or table.get("headers") or [])]
    rows: List[List[str]] = []
    for row in table.get("rows") or []:
        if isinstance(row, dict):
            if headers and isinstance(row.get("cells"), list):
                values: List[str] = []
                cell_values = [_text(cell) for cell in row.get("cells") or []]
                cell_idx = 0
                for header in headers:
                    normalized = header.strip().lower()
                    if normalized == "variable":
                        values.append(_text(row.get("variable")))
                    elif normalized == "type":
                        values.append(_text(row.get("type")))
                    elif normalized in {"p", "p-value", "p value"}:
                        values.append(_text(row.get("p") or row.get("p_value")))
                    else:
                        values.append(cell_values[cell_idx] if cell_idx < len(cell_values) else _text(row.get(header)))
                        cell_idx += 1
                rows.append(values)
            elif headers:
                rows.append([_text(row.get(header) or row.get(header.lower()) or row.get(header.replace(" ", "_").lower())) for header in headers])
            else:
                rows.extend([_text(k), _text(v)] for k, v in row.items())
        elif isinstance(row, (list, tuple)):
            rows.append([_text(cell) for cell in row])
        else:
            rows.append([_text(row)])
    if not headers and rows:
        headers = [f"Column {idx + 1}" for idx in range(max(len(row) for row in rows))]
    width = len(headers) or (max((len(row) for row in rows), default=1))
    normalized = [row + [""] * (width - len(row)) for row in rows]
    return headers or ["Value"], normalized


def _plain_docx_text(doc: Document, text: str, style: Optional[str] = None):
    para = doc.add_paragraph(style=style)
    para.add_run(_text(text))
    return para


def _set_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.bold = True


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(_text(text), level=level)


def _repeat_docx_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _format_docx_table(table, font_size: int = 8) -> None:
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def _add_table_docx(
    doc: Document,
    payload: Dict[str, Any],
    caption_no: int,
    label_ctx: Optional[Dict[str, Any]] = None,
    include_interpretation: bool = True,
    include_warnings: bool = True,
    keep_notes_column: bool = False,
) -> int:
    payload = _descriptive_export_table(payload, label_ctx or {})
    if isinstance(payload, list):
        payload = payload[0] if payload else {}
    payload = _normalise_table_for_export(payload, label_ctx or {})
    payload, warning_notes = _compact_table_for_export(payload, keep_notes_column=keep_notes_column)
    headers, rows = _table_rows(payload)
    if not rows:
        return caption_no
    caption = _display_value(_clean_table_title(payload.get("title") or f"Table {caption_no}"), label_ctx)
    _plain_docx_text(doc, f"Table {caption_no}. {caption}", style=None).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = _display_value(header, label_ctx)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    _repeat_docx_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[:len(headers)]):
            cells[idx].text = _display_value(value, label_ctx)
    _format_docx_table(table, font_size=8 if len(headers) > 5 else 9)
    if include_warnings:
        for warning in warning_notes:
            cleaned_warning = _clean_interpretation(warning, label_ctx)
            if cleaned_warning:
                _plain_docx_text(doc, cleaned_warning)
    interpretation = _clean_interpretation(payload.get("interpretation"), label_ctx)
    if include_interpretation and interpretation:
        _plain_docx_text(doc, interpretation)
    doc.add_paragraph()
    return caption_no + 1


def _strip_data_uri(data_uri: str) -> bytes:
    if not data_uri:
        return b""
    if "," in data_uri:
        data_uri = data_uri.split(",", 1)[1]
    try:
        return base64.b64decode(data_uri)
    except Exception:
        return b""


def _add_figure_docx(doc: Document, figure: Dict[str, Any], figure_no: int, label_ctx: Optional[Dict[str, Any]] = None) -> int:
    title = _display_value(figure.get("title") or f"Figure {figure_no}", label_ctx)
    png = _strip_data_uri(_text(figure.get("png_data_uri")))
    if not png:
        return figure_no
    if png:
        try:
            doc.add_picture(io.BytesIO(png), width=Inches(5.7))
        except Exception:
            return figure_no
    caption = _display_value(figure.get("caption") or title, label_ctx)
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run(f"Figure {figure_no}. {caption}")
    caption_run.italic = True
    interpretation = _clean_interpretation(figure.get("interpretation"), label_ctx)
    if interpretation:
        _plain_docx_text(doc, interpretation)
    doc.add_paragraph()
    return figure_no + 1


def _figure_matches_table(figure: Dict[str, Any], table: Dict[str, Any]) -> bool:
    table_ids = {str(item) for item in table.get("source_test_ids") or [] if item}
    fig_result = str(figure.get("source_result_id") or "")
    if table_ids and fig_result:
        return fig_result in table_ids
    table_vars = set(_source_variables(table))
    fig_vars = {str(item) for item in figure.get("source_variables") or [] if item}
    return bool(table_vars and fig_vars and table_vars.intersection(fig_vars))


def _enrich_figures_from_tables(
    figures: List[Dict[str, Any]],
    tables: List[Dict[str, Any]],
    label_ctx: Dict[str, Any],
) -> List[Dict[str, Any]]:
    enriched_figures: List[Dict[str, Any]] = []
    for figure in figures:
        clone = deepcopy(figure)
        matched = next((table for table in tables if _figure_matches_table(clone, table)), None)
        if matched:
            enriched_table = _normalise_table_for_export(matched, label_ctx)
            interpretation = _clean_interpretation(enriched_table.get("interpretation"), label_ctx)
            warnings = " ".join(_text(item) for item in enriched_table.get("warnings") or [] if _text(item))
            if warnings:
                interpretation = _clean_interpretation(f"{interpretation} {warnings}", label_ctx)
            if interpretation:
                clone["interpretation"] = interpretation
        enriched_figures.append(clone)
    return enriched_figures


def _render_section_docx(
    doc: Document,
    section: Dict[str, Any],
    table_no: int,
    figure_no: int,
    label_ctx: Dict[str, Any],
    include_optional_figures: bool = False,
) -> tuple[int, int]:
    section_id = str(section.get("section_id") or "")
    figures = [
        _normalise_figure_metadata(figure, label_ctx)
        for figure in _section_figures(section, include_optional_figures)
        if include_optional_figures or _is_core_figure(figure, label_ctx, section_id)
    ]
    used_figures: Set[str] = set()
    tables = _expand_export_tables(_section_tables(section), label_ctx)
    figures = _enrich_figures_from_tables(figures, tables, label_ctx)
    if section_id == "bivariate_associations":
        for figure in sorted(figures, key=lambda item: _association_figure_sort_key(item, label_ctx)):
            figure_id = _text(figure.get("figure_id") or figure.get("title"))
            if figure_id in used_figures:
                continue
            next_no = _add_figure_docx(doc, figure, figure_no, label_ctx)
            if next_no != figure_no:
                figure_no = next_no
                used_figures.add(figure_id)
        if tables:
            _add_heading(doc, "Detailed Association Tables", 2)
    for table in tables:
        matched_rendered_figure = section_id == "bivariate_associations" and any(
            _text(figure.get("figure_id") or figure.get("title")) in used_figures
            and _figure_matches_table(figure, table)
            for figure in figures
        )
        table_no = _add_table_docx(
            doc,
            table,
            table_no,
            label_ctx,
            include_interpretation=False,
            include_warnings=not matched_rendered_figure,
        )
        for figure in figures:
            figure_id = _text(figure.get("figure_id") or figure.get("title"))
            if figure_id in used_figures or not _figure_matches_table(figure, table):
                continue
            next_no = _add_figure_docx(doc, figure, figure_no, label_ctx)
            if next_no != figure_no:
                figure_no = next_no
                used_figures.add(figure_id)
        if not matched_rendered_figure:
            enriched = _normalise_table_for_export(table, label_ctx)
            interpretation = _clean_interpretation(enriched.get("interpretation"), label_ctx)
            if interpretation:
                _plain_docx_text(doc, interpretation)
    if include_optional_figures:
        for figure in figures[:4]:
            figure_id = _text(figure.get("figure_id") or figure.get("title"))
            if figure_id in used_figures:
                continue
            next_no = _add_figure_docx(doc, figure, figure_no, label_ctx)
            if next_no != figure_no:
                figure_no = next_no
                used_figures.add(figure_id)
    return table_no, figure_no


def _blueprint(results: Dict[str, Any]) -> Dict[str, Any]:
    blueprint = results.get("thesis_analysis_blueprint")
    if not isinstance(blueprint, dict) or not blueprint:
        raise ValueError("Chapter V export requires thesis_analysis_blueprint. Run analysis again before exporting.")
    return blueprint


def _study_summary(blueprint: Dict[str, Any], results: Dict[str, Any]) -> List[tuple[str, str]]:
    summary = dict(blueprint.get("study_summary") or {})
    rows = [
        ("Study design", _readable_study_design(blueprint.get("study_design") or summary.get("study_design"))),
        ("Sample size", summary.get("sample_size") or summary.get("n") or ""),
        ("Primary outcome", blueprint.get("primary_outcome") or summary.get("primary_outcome") or ""),
    ]
    return [(label, _text(value)) for label, value in rows if _text(value)]


def _add_summary_docx(doc: Document, blueprint: Dict[str, Any], results: Dict[str, Any]) -> None:
    rows = _study_summary(blueprint, results)
    if not rows:
        _plain_docx_text(doc, "Study summary metadata was not available.")
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def _main_sections(blueprint: Dict[str, Any]) -> List[Dict[str, Any]]:
    skipped = {"tested_associations_summary", "significant_findings_summary"}
    return [
        section for section in blueprint.get("analysis_sections") or []
        if isinstance(section, dict)
        and section.get("section_id") not in skipped
        and (_section_tables(section) or _section_figures(section) or _text(section.get("interpretation")))
    ]


def _section_export_title(section_id: str, fallback: Any, label_ctx: Dict[str, Any]) -> str:
    marker = _text(label_ctx.get("display") or "Marker")
    titles = {
        "baseline_characteristics": "Section I - Baseline Characteristics",
        "clinical_study_characteristics": "Section II - Nodal and Prognostic Pathology",
        "immunophenotype_characteristics": "Section III - Immunophenotype",
        "marker_outcome_components": f"Section IV - {marker} / Marker Expression",
        "primary_outcome_distribution": f"Section IV - {marker} / Marker Expression",
        "bivariate_associations": "Section V - Statistical Associations",
        "correlation_analysis": "Section V - Correlation Analysis",
        "regression_adjusted_analysis": "Section V - Regression / Adjusted Analysis",
        "diagnostic_accuracy": "Section V - Diagnostic Accuracy",
        "reliability_agreement": "Section V - Reliability and Agreement",
        "survival_analysis": "Section V - Survival Analysis",
        "repeated_measures": "Section V - Repeated-Measures Analysis",
        "other_analyses": "Section V - Statistical Analyses",
    }
    return titles.get(section_id, _text(fallback) or "Results")


def _user_facing_warning(value: Any) -> str:
    raw_text = re.sub(r"\s+", " ", str(value or "")).strip(" .")
    if raw_text.startswith("Duplicate raw ") and "category labels (" in raw_text:
        # Already fully formatted by thesis_blueprint._clean_report_warning.
        # Skip the generic _text() cleanup below — its blanket ">=14%"
        # category-value normalisation would otherwise collapse the
        # deliberately distinct raw duplicate labels quoted here (e.g.
        # ">=14%", ">=14", ">= 14%") into one repeated-looking value.
        return f"{raw_text}." if raw_text else ""
    text = re.sub(r"\s+", " ", _text(value)).strip(" .")
    lower = text.lower()
    if not text:
        return ""
    if "add only after confirming predictors" in lower:
        # Objective-routing/implementation detail: not clinically actionable
        # for a thesis reader. Keep out of main Word/PDF; raw text remains
        # in the Excel "Data Cleaning Log" audit sheet.
        return ""
    if "run separately under the correlation objective" in lower:
        return ""
    if "duplicate" in lower and "predictor" in lower:
        # A non-specific duplicate-predictor-labels warning reaching this far
        # (the blueprint-level cleanup normally already makes it specific or
        # drops it) is not actionable without naming the variable/labels
        # involved — keep it out of the main Word/PDF report rather than
        # showing a vague "category grouping should be reviewed" sentence.
        # The raw detail remains available in the Excel audit sheets.
        return ""
    if any(
        marker in lower
        for marker in (
            "phase_b", "phase-b", "_phase_b", "objective_routing", "debug:",
            "internal qa", "implementation hint", "trigger entry",
        )
    ):
        # Internal QA / implementation-detail wording is not actionable for
        # a thesis reader; keep it out of the main Word/PDF report. The raw
        # text remains available in the Excel audit sheets.
        return ""
    if not re.search(r"[.!?]$", text):
        text += "."
    return text


def _unavailable_warning(item: Any) -> str:
    if isinstance(item, dict):
        text = f"{_text(item.get('title') or item.get('analysis') or 'Analysis')}: {_text(item.get('reason'))}"
    else:
        text = _text(item)
    return _user_facing_warning(text)


def _tested_associations_table(blueprint: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    rows = [row for row in blueprint.get("tested_associations") or [] if isinstance(row, dict)]
    if not rows:
        return None
    return {
        "title": "Summary of tested associations",
        "columns": [
            "Predictor", "Test applied", "Test statistic", "p-value",
            "Adjusted p-value", "Effect size", "Significance status", "Notes",
        ],
        "rows": [[
            row.get("predictor") or "",
            row.get("test_applied") or "",
            row.get("test_statistic") or "-",
            row.get("p_value") or "-",
            row.get("adjusted_p_value") or "-",
            row.get("effect_size") or "-",
            row.get("significance_status") or "",
            row.get("notes_warnings") or "-",
        ] for row in rows],
        "interpretation": (
            "This table reports every completed predictor-versus-outcome association, including "
            "non-significant and nominally significant results."
        ),
        "warnings": [],
    }


def _apply_polish_overrides(
    blueprint: Dict[str, Any],
    overrides: Dict[str, str],
) -> Dict[str, Any]:
    """Return a deep copy of blueprint with prose overrides injected into interpretation fields.

    Tables also get a ``_polish_interpretation`` key so the override survives
    through ``_descriptive_export_table``, which regenerates ``interpretation``.
    """
    if not overrides:
        return blueprint
    bp = deepcopy(blueprint)
    if "results_synthesis" in overrides:
        bp["results_synthesis"] = overrides["results_synthesis"]
    for section in bp.get("analysis_sections") or []:
        section_id = str(section.get("section_id") or "")
        s_key = f"section:{section_id}"
        if s_key in overrides:
            section["interpretation"] = overrides[s_key]
        for table in section.get("tables") or []:
            t_key = f"table:{table.get('table_id', '')}"
            if t_key in overrides:
                table["interpretation"] = overrides[t_key]
                table["_polish_interpretation"] = overrides[t_key]
        for fig in section.get("figures") or []:
            f_key = f"figure:{fig.get('figure_id', '')}"
            if f_key in overrides:
                fig["interpretation"] = overrides[f_key]
    return bp


def _hydrate_blueprint_result_warnings(
    blueprint: Dict[str, Any], results: Dict[str, Any]
) -> Dict[str, Any]:
    """Restore executed-result cautions that compact blueprint tables omit."""
    warnings_by_id: Dict[str, List[str]] = {}
    for result in results.get("tests") or results.get("results") or []:
        result_id = _text(result.get("id") or result.get("result_id"))
        if not result_id:
            continue
        warnings: List[str] = []
        for key in ("warning", "note"):
            value = _text(result.get(key))
            if value and value not in {"-", "None"}:
                warnings.append(value)
        warnings.extend(_text(item) for item in result.get("warnings") or [] if _text(item))
        for table in result.get("tables") or []:
            headers, rows = _table_rows(table)
            warning_indexes = [
                idx for idx, header in enumerate(headers)
                if header.strip().lower() in {"warning", "warnings", "notes/warnings"}
            ]
            for row in rows:
                for idx in warning_indexes:
                    if idx < len(row) and _text(row[idx]) not in {"", "-", "None"}:
                        warnings.append(_text(row[idx]))
        if warnings:
            warnings_by_id[result_id] = list(dict.fromkeys(warnings))

    if not warnings_by_id:
        return blueprint
    hydrated = deepcopy(blueprint)
    for section in hydrated.get("analysis_sections") or []:
        for table in section.get("tables") or []:
            inherited: List[str] = list(table.get("warnings") or [])
            for source_id in table.get("source_test_ids") or []:
                inherited.extend(warnings_by_id.get(str(source_id), []))
            if inherited:
                table["warnings"] = list(dict.fromkeys(inherited))
    return hydrated


def _section_intro(
    section: Dict[str, Any],
    label_ctx: Dict[str, Any],
    overrides: Optional[Dict[str, str]] = None,
) -> str:
    section_id = str(section.get("section_id") or "")
    if overrides:
        s_key = f"section:{section_id}"
        if s_key in overrides:
            return overrides[s_key]
    marker = _text(label_ctx.get("display") or "the primary outcome")
    intros = {
        "baseline_characteristics": "This section describes the baseline profile of the analysed sample.",
        "clinical_study_characteristics": "This section summarises clinically relevant pathology and prognostic variables.",
        "immunophenotype_characteristics": "This section summarises receptor, immunophenotypic, and molecular marker variables.",
        "marker_outcome_components": f"This section summarises {marker} and related marker-expression components.",
        "primary_outcome_distribution": f"This section reports the distribution of {marker}.",
        "bivariate_associations": f"This section presents statistical associations between eligible predictors and {marker}.",
    }
    return intros.get(section_id, _clean_interpretation(section.get("interpretation"), label_ctx))


def _ai_polish_status(requested: bool, applied: bool) -> str:
    if applied:
        return "applied"
    if requested:
        return "fallback"
    return "deterministic"


def _ai_polish_audit_label(requested: bool, applied: bool) -> str:
    from app.services import ai_narrative
    return ai_narrative.audit_label(_ai_polish_status(requested, applied))


def _add_warnings_docx(
    doc: Document, blueprint: Dict[str, Any], *, ai_polish_requested: bool = False, ai_polish_applied: bool = False
) -> None:
    warnings = list(blueprint.get("warnings") or [])
    unavailable = list(blueprint.get("unavailable_or_recommended_only") or [])
    if not warnings and not unavailable:
        _plain_docx_text(doc, "No major thesis-reporting cautions were recorded.")
        _plain_docx_text(doc, _ai_polish_audit_label(ai_polish_requested, ai_polish_applied))
        return
    rendered = set()
    for warning in warnings:
        cleaned = _user_facing_warning(warning)
        if cleaned and cleaned not in rendered:
            if cleaned.startswith("Duplicate raw ") and "category labels (" in cleaned:
                doc.add_paragraph(style="List Bullet").add_run(_text_preserve_raw_labels(cleaned))
            else:
                _plain_docx_text(doc, cleaned, style="List Bullet")
            rendered.add(cleaned)
    for item in unavailable:
        cleaned = _unavailable_warning(item)
        if cleaned and cleaned not in rendered:
            _plain_docx_text(doc, cleaned, style="List Bullet")
            rendered.add(cleaned)
    _plain_docx_text(
        doc,
        "These are association analyses only; no causal, prognostic, or independent-effect conclusions should be drawn without an appropriate adjusted model and outcome data.",
    )
    _plain_docx_text(doc, _ai_polish_audit_label(ai_polish_requested, ai_polish_applied))


def generate_docx(
    results: Dict[str, Any],
    *,
    include_optional_figures: bool = False,
    polish_overrides: Optional[Dict[str, str]] = None,
    ai_polish_requested: bool = False,
) -> bytes:
    blueprint = _hydrate_blueprint_result_warnings(_blueprint(results), results)
    if polish_overrides:
        blueprint = _apply_polish_overrides(blueprint, polish_overrides)
    label_ctx = _outcome_label_context(blueprint)
    doc = Document()
    _set_docx_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CHAPTER V\nOBSERVATION AND RESULTS")
    run.bold = True
    run.font.size = Pt(14)

    _add_heading(doc, "5.1 Study Summary", 1)
    _add_summary_docx(doc, blueprint, results)

    _add_heading(doc, "5.2 Statistical Methods", 1)
    methods = _sanitize_thesis_claims(_format_statistical_display(blueprint.get("methods_text")))
    _plain_docx_text(doc, methods or "Statistical methods were generated from the executed Sigma analysis plan.")

    table_no = 1
    figure_no = 1
    main_sections = _main_sections(blueprint)

    baseline_sections = _dedupe_descriptive_sections([
        section for section in main_sections
        if section.get("section_id") in {
            "baseline_characteristics",
            "clinical_study_characteristics",
            "immunophenotype_characteristics",
            "marker_outcome_components",
            "descriptive_results",
        }
    ], label_ctx)
    if not baseline_sections:
        _plain_docx_text(doc, "Baseline and study characteristics were not available in the blueprint.")
    for section in [s for s in baseline_sections if s.get("section_id") not in {"marker_outcome_components"}]:
        _add_heading(doc, _section_export_title(str(section.get("section_id") or ""), section.get("title"), label_ctx), 1)
        intro = _section_intro(section, label_ctx, polish_overrides)
        if intro:
            _plain_docx_text(doc, intro)
        table_no, figure_no = _render_section_docx(
            doc, section, table_no, figure_no, label_ctx, include_optional_figures
        )

    _add_heading(doc, _section_export_title("primary_outcome_distribution", "Marker Expression", label_ctx), 1)
    outcome_sections = [section for section in main_sections if section.get("section_id") == "primary_outcome_distribution"]
    marker_sections = [section for section in baseline_sections if section.get("section_id") == "marker_outcome_components"]
    if outcome_sections:
        for section in outcome_sections:
            intro = _section_intro(section, label_ctx, polish_overrides)
            if intro:
                _plain_docx_text(doc, intro)
            outcome_tables = _section_tables(section) or _primary_outcome_tables(blueprint, label_ctx)
            display_tables = _expand_export_tables(outcome_tables, label_ctx)
            for table in display_tables:
                table_no = _add_table_docx(doc, table, table_no, label_ctx, include_interpretation=False)
            rendered_outcome_figure = False
            if outcome_tables:
                generated = _primary_outcome_figure(blueprint, outcome_tables[0], label_ctx)
                if generated:
                    next_no = _add_figure_docx(doc, generated, figure_no, label_ctx)
                    rendered_outcome_figure = rendered_outcome_figure or next_no != figure_no
                    figure_no = next_no
            for fig in _section_figures(section, include_optional_figures)[:1]:
                if _text(fig.get("figure_id")) == "primary_outcome_distribution_generated":
                    continue
                next_no = _add_figure_docx(doc, _normalise_figure_metadata(fig, label_ctx), figure_no, label_ctx)
                rendered_outcome_figure = rendered_outcome_figure or next_no != figure_no
                figure_no = next_no
            for table in display_tables:
                interpretation = _clean_interpretation(table.get("interpretation"), label_ctx)
                if interpretation:
                    _plain_docx_text(doc, interpretation)
    else:
        _plain_docx_text(doc, "Primary outcome distribution was not available in the blueprint.")
    for section in marker_sections:
        table_no, figure_no = _render_section_docx(
            doc, section, table_no, figure_no, label_ctx, include_optional_figures
        )

    inferential_ids = {
        "bivariate_associations", "correlation_analysis", "regression_adjusted_analysis",
        "diagnostic_accuracy", "reliability_agreement", "survival_analysis",
        "repeated_measures", "other_analyses",
    }
    inferential_sections = [section for section in main_sections if section.get("section_id") in inferential_ids]
    if not inferential_sections:
        _add_heading(doc, "Section V - Statistical Associations", 1)
        _plain_docx_text(doc, "No thesis-ready inferential analysis tables were available.")
    for section in inferential_sections:
        _add_heading(doc, _section_export_title(str(section.get("section_id") or ""), section.get("title"), label_ctx), 1)
        intro = _section_intro(section, label_ctx, polish_overrides)
        if intro:
            _plain_docx_text(doc, intro)
        table_no, figure_no = _render_section_docx(
            doc, section, table_no, figure_no, label_ctx, include_optional_figures
        )

    if inferential_sections:
        _plain_docx_text(doc, _percentage_denominator_note(blueprint))

    _add_heading(doc, "Section VI - Summary of Tested Associations", 1)
    association_summary = _tested_associations_table(blueprint)
    if association_summary:
        table_no = _add_table_docx(doc, association_summary, table_no, label_ctx)
    else:
        _plain_docx_text(doc, "No completed predictor-versus-outcome associations were available.")
    findings = list(blueprint.get("significant_findings") or [])
    if findings:
        _add_heading(doc, "Significant Findings Highlight", 2)
        sig_table = {
            "title": "Final thesis significant findings",
            "columns": [
                "Variable / parameter", "Key finding", "Test statistic", "p-value",
                "Adjusted p-value", "Test applied", "Effect size", "Notes/warnings",
            ],
            "rows": [
                [
                    _display_value(row.get("variable") or "", label_ctx),
                    _clean_finding_text(row.get("key_finding") or ""),
                    row.get("test_statistic") or "",
                    row.get("p_value") or "",
                    row.get("adjusted_p_value") or "",
                    row.get("test_applied") or "",
                    row.get("effect_size") or "",
                    row.get("notes_warnings") or "",
                ]
                for row in findings if isinstance(row, dict)
            ],
            "interpretation": "These findings are filtered for thesis-facing reporting and exclude marker/outcome components by default.",
            "warnings": [],
        }
        table_no = _add_table_docx(doc, sig_table, table_no, label_ctx, keep_notes_column=True)
    else:
        _plain_docx_text(doc, "No final thesis significant findings were identified among the completed tests.")

    results_synthesis = _clean_interpretation(blueprint.get("results_synthesis"), label_ctx)
    if results_synthesis:
        _add_heading(doc, "Results Synthesis", 2)
        _plain_docx_text(doc, results_synthesis)

    _add_heading(doc, "Warnings and Interpretation Notes", 1)
    _add_warnings_docx(
        doc, blueprint,
        ai_polish_requested=ai_polish_requested or bool(polish_overrides),
        ai_polish_applied=bool(polish_overrides),
    )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _pdf_escape(value: Any) -> str:
    return _text(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_table(
    payload: Dict[str, Any],
    label_ctx: Optional[Dict[str, Any]] = None,
    width: float = 10.4 * inch,
    keep_notes_column: bool = False,
) -> Tuple[Table, List[str]]:
    payload = _normalise_table_for_export(payload, label_ctx or {})
    payload, warning_notes = _compact_table_for_export(payload, keep_notes_column=keep_notes_column)
    headers, rows = _table_rows(payload)
    font_size = 6 if len(headers) > 6 else 7
    style = ParagraphStyle("Cell", fontName="Helvetica", fontSize=font_size, leading=font_size + 2)
    header_style = ParagraphStyle("CellHeader", parent=style, fontName="Helvetica-Bold")
    data = [[Paragraph(_pdf_escape(_display_value(h, label_ctx)), header_style) for h in headers]] + [
        [Paragraph(_pdf_escape(_display_value(cell, label_ctx)), style) for cell in row] for row in rows
    ]
    col_width = width / max(len(headers), 1)
    table = Table(data, colWidths=[col_width] * max(len(headers), 1), repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E9EEF7")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return table, warning_notes


def _add_pdf_figure(flow: List[Any], figure: Dict[str, Any], figure_no: int, label_ctx: Dict[str, Any], body, small) -> int:
    png = _strip_data_uri(_text(figure.get("png_data_uri")))
    if not png:
        return figure_no
    try:
        flow.append(Image(io.BytesIO(png), width=5.5 * inch, height=3.4 * inch))
    except Exception:
        return figure_no
    flow.append(Paragraph(f"<i>Figure {figure_no}. {_pdf_escape(_display_value(figure.get('caption') or figure.get('title'), label_ctx))}</i>", small))
    interpretation = _clean_interpretation(figure.get("interpretation"), label_ctx)
    if interpretation:
        flow.append(Paragraph(_pdf_escape(interpretation), body))
    return figure_no + 1


def _render_section_pdf(
    flow: List[Any],
    section: Dict[str, Any],
    figure_no: int,
    label_ctx: Dict[str, Any],
    available_width: float,
    body,
    small,
    include_optional_figures: bool = False,
    table_no: int = 1,
    h2: Any = None,
) -> Tuple[int, int]:
    section_id = str(section.get("section_id") or "")
    figures = [
        _normalise_figure_metadata(figure, label_ctx)
        for figure in _section_figures(section, include_optional_figures)
        if include_optional_figures or _is_core_figure(figure, label_ctx, section_id)
    ]
    used_figures: Set[str] = set()
    tables = _expand_export_tables(_section_tables(section), label_ctx)
    figures = _enrich_figures_from_tables(figures, tables, label_ctx)
    if section_id == "bivariate_associations":
        for figure in sorted(figures, key=lambda item: _association_figure_sort_key(item, label_ctx)):
            figure_id = _text(figure.get("figure_id") or figure.get("title"))
            if figure_id in used_figures:
                continue
            next_no = _add_pdf_figure(flow, figure, figure_no, label_ctx, body, small)
            if next_no != figure_no:
                figure_no = next_no
                used_figures.add(figure_id)
        if tables and h2 is not None:
            flow.append(Paragraph("Detailed Association Tables", h2))
    for table in tables:
        matched_rendered_figure = section_id == "bivariate_associations" and any(
            _text(figure.get("figure_id") or figure.get("title")) in used_figures
            and _figure_matches_table(figure, table)
            for figure in figures
        )
        clean_title = _display_value(_clean_table_title(table.get('title')), label_ctx)
        flow.append(Paragraph(f"<b>{_pdf_escape(f'Table {table_no}. {clean_title}')}</b>", small))
        table_no += 1
        pdf_table, warning_notes = _pdf_table(table, label_ctx, available_width)
        flow.append(pdf_table)
        if not matched_rendered_figure:
            for warning in warning_notes:
                cleaned_warning = _clean_interpretation(warning, label_ctx)
                if cleaned_warning:
                    flow.append(Paragraph(_pdf_escape(cleaned_warning), small))
        for figure in figures:
            figure_id = _text(figure.get("figure_id") or figure.get("title"))
            if figure_id in used_figures or not _figure_matches_table(figure, table):
                continue
            next_no = _add_pdf_figure(flow, figure, figure_no, label_ctx, body, small)
            if next_no != figure_no:
                figure_no = next_no
                used_figures.add(figure_id)
        if not matched_rendered_figure:
            enriched = _normalise_table_for_export(table, label_ctx)
            interpretation = _clean_interpretation(enriched.get("interpretation"), label_ctx)
            if interpretation:
                flow.append(Paragraph(_pdf_escape(interpretation), body))
        flow.append(Spacer(1, 6))
    if include_optional_figures:
        for figure in figures[:4]:
            figure_id = _text(figure.get("figure_id") or figure.get("title"))
            if figure_id in used_figures:
                continue
            next_no = _add_pdf_figure(flow, figure, figure_no, label_ctx, body, small)
            if next_no != figure_no:
                figure_no = next_no
                used_figures.add(figure_id)
    return figure_no, table_no


def generate_pdf(
    results: Dict[str, Any],
    *,
    include_optional_figures: bool = False,
    polish_overrides: Optional[Dict[str, str]] = None,
    ai_polish_requested: bool = False,
) -> bytes:
    blueprint = _hydrate_blueprint_result_warnings(_blueprint(results), results)
    if polish_overrides:
        blueprint = _apply_polish_overrides(blueprint, polish_overrides)
    label_ctx = _outcome_label_context(blueprint)
    out = io.BytesIO()
    page_size = landscape(A4)
    available_width = page_size[0] - 72
    doc = SimpleDocTemplate(out, pagesize=page_size, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]
    small = ParagraphStyle("Small", parent=body, fontSize=8, leading=10)
    flow: List[Any] = []
    figure_no = 1
    table_no = 1

    flow.append(Paragraph("CHAPTER V<br/>OBSERVATION AND RESULTS", styles["Title"]))
    flow.append(Spacer(1, 10))
    flow.append(Paragraph("5.1 Study Summary", h1))
    for label, value in _study_summary(blueprint, results):
        flow.append(Paragraph(f"<b>{_pdf_escape(label)}:</b> {_pdf_escape(value)}", body))
    flow.append(Paragraph("5.2 Statistical Methods", h1))
    methods = _sanitize_thesis_claims(_format_statistical_display(blueprint.get("methods_text")))
    flow.append(Paragraph(_pdf_escape(methods or "Statistical methods were generated from the executed Sigma analysis plan."), body))

    main_sections = _main_sections(blueprint)
    descriptive_sections = _dedupe_descriptive_sections([
        section for section in main_sections
        if section.get("section_id") in {"baseline_characteristics", "clinical_study_characteristics", "immunophenotype_characteristics", "marker_outcome_components", "descriptive_results"}
    ], label_ctx)
    for section in [s for s in descriptive_sections if s.get("section_id") != "marker_outcome_components"]:
        flow.append(Paragraph(_pdf_escape(_section_export_title(str(section.get("section_id") or ""), section.get("title"), label_ctx)), h1))
        intro = _section_intro(section, label_ctx, polish_overrides)
        if intro:
            flow.append(Paragraph(_pdf_escape(intro), body))
        figure_no, table_no = _render_section_pdf(
            flow, section, figure_no, label_ctx, available_width, body, small, include_optional_figures,
            table_no=table_no,
        )

    flow.append(Paragraph(_pdf_escape(_section_export_title("primary_outcome_distribution", "Marker Expression", label_ctx)), h1))
    outcome_sections = [section for section in main_sections if section.get("section_id") == "primary_outcome_distribution"]
    for section in outcome_sections:
        intro = _section_intro(section, label_ctx, polish_overrides)
        if intro:
            flow.append(Paragraph(_pdf_escape(intro), body))
        outcome_tables = _primary_outcome_tables(blueprint, label_ctx) or _section_tables(section)
        display_tables = _expand_export_tables(outcome_tables, label_ctx)
        for table in display_tables:
            clean_title = _display_value(_clean_table_title(table.get('title')), label_ctx)
            flow.append(Paragraph(f"<b>{_pdf_escape(f'Table {table_no}. {clean_title}')}</b>", small))
            table_no += 1
            pdf_table, warning_notes = _pdf_table(table, label_ctx, available_width)
            flow.append(pdf_table)
            for warning in warning_notes:
                cleaned_warning = _clean_interpretation(warning, label_ctx)
                if cleaned_warning:
                    flow.append(Paragraph(_pdf_escape(cleaned_warning), small))
        rendered_outcome_figure = False
        if outcome_tables:
            generated = _primary_outcome_figure(blueprint, outcome_tables[0], label_ctx)
            if generated:
                next_no = _add_pdf_figure(flow, generated, figure_no, label_ctx, body, small)
                rendered_outcome_figure = rendered_outcome_figure or next_no != figure_no
                figure_no = next_no
        for fig in _section_figures(section, include_optional_figures)[:1]:
            if _text(fig.get("figure_id")) == "primary_outcome_distribution_generated":
                continue
            next_no = _add_pdf_figure(flow, _normalise_figure_metadata(fig, label_ctx), figure_no, label_ctx, body, small)
            rendered_outcome_figure = rendered_outcome_figure or next_no != figure_no
            figure_no = next_no
        for table in display_tables:
            interpretation = _clean_interpretation(table.get("interpretation"), label_ctx)
            if interpretation:
                flow.append(Paragraph(_pdf_escape(interpretation), body))
    for section in [s for s in descriptive_sections if s.get("section_id") == "marker_outcome_components"]:
        figure_no, table_no = _render_section_pdf(
            flow, section, figure_no, label_ctx, available_width, body, small, include_optional_figures,
            table_no=table_no,
        )

    inferential_ids = {"bivariate_associations", "correlation_analysis", "regression_adjusted_analysis", "diagnostic_accuracy", "reliability_agreement", "survival_analysis", "repeated_measures", "other_analyses"}
    pdf_inferential_sections = [section for section in main_sections if section.get("section_id") in inferential_ids]
    for section in pdf_inferential_sections:
        flow.append(Paragraph(_pdf_escape(_section_export_title(str(section.get("section_id") or ""), section.get("title"), label_ctx)), h1))
        intro = _section_intro(section, label_ctx, polish_overrides)
        if intro:
            flow.append(Paragraph(_pdf_escape(intro), body))
        figure_no, table_no = _render_section_pdf(
            flow, section, figure_no, label_ctx, available_width, body, small, include_optional_figures,
            table_no=table_no, h2=h2,
        )
    if pdf_inferential_sections:
        flow.append(Paragraph(_pdf_escape(_percentage_denominator_note(blueprint)), small))
    flow.append(Paragraph("Section VI - Summary of Tested Associations", h1))
    association_summary = _tested_associations_table(blueprint)
    if association_summary:
        flow.append(Paragraph(f"<b>{_pdf_escape(f'Table {table_no}. Summary of tested associations')}</b>", small))
        table_no += 1
        pdf_table, warning_notes = _pdf_table(association_summary, label_ctx, available_width)
        flow.append(pdf_table)
        for warning in warning_notes:
            cleaned_warning = _clean_interpretation(warning, label_ctx)
            if cleaned_warning:
                flow.append(Paragraph(_pdf_escape(cleaned_warning), small))
    else:
        flow.append(Paragraph("No completed predictor-versus-outcome associations were available.", body))
    if blueprint.get("significant_findings"):
        flow.append(Paragraph("Significant Findings Highlight", h2))
        flow.append(Paragraph(f"<b>{_pdf_escape(f'Table {table_no}. Final thesis significant findings')}</b>", small))
        table_no += 1
        pdf_table, warning_notes = _pdf_table({
            "columns": [
                "Variable / parameter", "Key finding", "Test statistic", "p-value",
                "Adjusted p-value", "Test applied", "Effect size", "Notes/warnings",
            ],
            "rows": [
                [
                    _display_value(row.get("variable") or "", label_ctx),
                    _clean_finding_text(row.get("key_finding") or ""),
                    row.get("test_statistic") or "",
                    row.get("p_value") or "",
                    row.get("adjusted_p_value") or "",
                    row.get("test_applied") or "",
                    row.get("effect_size") or "",
                    row.get("notes_warnings") or "",
                ]
                for row in blueprint.get("significant_findings") or []
            ],
        }, label_ctx, available_width, keep_notes_column=True)
        flow.append(pdf_table)
        for warning in warning_notes:
            cleaned_warning = _clean_interpretation(warning, label_ctx)
            if cleaned_warning:
                flow.append(Paragraph(_pdf_escape(cleaned_warning), small))
    else:
        flow.append(Paragraph("No final thesis significant findings were identified among the completed tests.", body))
    results_synthesis = _clean_interpretation(blueprint.get("results_synthesis"), label_ctx)
    if results_synthesis:
        flow.append(Paragraph("Results Synthesis", h2))
        flow.append(Paragraph(_pdf_escape(results_synthesis), body))
    flow.append(Paragraph("Warnings and Interpretation Notes", h1))
    warnings = list(blueprint.get("warnings") or [])
    unavailable = list(blueprint.get("unavailable_or_recommended_only") or [])
    if not warnings and not unavailable:
        flow.append(Paragraph("No major thesis-reporting cautions were recorded.", body))
    rendered_warnings = set()
    for warning in warnings:
        cleaned = _user_facing_warning(warning)
        if cleaned and cleaned not in rendered_warnings:
            if cleaned.startswith("Duplicate raw ") and "category labels (" in cleaned:
                escaped = _text_preserve_raw_labels(cleaned).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                flow.append(Paragraph(f"• {escaped}", body))
            else:
                flow.append(Paragraph(f"• {_pdf_escape(cleaned)}", body))
            rendered_warnings.add(cleaned)
    for item in unavailable:
        cleaned = _unavailable_warning(item)
        if cleaned and cleaned not in rendered_warnings:
            flow.append(Paragraph(f"• {_pdf_escape(cleaned)}", body))
            rendered_warnings.add(cleaned)
    flow.append(Paragraph("These are association analyses only; no causal, prognostic, or independent-effect conclusions should be drawn without an appropriate adjusted model and outcome data.", body))
    flow.append(Paragraph(
        _ai_polish_audit_label(ai_polish_requested or bool(polish_overrides), bool(polish_overrides)),
        body,
    ))
    doc.build(flow)
    return out.getvalue()
