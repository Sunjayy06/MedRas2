"""Build a blank standard Indian MD / MS thesis proforma DOCX.

The template follows the conventional 14-section format used across NBEMS
(MD / MS / DNB) programmes.  All field lines are left blank for the
researcher to fill by hand or on-screen.
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _set_spacing(p, before: int = 0, after: int = 2) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _normal(doc: Document, text: str = "", bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _section_head(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_spacing(p, before=8, after=2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _field(doc: Document, label: str, filler: str = "") -> None:
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2)
    lr = p.add_run(label + ": ")
    lr.bold = True
    lr.font.size = Pt(12)
    vr = p.add_run(filler if filler else ("_" * 55))
    vr.font.size = Pt(12)


def _blank_lines(doc: Document, label: str = "", count: int = 4) -> None:
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2)
    if label:
        lr = p.add_run(label + ":  ")
        lr.bold = True
        lr.font.size = Pt(12)
    p.add_run(("_" * 80 + "\n") * count).font.size = Pt(12)


def _inline(doc: Document, *pairs: tuple[str, str]) -> None:
    """Render multiple label: _______ pairs on one paragraph line."""
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2)
    for label, filler in pairs:
        lr = p.add_run(label + ": ")
        lr.bold = True
        lr.font.size = Pt(12)
        p.add_run((filler if filler else ("_" * 20)) + "    ").font.size = Pt(12)


def _spacer(doc: Document) -> None:
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=0)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_proforma_docx() -> BytesIO:
    """Return a BytesIO containing a blank proforma DOCX."""

    doc = Document()

    # ── Page setup: A4, 1-inch margins ─────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1.0))

    # ── Default style: Times New Roman 12 pt ───────────────────────────────
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    # ── Title block ─────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after = Pt(4)
    tr = t.add_run("PROFORMA")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(14)
    sr = sub.add_run("(One proforma per patient — fill all fields legibly)")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.name = "Times New Roman"

    # ── A. Patient Identification ───────────────────────────────────────────
    _section_head(doc, "A.  PATIENT IDENTIFICATION")
    _inline(doc,
            ("Serial No.", "_______"),
            ("Date of Admission", "_______________"),
            ("Date of Discharge", "_______________"))
    _field(doc, "Name")
    _inline(doc, ("Age", "_____ yrs"), ("Sex", "Male / Female / Other"))
    _field(doc, "I.P. / O.P. No.")
    _inline(doc, ("Ward", "_______________________"), ("Bed No.", "_____________"))
    _field(doc, "Occupation")
    _inline(doc, ("Religion", "_______________________"), ("Caste", "_______________________"))
    _field(doc, "Address")
    p = doc.add_paragraph()
    _set_spacing(p, before=1, after=2)
    p.add_run("_" * 80).font.size = Pt(12)
    _inline(doc, ("Contact No.", "___________________________"))
    _spacer(doc)

    # ── B. Chief Complaints ─────────────────────────────────────────────────
    _section_head(doc, "B.  CHIEF COMPLAINTS WITH DURATION")
    for i in range(1, 6):
        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=2)
        p.add_run(f"{i}.  ").font.size = Pt(12)
        p.add_run("_" * 58 + "  since  " + "_" * 14).font.size = Pt(12)
    _spacer(doc)

    # ── C. History of Presenting Illness ────────────────────────────────────
    _section_head(doc, "C.  HISTORY OF PRESENTING ILLNESS")
    for _ in range(6):
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("_" * 90).font.size = Pt(12)
    _spacer(doc)

    # ── D. Past History ─────────────────────────────────────────────────────
    _section_head(doc, "D.  PAST HISTORY")
    for label in ["Medical", "Surgical", "Drug history", "Drug allergy"]:
        _field(doc, label)
    _spacer(doc)

    # ── E. Personal History ─────────────────────────────────────────────────
    _section_head(doc, "E.  PERSONAL HISTORY")
    _field(doc, "Diet", "Vegetarian  /  Non-vegetarian")
    _inline(doc,
            ("Smoking", "Yes / No"),
            ("Alcohol", "Yes / No"),
            ("Tobacco / Gutkha", "Yes / No"))
    _field(doc, "Menstrual history (female)", "(LMP / regularity / flow / dysmenorrhoea)")
    _field(doc, "Obstetric history (female)", "G ___ P ___ L ___ A ___")
    _spacer(doc)

    # ── F. Family History ───────────────────────────────────────────────────
    _section_head(doc, "F.  FAMILY HISTORY")
    for _ in range(3):
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("_" * 90).font.size = Pt(12)
    _spacer(doc)

    # ── G. General Examination ──────────────────────────────────────────────
    _section_head(doc, "G.  GENERAL EXAMINATION")
    _field(doc, "Built and nourishment")
    _inline(doc, ("Pallor", "___"), ("Icterus", "___"), ("Cyanosis", "___"), ("Clubbing", "___"))
    _inline(doc, ("Lymphadenopathy", "_____________"), ("Pedal oedema", "_____________"), ("JVP", "_______"))
    _inline(doc, ("Temperature", "_______°F"), ("SpO₂", "_______%"))
    _inline(doc,
            ("Pulse", "_____ /min"),
            ("Blood Pressure", "_____/_____ mmHg"),
            ("Respiratory Rate", "_____ /min"))
    _spacer(doc)

    # ── H. Systemic Examination ─────────────────────────────────────────────
    _section_head(doc, "H.  SYSTEMIC EXAMINATION")
    systems = [
        ("Cardiovascular System",
         "S\u2081 S\u2082: _________________   Murmurs: _________________   JVP: _______________"),
        ("Respiratory System",
         "Air entry: Bilateral equal / Unequal   Added sounds: _____________________________"),
        ("Per Abdomen",
         "Tenderness: ____________   Guarding: ___________   Organomegaly: ________________"),
        ("Central Nervous System",
         "Oriented: Yes / No   GCS: E___ V___ M___   Focal deficit: _____________________"),
        ("Other systems / Local examination",
         "_____________________________________________________________________________"),
    ]
    for system, detail in systems:
        p = doc.add_paragraph()
        _set_spacing(p, before=3, after=2)
        lr = p.add_run(system + ":  ")
        lr.bold = True
        lr.font.size = Pt(12)
        p.add_run(detail).font.size = Pt(11)
    _spacer(doc)

    # ── I. Investigations ───────────────────────────────────────────────────
    _section_head(doc, "I.  INVESTIGATIONS")

    p = doc.add_paragraph()
    _set_spacing(p, before=3, after=1)
    p.add_run("Haematology:").bold = True
    p.runs[-1].font.size = Pt(12)
    for line in [
        "Hb: ______ g/dL    TLC: ______ cells/mm³    DLC: N___ L___ M___ E___ B___",
        "Platelets: ________________    ESR: ________________    Blood group: ________",
        "Peripheral smear: __________________________________________________________",
    ]:
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("  " + line).font.size = Pt(12)

    p = doc.add_paragraph()
    _set_spacing(p, before=5, after=1)
    p.add_run("Biochemistry:").bold = True
    p.runs[-1].font.size = Pt(12)
    for line in [
        "Blood glucose (F/PP/R): __________   BUN: __________   S. Creatinine: __________",
        "Serum Na\u207a: __________   Serum K\u207a: __________   Serum Cl\u207b: __________",
        "LFT \u2014 T. Bil: ____  D. Bil: ____  S. AST: ____  S. ALT: ____  Albumin: ____  T. Protein: ____",
        "Urine routine / microscopy: ________________________________________________",
    ]:
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("  " + line).font.size = Pt(12)

    p = doc.add_paragraph()
    _set_spacing(p, before=5, after=1)
    p.add_run("Radiology and Cardiology:").bold = True
    p.runs[-1].font.size = Pt(12)
    for label in ["ECG", "X-ray Chest (PA)", "2D Echocardiography", "USG Abdomen"]:
        _field(doc, "  " + label)

    p = doc.add_paragraph()
    _set_spacing(p, before=5, after=1)
    p.add_run("Special investigations (as per study protocol):").bold = True
    p.runs[-1].font.size = Pt(12)
    for _ in range(5):
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("  " + "_" * 85).font.size = Pt(12)
    _spacer(doc)

    # ── J. Diagnosis ────────────────────────────────────────────────────────
    _section_head(doc, "J.  DIAGNOSIS")
    _field(doc, "Provisional diagnosis")
    _field(doc, "Final diagnosis")
    _spacer(doc)

    # ── K. Treatment Given ──────────────────────────────────────────────────
    _section_head(doc, "K.  TREATMENT GIVEN")
    for _ in range(5):
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("_" * 90).font.size = Pt(12)
    _spacer(doc)

    # ── L. Course in Hospital ───────────────────────────────────────────────
    _section_head(doc, "L.  COURSE IN HOSPITAL")
    for _ in range(5):
        p = doc.add_paragraph()
        _set_spacing(p, before=1, after=1)
        p.add_run("_" * 90).font.size = Pt(12)
    _spacer(doc)

    # ── M. Outcome ──────────────────────────────────────────────────────────
    _section_head(doc, "M.  OUTCOME AT DISCHARGE")
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=2)
    p.add_run("\u2610 Improved    \u2610 Cured    \u2610 LAMA    \u2610 Referred    \u2610 Expired").font.size = Pt(12)
    _field(doc, "Condition at discharge")
    _inline(doc, ("Discharge date", "_______________________"))
    _spacer(doc)

    # ── N. Follow-up ────────────────────────────────────────────────────────
    _section_head(doc, "N.  FOLLOW-UP")
    for i in range(1, 4):
        p = doc.add_paragraph()
        _set_spacing(p, before=2, after=2)
        lr = p.add_run(f"Visit {i}:  ")
        lr.bold = True
        lr.font.size = Pt(12)
        p.add_run("Date: ___________________    Findings: " + "_" * 45).font.size = Pt(12)
    _spacer(doc)
    _spacer(doc)

    # Signature line (right-aligned)
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.space_before = Pt(20)
    sig.add_run("Signature of Investigator: _____________________________").font.size = Pt(12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
