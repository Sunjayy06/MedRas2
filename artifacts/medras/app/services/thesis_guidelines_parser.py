"""Parse a university thesis-guidelines PDF / DOCX and extract enforceable
rules: page limits, font, line spacing, margins, reference minimum,
citation style and required declarations.

Library-backed text extraction only (PyPDF2 / python-docx). No LLM — the
rules are matched via deterministic regex so the result is reproducible
and auditable.

Public surface
--------------
* ``parse_guidelines(filename, data)`` -> ``{rules, evidence}``
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

from app.services.thesis_formats import DEFAULT_RULES

log = logging.getLogger(__name__)

MAX_BYTES = 20 * 1024 * 1024  # 20 MB cap; guidelines docs are small


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_text(filename: str, data: bytes) -> str:
    """Return concatenated text from a PDF or DOCX. Empty string on failure."""
    name = (filename or "").lower()
    if len(data) > MAX_BYTES:
        raise ValueError(f"File too large (>{MAX_BYTES // (1024*1024)} MB).")
    try:
        if name.endswith(".pdf"):
            return _pdf_text(data)
        if name.endswith(".docx"):
            return _docx_text(data)
        if name.endswith(".txt") or name.endswith(".md"):
            return data.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001 — surface a user-friendly error upstream
        log.info("thesis_guidelines: text extraction failed: %s", exc)
        return ""
    return ""


def _pdf_text(data: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        from pypdf import PdfReader  # type: ignore[no-redef]
    reader = PdfReader(io.BytesIO(data))
    parts: List[str] = []
    for p in reader.pages[:200]:  # 200-page cap mirrors the upload module
        try:
            parts.append(p.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _docx_text(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Rule extraction — pure regex, no LLM
# ---------------------------------------------------------------------------

_NUM = r"(\d{1,4}(?:\.\d{1,2})?)"


def _first_match(text: str, patterns: List[str]) -> Optional[str]:
    for pat in patterns:
        m = re.search(pat, text, re.I)
        if m:
            return m.group(1)
    return None


def _extract_rules(text: str) -> Tuple[Dict[str, Any], List[str]]:
    """Return ``(rules, evidence)``. ``evidence`` is a list of human-readable
    snippets explaining where each value came from — surfaced in the UI so
    the researcher can verify autofill before accepting it.
    """
    rules: Dict[str, Any] = {}
    evidence: List[str] = []

    # Page count — "size of 80 pages", "should not exceed 100 pages", etc.
    pages = _first_match(text, [
        rf"(?:size of|maximum of|not exceed|up to|restricted to)\s+{_NUM}\s+pages",
        rf"page\s+(?:limit|count)\s*(?:of|:)?\s*{_NUM}",
        rf"{_NUM}\s+pages?\s+(?:maximum|max\.?|limit)",
    ])
    if pages:
        rules["max_pages"] = int(float(pages))
        evidence.append(f"Page limit: {pages}")

    # Font family
    fam_match = re.search(
        r"\b(Times New Roman|Arial|Calibri|Garamond|Georgia|Helvetica|Cambria)\b",
        text, re.I)
    if fam_match:
        rules["font_family"] = fam_match.group(1).title()
        evidence.append(f"Font family: {rules['font_family']}")

    # Font size
    fs = _first_match(text, [
        rf"font\s*size\s*(?:of|:)?\s*{_NUM}\s*(?:pt|point)",
        rf"{_NUM}\s*(?:pt|point)\s+(?:font|type)",
        rf"size\s+{_NUM}\s+font",
    ])
    if fs:
        rules["font_size_pt"] = int(float(fs))
        evidence.append(f"Font size: {fs} pt")

    # Line spacing — "1.5 space", "double-spaced", "single spaced"
    if re.search(r"\bdouble[- ]spac", text, re.I):
        rules["line_spacing"] = 2.0
        evidence.append("Line spacing: double (2.0)")
    elif re.search(r"\bsingle[- ]spac", text, re.I):
        rules["line_spacing"] = 1.0
        evidence.append("Line spacing: single (1.0)")
    else:
        ls = _first_match(text, [
            rf"{_NUM}\s*(?:space|spacing|line spacing)",
            rf"line\s*spacing\s*(?:of|:)?\s*{_NUM}",
        ])
        if ls:
            try:
                rules["line_spacing"] = float(ls)
                evidence.append(f"Line spacing: {ls}")
            except ValueError:
                pass

    # Margins — "1\" margins", "2.5 cm margin"
    mm = re.search(rf"{_NUM}\s*(?:cm|inch|\"|'')\s+margin", text, re.I)
    if mm:
        val = float(mm.group(1))
        unit = mm.group(0).lower()
        inches = val if ("inch" in unit or '"' in unit or "''" in unit) else round(val / 2.54, 2)
        rules["margin_inches"] = inches
        evidence.append(f"Margin: {val} ({inches}\")")

    # Citation style
    style_lookup = [
        ("vancouver", r"\b(vancouver|ICMJE)\b"),
        ("apa",       r"\bAPA\b"),
        ("harvard",   r"\bHarvard\b"),
        ("chicago",   r"\bChicago\b"),
        ("ieee",      r"\bIEEE\b"),
        ("mla",       r"\bMLA\b"),
    ]
    for code, pat in style_lookup:
        if re.search(pat, text, re.I):
            rules["citation_style"] = code
            evidence.append(f"Citation style: {code.upper()}")
            break

    # References minimum
    refs = _first_match(text, [
        rf"(?:minimum|min\.?|at least)\s+{_NUM}\s+references?",
        rf"{_NUM}\s+references?\s+(?:minimum|min\.?|required)",
    ])
    if refs:
        rules["min_references"] = int(float(refs))
        evidence.append(f"Minimum references: {refs}")

    # Plagiarism cap
    plag = _first_match(text, [
        rf"plagiarism\s+(?:should\s+be\s+)?(?:less than|<|below|under|not exceed|max\.?)\s+{_NUM}\s*%",
        rf"similarity\s+index\s+(?:less than|<|below|under|not exceed|max\.?)\s+{_NUM}\s*%",
        rf"{_NUM}\s*%\s+plagiarism",
    ])
    if plag:
        rules["max_plagiarism_pct"] = float(plag)
        evidence.append(f"Plagiarism cap: {plag}%")

    # IEC / ethics committee approval mention
    if re.search(r"\b(IEC|institutional ethics|ethics committee)\b", text, re.I):
        rules["iec_required"] = True

    # Paper size
    if re.search(r"\bA4\b", text, re.I):
        rules["paper"] = "A4"
    elif re.search(r"\bletter\b", text, re.I):
        rules["paper"] = "Letter"

    return rules, evidence


def parse_guidelines(filename: str, data: bytes) -> Dict[str, Any]:
    """Parse an uploaded guidelines file and return autofilled rules.

    Returns
    -------
    ``{"rules": <merged dict>, "evidence": [...], "extracted": <only what we
    found in the file>, "defaults_used": [<rule names that fell back>]}``.
    The ``rules`` dict is always a complete set safe to use downstream
    because anything missing is filled from ``DEFAULT_RULES``.
    """
    text = _extract_text(filename, data)
    if not text.strip():
        return {
            "rules": dict(DEFAULT_RULES),
            "evidence": ["Could not read text from this file — using NBEMS defaults."],
            "extracted": {},
            "defaults_used": list(DEFAULT_RULES.keys()),
        }
    extracted, evidence = _extract_rules(text)
    merged = dict(DEFAULT_RULES)
    merged.update(extracted)
    defaults_used = [k for k in DEFAULT_RULES if k not in extracted]
    return {
        "rules": merged,
        "evidence": evidence,
        "extracted": extracted,
        "defaults_used": defaults_used,
    }
